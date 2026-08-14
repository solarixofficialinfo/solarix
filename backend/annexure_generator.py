"""
annexure_generator.py
---------------------
DOCX-template-based Annexure-I generator for Solarix.

Flow:
    onboarding/client data
        ↓
    resolve_annexure_values()
        ↓
    docx_template_engine.render_docx()  (placeholder → value substitution)
        ↓
    _convert_docx_to_pdf()              (LibreOffice headless)
        ↓
    PDF bytes returned to caller

If LibreOffice is not available, the filled DOCX bytes are returned instead and
the caller receives a DOCX content-type.

DO NOT TOUCH any other generator (WCR, SLDR, Net Meter, Vendor, etc.).
"""
from __future__ import annotations

import io
import logging
import os
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

from docx import Document

logger = logging.getLogger(__name__)

# ── Template location ──────────────────────────────────────────────────────────
_HERE = Path(__file__).parent
TEMPLATE_PATH = _HERE / "annexure_template.docx"


def _load_template_bytes() -> bytes:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Annexure DOCX template not found at {TEMPLATE_PATH}. "
            "Run backend/instrument_annexure_template.py to generate it."
        )
    return TEMPLATE_PATH.read_bytes()


# ── Value resolution ───────────────────────────────────────────────────────────

def _s(v: Any) -> str:
    """Safe string – never returns None."""
    if v is None:
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v).strip()


def _resolve_installation_date(client: dict) -> str:
    raw = (
        client.get("installation_date")
        or client.get("install_date")
        or client.get("commissioning_date")
    )
    if raw:
        raw = _s(raw)
        # Accept ISO date (YYYY-MM-DD) → reformat to DD/MM/YYYY
        if len(raw) >= 10 and raw[4] == "-":
            try:
                dt = datetime.strptime(raw[:10], "%Y-%m-%d")
                return dt.strftime("%d/%m/%Y")
            except ValueError:
                pass
        return raw
    # Fallback: today
    return datetime.now(timezone.utc).strftime("%d/%m/%Y")


def resolve_annexure_values(client: dict, company: dict) -> Dict[str, str]:
    """
    Return a flat dict mapping every placeholder name (without {{ }}) → resolved string.
    All values come from client or company data. No hardcoded defaults.
    If a field is empty, the placeholder is replaced with an empty string.
    """
    ob: dict = {}
    stages = client.get("stages") or {}
    if isinstance(stages, dict):
        ob = dict(stages.get("onboarding_data") or {})

    def _get(*keys) -> str:
        """Try multiple keys across client dict and onboarding_data."""
        for k in keys:
            v = client.get(k) or ob.get(k)
            if v:
                return _s(v)
        return ""

    # ── Basic client fields ────────────────────────────────────────────────────
    client_full_name  = _get("full_name", "name", "client_name")
    consumer_no       = _get("consumer_number", "consumer_no")
    mobile_no         = _get("mobile", "mobile_no")
    email_id          = _get("email", "email_id")

    # Address
    addr = _get("address")
    city = _get("city")
    state = _get("state")
    pincode = _get("pincode")
    address_parts = [addr]
    if city:    address_parts.append(city)
    if state:   address_parts.append(state)
    if pincode: address_parts.append(pincode)
    address = ", ".join(p for p in address_parts if p)

    # ── Fixed / enum values ───────────────────────────────────────────────────
    net_metering_arrangement = "Net Metering Arrangement"
    re_source                = "Solar"
    capacity_type            = "Rooftop"
    project_model            = "Capex"

    # ── Solar system ──────────────────────────────────────────────────────────
    system_kw = _get("system_kw", "capacity", "solar_capacity")
    panel_in_kw_kw = f"{system_kw} KW" if system_kw else ""

    panel_brand = _get("panel_brand", "panel_make")
    panel_tech  = _get("panel_technology", "panel_tech")
    panel_wp    = _get("panel_wattage", "panel_watt", "panel_wp")

    # Solar PV Details: "<Brand> <Tech> <Wattage> WP"
    pv_parts = [p for p in [panel_brand, panel_tech, (f"{panel_wp} WP" if panel_wp else "")] if p]
    solar_pv_details = " ".join(pv_parts) if pv_parts else ""

    panel_in_wp_wp = f"{panel_wp} WP" if panel_wp else ""

    num_panels = _get("num_panels", "number_of_panels", "panel_quantity")
    solar_panels_in_nos_nos = f"{num_panels} NOS" if num_panels else ""

    raw_inverters = client.get("inverters") or ob.get("inverters")
    inverter_list = []
    if isinstance(raw_inverters, list) and len(raw_inverters) > 0:
        for inv in raw_inverters:
            if isinstance(inv, dict):
                brand = _s(inv.get("brand") or inv.get("make") or "")
                model = _s(inv.get("model") or "")
                cap = _s(inv.get("capacity") or "")
                qty = _s(inv.get("quantity") or inv.get("qty") or "1")
                if brand or model or cap:
                    inverter_list.append({"brand": brand, "model": model, "capacity": cap, "quantity": qty})

    manual_inverter_kw = _get("inverter_capacity", "inverter_kw", "system_kw")
    inverter_in_kw_kw = f"{manual_inverter_kw} KW" if (manual_inverter_kw and "kw" not in manual_inverter_kw.lower()) else (manual_inverter_kw or "")

    unique_brands = []
    if inverter_list:
        for inv in inverter_list:
            b = _s(inv.get("brand") or inv.get("make") or "").strip()
            if b and b not in unique_brands:
                unique_brands.append(b)
    if not unique_brands:
        fb = _get("inverter_brand", "inverter_make")
        if fb:
            unique_brands.append(fb)

    inverter_brand = ", ".join(unique_brands)

    installation_date = _resolve_installation_date(client)

    # ── Company / vendor ──────────────────────────────────────────────────────
    company_name    = _s(company.get("company_name") or company.get("name") or company.get("legal_business_name") or company.get("vendor_name") or "")
    company_city    = _s(company.get("city") or "")
    company_details = f"{company_name}, {company_city}".strip(", ") if (company_name or company_city) else company_name

    return {
        # Exact placeholder names (without {{ }}) as written in annexure_template.docx
        "client_full_name":          client_full_name,
        "consumer_no":               consumer_no,
        "mobile_no":                 mobile_no,
        "email_id":                  email_id,
        "address":                   address,
        "net_metering_arrangement":  net_metering_arrangement,
        "re_source":                 re_source,
        "capacity_type":             capacity_type,
        "project_model":             project_model,
        "panel_in_kw_kw":            panel_in_kw_kw,
        "installation_date":         installation_date,
        "solar_pv_details":          solar_pv_details,
        "inverter_in_kw_kw":         inverter_in_kw_kw,
        "inverter_brand":            inverter_brand,
        "solar_panels_in_nos_nos":   solar_panels_in_nos_nos,
        "panel_in_wp_wp":            panel_in_wp_wp,
        "company_details":           company_details,
    }


# ── DOCX placeholder replacement ──────────────────────────────────────────────

import re as _re

_PH_RE = _re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def _replace_in_paragraph(paragraph, replacements: Dict[str, str]):
    """Replace {{ key }} → value in a paragraph while preserving run formatting."""
    runs = paragraph.runs
    if not runs:
        return
    full = "".join(r.text for r in runs)
    if "{{" not in full:
        return
    new_text = full
    for key, val in replacements.items():
        tokens = key.split()
        body = r"\s+".join(_re.escape(t) for t in tokens)
        pattern = _re.compile(r"\{\{\s*" + body + r"\s*\}\}")
        new_text = pattern.sub(lambda m: val, new_text)
    if new_text == full:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _walk_and_replace(container, replacements: Dict[str, str]):
    if hasattr(container, "paragraphs"):
        for p in container.paragraphs:
            _replace_in_paragraph(p, replacements)
    if hasattr(container, "tables"):
        for tbl in container.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    _walk_and_replace(cell, replacements)


def _fill_template(template_bytes: bytes, replacements: Dict[str, str]) -> bytes:
    doc = Document(io.BytesIO(template_bytes))
    _walk_and_replace(doc, replacements)
    for section in doc.sections:
        _walk_and_replace(section.header, replacements)
        _walk_and_replace(section.footer, replacements)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ── PDF conversion (LibreOffice headless) ─────────────────────────────────────

def _libreoffice_path() -> Optional[str]:
    for candidate in [
        "libreoffice", "soffice",
        "/usr/bin/libreoffice", "/usr/bin/soffice",
        "/usr/local/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]:
        if shutil.which(candidate):
            return candidate
        if os.path.isfile(candidate):
            return candidate
    return None


def _convert_docx_to_pdf(docx_bytes: bytes) -> Optional[bytes]:
    """Convert DOCX bytes → PDF bytes using LibreOffice headless. Returns None on failure."""
    lo = _libreoffice_path()
    if not lo:
        logger.warning("LibreOffice not found – returning DOCX instead of PDF.")
        return None
    with tempfile.TemporaryDirectory() as tmpdir:
        in_path = os.path.join(tmpdir, "annexure.docx")
        out_path = os.path.join(tmpdir, "annexure.pdf")
        with open(in_path, "wb") as f:
            f.write(docx_bytes)
        try:
            result = subprocess.run(
                [lo, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, in_path],
                capture_output=True, timeout=60,
            )
            if result.returncode != 0:
                logger.error(f"LibreOffice conversion failed: {result.stderr.decode()}")
                return None
            if os.path.exists(out_path):
                with open(out_path, "rb") as f:
                    return f.read()
            # LibreOffice sometimes names file differently
            for fn in os.listdir(tmpdir):
                if fn.endswith(".pdf"):
                    with open(os.path.join(tmpdir, fn), "rb") as f:
                        return f.read()
            logger.error("LibreOffice ran but no PDF output found.")
            return None
        except subprocess.TimeoutExpired:
            logger.error("LibreOffice conversion timed out.")
            return None
        except Exception as e:
            logger.error(f"LibreOffice conversion error: {e}")
            return None


def _convert_via_docx2pdf(docx_bytes: bytes) -> Optional[bytes]:
    """Try converting using the docx2pdf package (uses MS Word on macOS, LibreOffice on Linux)."""
    try:
        import docx2pdf  # type: ignore[import]
    except ImportError:
        return None
    with tempfile.TemporaryDirectory() as tmpdir:
        in_path = os.path.join(tmpdir, "annexure.docx")
        out_path = os.path.join(tmpdir, "annexure.pdf")
        with open(in_path, "wb") as f:
            f.write(docx_bytes)
        try:
            docx2pdf.convert(in_path, out_path)
            if os.path.exists(out_path):
                with open(out_path, "rb") as f:
                    return f.read()
        except Exception as e:
            logger.warning(f"docx2pdf conversion failed: {e}")
    return None


def _convert_via_unoconv(docx_bytes: bytes) -> Optional[bytes]:
    """Try converting using unoconv if installed."""
    unoconv = shutil.which("unoconv")
    if not unoconv:
        return None
    with tempfile.TemporaryDirectory() as tmpdir:
        in_path = os.path.join(tmpdir, "annexure.docx")
        out_path = os.path.join(tmpdir, "annexure.pdf")
        with open(in_path, "wb") as f:
            f.write(docx_bytes)
        try:
            result = subprocess.run(
                [unoconv, "-f", "pdf", "-o", out_path, in_path],
                capture_output=True, timeout=60,
            )
            if os.path.exists(out_path):
                with open(out_path, "rb") as f:
                    return f.read()
        except Exception as e:
            logger.warning(f"unoconv conversion failed: {e}")
    return None


def _docx_to_pdf(docx_bytes: bytes) -> Optional[bytes]:
    """Try headless DOCX → PDF conversion methods (LibreOffice CLI / unoconv)."""
    # 1. LibreOffice headless (best on Linux servers)
    result = _convert_docx_to_pdf(docx_bytes)
    if result:
        return result
    # 2. unoconv
    result = _convert_via_unoconv(docx_bytes)
    if result:
        return result
    return None


def _generate_annexure_pdf_reportlab(replacements: Dict[str, str], company: dict) -> bytes:
    """Generate high-fidelity MSEDCL Annexure-I PDF directly using ReportLab as fallback when LibreOffice is absent."""
    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )
    styles = getSampleStyleSheet()

    title_center_bold = ParagraphStyle(
        'ann_title_center_bold',
        parent=styles['Normal'],
        fontSize=12,
        leading=15,
        alignment=1,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#0f172a')
    )
    title_center = ParagraphStyle(
        'ann_title_center',
        parent=styles['Normal'],
        fontSize=10,
        leading=13,
        alignment=1,
        fontName='Helvetica',
        textColor=colors.HexColor('#1e293b')
    )
    cell_style = ParagraphStyle(
        'ann_cell_style',
        parent=styles['Normal'],
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor('#0f172a')
    )
    cell_bold = ParagraphStyle(
        'ann_cell_bold',
        parent=styles['Normal'],
        fontSize=8.5,
        leading=11,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#0f172a')
    )

    story = []

    # Header matching MSEDCL Annexure-I
    story.append(Paragraph("Maharashtra State Electricity", title_center_bold))
    story.append(Paragraph("Distribution Company Limited", title_center_bold))
    story.append(Spacer(1, 4))
    story.append(Paragraph("Renewable Energy Generating System", title_center))
    story.append(Paragraph("Annexure-I", title_center_bold))
    story.append(Paragraph("(Commissioning Report for RE System)", title_center))
    story.append(Spacer(1, 8))

    # Table matching annexure_template.docx Table 0
    headers = [
        Paragraph("<b>S/N</b>", cell_bold),
        Paragraph("<b>Particulars</b>", cell_bold),
        Paragraph("<b>As Commissioned</b>", cell_bold)
    ]

    table_rows = [headers]

    raw_rows = [
        ("1", "Name of Consumer", replacements.get("client_full_name", "")),
        ("2", "Consumer Number", replacements.get("consumer_no", "")),
        ("3", "Mobile Number", replacements.get("mobile_no", "")),
        ("4", "Email Id", replacements.get("email_id", "")),
        ("5", "Address of Installation", replacements.get("address", "")),
        ("6", "RE Arrangement Type", replacements.get("net_metering_arrangement", "")),
        ("7", "RE Source", replacements.get("re_source", "")),
        ("8", "", ""),
        ("9", "Capacity Type", replacements.get("capacity_type", "")),
        ("10", "Project Model", replacements.get("project_model", "")),
        ("11", "RE Installed Capacity (Rooftop) (Kw)", replacements.get("panel_in_kw_kw", "")),
        ("12", "RE Installed Capacity (Rooftop + Ground) (Kw)", "NA"),
        ("13", "RE Installed Capacity (Ground) (Kw)", "NA"),
        ("14", "Installation Date", replacements.get("installation_date", "")),
        ("15", "Solar PV Details", replacements.get("solar_pv_details", "")),
        ("16", "Inverter Capacity (Kw)", replacements.get("inverter_in_kw_kw", "")),
        ("17", "Inverter Make", replacements.get("inverter_brand", "")),
        ("18", "No. of PV Modules", replacements.get("solar_panels_in_nos_nos", "")),
        ("19", "Module Capacity (Kw)", replacements.get("panel_in_wp_wp", "")),
    ]

    for sn, part, val in raw_rows:
        table_rows.append([
            Paragraph(sn, cell_bold if sn else cell_style),
            Paragraph(part, cell_bold if part else cell_style),
            Paragraph(val, cell_style)
        ])

    table = Table(table_rows, colWidths=[1.2 * cm, 8.5 * cm, 8.3 * cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f1f5f9')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 8))

    # Notes section
    story.append(Paragraph("<b>To be uploaded separately:</b>", cell_bold))
    notes = [
        "Self-certification of safety of the installation of the RTS system along with the test report of the Licensed Electrical Contractor.",
        "Electrical Inspector permission if system is above 200 Kw. (Mandatory if system is above 200 Kw only)",
        "Third party leasing agreement if Model selected is RESCO.",
        "Photograph of the system commissioned."
    ]
    for n in notes:
        story.append(Paragraph(f"• {n}", cell_style))

    story.append(Spacer(1, 12))

    # Footer / Signatures
    inst_date = replacements.get("installation_date", "")
    client_name = replacements.get("client_full_name", "")
    comp_details = replacements.get("company_details", "")

    sig_data = [
        [
            Paragraph(f"<b>Date:</b> {inst_date}", cell_style),
            Paragraph(f"<b>Consumer:</b> {client_name}", cell_style),
            Paragraph(f"<b>Vendor / Company:</b> {comp_details}", cell_style)
        ]
    ]
    sig_table = Table(sig_data, colWidths=[4.0 * cm, 7.0 * cm, 7.0 * cm])
    sig_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(sig_table)

    doc.build(story)
    return buf.getvalue()


# ── Public API ─────────────────────────────────────────────────────────────────

def generate_annexure(client: dict, company: dict) -> Tuple[bytes, str]:
    """
    Generate Annexure-I document.

    Returns:
        (content_bytes, content_type)
        content_type is "application/pdf"
    """
    template_bytes = _load_template_bytes()
    replacements = resolve_annexure_values(client, company)
    logger.info(f"Annexure replacements: { {k: v for k, v in replacements.items()} }")

    filled_docx = _fill_template(template_bytes, replacements)

    pdf_bytes = _docx_to_pdf(filled_docx)
    if pdf_bytes:
        return pdf_bytes, "application/pdf"
    else:
        # Fall back to ReportLab PDF generator so PDF request always returns application/pdf
        logger.info("LibreOffice not available for DOCX->PDF; generating Annexure PDF via ReportLab.")
        pdf_bytes = _generate_annexure_pdf_reportlab(replacements, company or {})
        return pdf_bytes, "application/pdf"

