from __future__ import annotations
from typing import Any, Literal
"""PDF generators for Solarix documents."""
from io import BytesIO
from datetime import datetime, timezone
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, Image as RLImage
from reportlab.pdfgen import canvas
from reportlab.graphics.shapes import Drawing, Rect, String, Line, Group, Circle, PolyLine

styles = getSampleStyleSheet()
H1 = ParagraphStyle('h1', parent=styles['Heading1'], fontSize=16, textColor=colors.HexColor('#1d4ed8'), spaceAfter=8, alignment=1)
H2 = ParagraphStyle('h2', parent=styles['Heading2'], fontSize=11, textColor=colors.HexColor('#0f172a'), spaceAfter=6)
BODY = ParagraphStyle('body', parent=styles['BodyText'], fontSize=9, leading=13, textColor=colors.HexColor('#1f2937'))
SMALL = ParagraphStyle('small', parent=styles['BodyText'], fontSize=8, leading=11, textColor=colors.HexColor('#475569'))
BOLD_SMALL = ParagraphStyle('bold_small', parent=styles['BodyText'], fontSize=8, leading=11, fontName='Helvetica-Bold', textColor=colors.HexColor('#1f2937'))
HEADER_TEXT_STYLE = ParagraphStyle('header_text_style', parent=styles['BodyText'], fontSize=8, leading=11, fontName='Helvetica-Bold', textColor=colors.white)


def _header(company: dict, prepared_by: str | None = None, show_owner: bool = True):
    company_name = company.get('company_name') or company.get('name') or company.get('legal_business_name') or ''
    owner_name = company.get('owner_name') or company.get('proprietor_name') or company.get('authorized_signatory') or ''
    mobile = company.get('mobile') or company.get('phone') or company.get('phone_number') or ''
    email = company.get('email') or ''
    gst = company.get('gst_number') or company.get('gstin') or company.get('gst') or ''
    address = company.get('address') or company.get('address_line_1') or ''
    city = company.get('city') or ''
    state = company.get('state') or ''
    pincode = company.get('pincode') or ''
    website = company.get('website') or ''
    
    full_address = f"{address}"
    if city or state or pincode:
        full_address += f", {city}" if city else ""
        full_address += f", {state}" if state else ""
        full_address += f" - {pincode}" if pincode else ""
        
    lines = [
        f"<b><font size='14' color='#1d4ed8'>{company_name}</font></b>",
    ]
    if show_owner and owner_name:
        lines.append(f"Owner: {owner_name}")
    if prepared_by:
        lines.append(f"Prepared By: {prepared_by}")
    if mobile:
        lines.append(f"Mobile: {mobile}")
    if email:
        lines.append(f"Email: {email}")
    if website:
        lines.append(f"Website: {website}")
    if gst:
        lines.append(f"GSTIN: {gst}")
    if full_address:
        lines.append(f"Address: {full_address}")
        
    header_text = "<br/>".join(lines)
    header_p = Paragraph(header_text, ParagraphStyle('header_p', parent=styles['BodyText'], fontSize=9, leading=12))
    
    # Render Company Logo at TOP-RIGHT CORNER
    logo_bytes = company.get("logo_bytes")
    logo_d = None
    if logo_bytes:
        try:
            from PIL import Image as PILImage
            img = PILImage.open(BytesIO(logo_bytes))
            w, h = img.size
            if w > 0 and h > 0:
                aspect = h / w
                max_w = 9.0 * cm
                max_h = 4.4 * cm
                target_w = max_w
                target_h = target_w * aspect
                if target_h > max_h:
                    target_h = max_h
                    target_w = target_h / aspect
                resample_filter = getattr(getattr(PILImage, "Resampling", PILImage), "LANCZOS", getattr(PILImage, "LANCZOS", 1))
                resampled = img.resize((round(target_w * 4), round(target_h * 4)), resample_filter)
                res_buf = BytesIO()
                resampled.save(res_buf, format='PNG')
                logo_d = RLImage(BytesIO(res_buf.getvalue()), width=target_w, height=target_h)
        except Exception:
            logo_d = None

    if logo_d:
        header_table = Table([[header_p, logo_d]], colWidths=[13.0 * cm, 5.0 * cm])
        header_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ]))
    else:
        header_table = Table([[header_p]], colWidths=[18 * cm])
        header_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
    return [header_table, Spacer(1, 0.3 * cm)]


def _kv_table(rows):
    t = Table(rows, colWidths=[5 * cm, 13 * cm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f1f5f9')),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#475569')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#6b7280')),
        ('PADDING', (0, 0), (-1, -1), 5),
    ]))
    return t


def _format_currency(value: float | None) -> str:
    if value is None:
        return "Rs. 0.00"
    try:
        return f"Rs. {value:,.2f}"
    except Exception:
        return "Rs. 0.00"


def _amount_to_words(amount: float) -> str:
    words = ["zero", "one", "two", "three", "four", "five", "six", "seven", "eight", "nine", "ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen", "sixteen", "seventeen", "eighteen", "nineteen"]
    tens = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy", "eighty", "ninety"]

    def convert(num: int) -> str:
        if num < 20:
            return words[num]
        if num < 100:
            return tens[num // 10] + (" " + words[num % 10] if num % 10 else "")
        if num < 1000:
            return words[num // 100] + " hundred" + (" " + convert(num % 100) if num % 100 else "")
        if num < 100000:
            return convert(num // 1000) + " thousand" + (" " + convert(num % 1000) if num % 1000 else "")
        return convert(num // 100000) + " lakh" + (" " + convert(num % 100000) if num % 100000 else "")

    integer_part = int(amount)
    paise_part = round((amount - integer_part) * 100)
    words_out = convert(integer_part) + " rupees"
    if paise_part:
        words_out += " and " + convert(paise_part) + " paise"
    return words_out.replace("  ", " ").strip().capitalize() + " only"


def _safe_client_name(client: dict) -> str:
    return client.get("full_name") or client.get("name") or "Customer"


def _table(rows, col_widths=None, header_row=False):
    t = Table(rows, colWidths=col_widths, repeatRows=1 if header_row else 0)
    style = TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#6b7280")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ])
    if header_row:
        style.add("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e3a8a"))
        style.add("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold")
        style.add("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#ffffff"))
    t.setStyle(style)
    return t


def _client_table(client: dict):
    rows = [
        [Paragraph("<b>Client Name</b>", BOLD_SMALL), Paragraph(_safe_client_name(client), SMALL)],
        [Paragraph("<b>Address</b>", BOLD_SMALL), Paragraph(client.get("address", ""), SMALL)],
        [Paragraph("<b>Mobile</b>", BOLD_SMALL), Paragraph(client.get("mobile", ""), SMALL)],
        [Paragraph("<b>GSTIN</b>", BOLD_SMALL), Paragraph(client.get("gst_number", "") or "—", SMALL)],
        [Paragraph("<b>Email</b>", BOLD_SMALL), Paragraph(client.get("email", "") or "—", SMALL)],
        [Paragraph("<b>Site Address</b>", BOLD_SMALL), Paragraph(client.get("site_address", "") or client.get("address", ""), SMALL)],
    ]
    return _table(rows, col_widths=[5 * cm, 13 * cm])


def _vendor_table(vendor: dict):
    name = vendor.get("name") or vendor.get("vendor_name") or "Vendor"
    address = vendor.get("address") or vendor.get("vendor_address") or "—"
    phone = vendor.get("phone") or vendor.get("vendor_phone") or "—"
    gstin = vendor.get("gstin") or vendor.get("gst_number") or vendor.get("vendor_gstin") or "—"
    email = vendor.get("email") or vendor.get("vendor_email") or "—"
    rows = [
        [Paragraph("<b>Vendor Name</b>", BOLD_SMALL), Paragraph(name, SMALL)],
        [Paragraph("<b>Address</b>", BOLD_SMALL), Paragraph(address, SMALL)],
        [Paragraph("<b>Phone / Contact</b>", BOLD_SMALL), Paragraph(phone, SMALL)],
        [Paragraph("<b>GSTIN</b>", BOLD_SMALL), Paragraph(gstin, SMALL)],
        [Paragraph("<b>Email</b>", BOLD_SMALL), Paragraph(email, SMALL)],
    ]
    return _table(rows, col_widths=[5 * cm, 13 * cm])


def _render_items_table(doc_type: str, items: list[dict], data: dict, apply_gst: bool = True) -> Table:
    if doc_type == "tax_invoice":
        if apply_gst:
            headers = [
                Paragraph('<font color="#ffffff"><b>S.No</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>Description</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>Qty</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>Unit</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>Rate</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>Taxable Value</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>CGST</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>SGST</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>IGST</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>Total</b></font>', HEADER_TEXT_STYLE),
            ]
            rows = [headers]
            for idx, row in enumerate(items, 1):
                qty = float(row.get("quantity") or 0)
                rate = float(row.get("rate") or 0)
                discount = float(row.get("discount") or 0)
                taxable = max(0.0, qty * rate - discount)
                
                cgst = float(row.get("cgst") or 0)
                sgst = float(row.get("sgst") or 0)
                igst = float(row.get("igst") or 0)
                total = float(row.get("amount") or (taxable + cgst + sgst + igst))
                
                desc_text = row.get("product", "")
                serials = row.get("serial_numbers") or row.get("serials")
                if serials:
                    desc_text += f'<br/><font size="7.5" color="#64748b">Serial: {serials}</font>'
                
                rows.append([
                    Paragraph(str(idx), SMALL),
                    Paragraph(desc_text, SMALL),
                    Paragraph(str(qty), SMALL),
                    Paragraph(row.get("unit", ""), SMALL),
                    Paragraph(_format_currency(rate), SMALL),
                    Paragraph(_format_currency(taxable), SMALL),
                    Paragraph(_format_currency(cgst), SMALL),
                    Paragraph(_format_currency(sgst), SMALL),
                    Paragraph(_format_currency(igst), SMALL),
                    Paragraph(_format_currency(total), SMALL),
                ])
            col_widths = [0.8 * cm, 4.0 * cm, 0.9 * cm, 1.0 * cm, 2.0 * cm, 2.1 * cm, 1.6 * cm, 1.6 * cm, 1.6 * cm, 2.4 * cm]
            return _table(rows, col_widths=col_widths, header_row=True)
        else:
            headers = [
                Paragraph('<font color="#ffffff"><b>S.No</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>Description</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>Qty</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>Unit</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>Rate</b></font>', HEADER_TEXT_STYLE),
                Paragraph('<font color="#ffffff"><b>Total</b></font>', HEADER_TEXT_STYLE),
            ]
            rows = [headers]
            for idx, row in enumerate(items, 1):
                qty = float(row.get("quantity") or 0)
                rate = float(row.get("rate") or 0)
                discount = float(row.get("discount") or 0)
                taxable = max(0.0, qty * rate - discount)
                
                desc_text = row.get("product", "")
                serials = row.get("serial_numbers") or row.get("serials")
                if serials:
                    desc_text += f'<br/><font size="7.5" color="#64748b">Serial: {serials}</font>'
                
                rows.append([
                    Paragraph(str(idx), SMALL),
                    Paragraph(desc_text, SMALL),
                    Paragraph(str(qty), SMALL),
                    Paragraph(row.get("unit", ""), SMALL),
                    Paragraph(_format_currency(rate), SMALL),
                    Paragraph(_format_currency(taxable), SMALL),
                ])
            col_widths = [1.0 * cm, 9.0 * cm, 1.5 * cm, 1.5 * cm, 2.5 * cm, 2.5 * cm]
            return _table(rows, col_widths=col_widths, header_row=True)

    if doc_type == "purchase_order":
        headers = [
            Paragraph('<font color="#ffffff"><b>No.</b></font>', HEADER_TEXT_STYLE),
            Paragraph('<font color="#ffffff"><b>Product Name / Description</b></font>', HEADER_TEXT_STYLE),
            Paragraph('<font color="#ffffff"><b>Qty</b></font>', HEADER_TEXT_STYLE),
            Paragraph('<font color="#ffffff"><b>Unit</b></font>', HEADER_TEXT_STYLE),
            Paragraph('<font color="#ffffff"><b>Unit Price</b></font>', HEADER_TEXT_STYLE),
            Paragraph('<font color="#ffffff"><b>Total</b></font>', HEADER_TEXT_STYLE),
        ]
        rows = [headers]
        for idx, row in enumerate(items, 1):
            p_name = row.get("product_name") or row.get("product") or ""
            size = row.get("size") or ""
            if size and size not in p_name:
                p_name += f" ({size})"
            qty = float(row.get("quantity") or 0)
            rate = float(row.get("unit_price") or row.get("rate") or 0)
            amount = float(row.get("amount") or (qty * rate))
            
            rows.append([
                Paragraph(str(idx), SMALL),
                Paragraph(p_name, SMALL),
                Paragraph(f"{qty:g}", SMALL),
                Paragraph(str(row.get("unit") or "Nos"), SMALL),
                Paragraph(_format_currency(rate), SMALL),
                Paragraph(_format_currency(amount), SMALL),
            ])
        col_widths = [1.0 * cm, 8.5 * cm, 1.8 * cm, 1.5 * cm, 2.6 * cm, 2.6 * cm]
        return _table(rows, col_widths=col_widths, header_row=True)
            
    if doc_type == "delivery_bill":
        show_rate = data.get("show_rate", True)
        show_amount = data.get("show_amount", True)
        
        headers = [
            Paragraph('<font color="#ffffff"><b>S.No</b></font>', HEADER_TEXT_STYLE),
            Paragraph('<font color="#ffffff"><b>Description</b></font>', HEADER_TEXT_STYLE),
            Paragraph('<font color="#ffffff"><b>Size</b></font>', HEADER_TEXT_STYLE),
            Paragraph('<font color="#ffffff"><b>Unit</b></font>', HEADER_TEXT_STYLE),
            Paragraph('<font color="#ffffff"><b>Dispatch Qty</b></font>', HEADER_TEXT_STYLE),
        ]
        if show_rate:
            headers.append(Paragraph('<font color="#ffffff"><b>Rate</b></font>', HEADER_TEXT_STYLE))
        if show_amount:
            headers.append(Paragraph('<font color="#ffffff"><b>Amount</b></font>', HEADER_TEXT_STYLE))
            
        hv_keywords = ["SOLAR PANEL", "PANEL", "INVERTER", "ACDB", "DCDB", "METER", "BATTERY"]
        def _is_hv_db(row):
            pn = (row.get("product") or "").upper()
            return row.get("high_value_goods") or row.get("high_value_asset") or any(kw in pn for kw in hv_keywords)

        sorted_items = list(items or [])
        sorted_items.sort(key=lambda r: (0 if _is_hv_db(r) else 1, (r.get("product") or "").lower(), (r.get("size") or "").lower()))

        rows = [headers]
        for idx, row in enumerate(sorted_items, 1):
            qty = float(row.get("dispatch_qty") or 0)
            rate = float(row.get("rate") or 0)
            amount = qty * rate
            
            desc_text = row.get("product", "")
            serials = row.get("serial_numbers") or row.get("serials")
            if serials:
                desc_text += f'<br/><font size="7.5" color="#64748b">Serial: {serials}</font>'
                
            r_data = [
                Paragraph(str(idx), SMALL),
                Paragraph(desc_text, SMALL),
                Paragraph(row.get("size", ""), SMALL),
                Paragraph(row.get("unit", ""), SMALL),
                Paragraph(str(qty), SMALL),
            ]
            if show_rate:
                r_data.append(Paragraph(_format_currency(rate), SMALL))
            if show_amount:
                r_data.append(Paragraph(_format_currency(amount), SMALL))
            rows.append(r_data)
            
        if show_rate and show_amount:
            col_widths = [1.0 * cm, 6.5 * cm, 2.0 * cm, 1.5 * cm, 2.0 * cm, 2.5 * cm, 2.5 * cm]
        elif show_rate:
            col_widths = [1.0 * cm, 9.0 * cm, 2.0 * cm, 1.5 * cm, 2.0 * cm, 2.5 * cm]
        elif show_amount:
            col_widths = [1.0 * cm, 9.0 * cm, 2.0 * cm, 1.5 * cm, 2.0 * cm, 2.5 * cm]
        else:
            col_widths = [1.0 * cm, 11.5 * cm, 2.0 * cm, 1.5 * cm, 2.0 * cm]
            
        return _table(rows, col_widths=col_widths, header_row=True)

    # Quotation
    custom_cols = data.get("custom_columns") or []
    formula_cols = data.get("formula_columns") or []
    num_extra = len(custom_cols) + len(formula_cols)
    extra_width = 1.5 * cm
    
    headers = [
        Paragraph('<font color="#ffffff"><b>S.No</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Description</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Size</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Unit</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Qty</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Rate</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Discount</b></font>', HEADER_TEXT_STYLE),
    ]
    if apply_gst:
        headers.append(Paragraph('<font color="#ffffff"><b>GST %</b></font>', HEADER_TEXT_STYLE))
    headers.append(Paragraph('<font color="#ffffff"><b>Amount</b></font>', HEADER_TEXT_STYLE))
    
    for c in custom_cols:
        headers.append(Paragraph(f'<font color="#ffffff"><b>{c.get("label", "Custom")}</b></font>', HEADER_TEXT_STYLE))
    for f in formula_cols:
        headers.append(Paragraph(f'<font color="#ffffff"><b>{f.get("label", "Formula")}</b></font>', HEADER_TEXT_STYLE))
        
    rows = [headers]
    for idx, row in enumerate(items, 1):
        qty = float(row.get("quantity") or 0)
        rate = float(row.get("rate") or 0)
        discount = float(row.get("discount") or 0)
        gst = float(row.get("gst") or 0)
        
        taxable = max(0.0, qty * rate - discount)
        gst_amount = taxable * gst / 100 if apply_gst else 0
        amount = float(row.get("amount") or (taxable + gst_amount))
        
        desc_text = row.get("product", "")
        serials = row.get("serial_numbers") or row.get("serials")
        if serials:
            desc_text += f'<br/><font size="7.5" color="#64748b">Serial: {serials}</font>'
            
        r_cols = [
            Paragraph(str(idx), SMALL),
            Paragraph(desc_text, SMALL),
            Paragraph(row.get("size", ""), SMALL),
            Paragraph(row.get("unit", ""), SMALL),
            Paragraph(str(qty), SMALL),
            Paragraph(_format_currency(rate), SMALL),
            Paragraph(_format_currency(discount), SMALL),
        ]
        if apply_gst:
            r_cols.append(Paragraph(f"{gst:.0f}%", SMALL))
        r_cols.append(Paragraph(_format_currency(amount), SMALL))
        
        # custom fields
        custom_data = row.get("custom") or {}
        for c in custom_cols:
            val = custom_data.get(c.get("id"), "")
            r_cols.append(Paragraph(str(val), SMALL))
            
        # formula fields
        formula_data = row.get("formula") or {}
        for f in formula_cols:
            val = formula_data.get(f.get("id"), 0)
            r_cols.append(Paragraph(_format_currency(val), SMALL))
            
        rows.append(r_cols)

    rem_width = 18.0 * cm - (num_extra * extra_width)
    if apply_gst:
        base_widths = [0.8 * cm, 4.0 * cm, 1.5 * cm, 1.0 * cm, 1.0 * cm, 2.0 * cm, 1.8 * cm, 1.4 * cm, 2.5 * cm]
    else:
        base_widths = [1.0 * cm, 4.5 * cm, 2.0 * cm, 1.2 * cm, 1.2 * cm, 2.5 * cm, 2.1 * cm, 2.5 * cm]
        
    total_base = sum(base_widths)
    scale = rem_width / total_base
    col_widths = [w * scale for w in base_widths] + [extra_width] * num_extra
    
    return _table(rows, col_widths=col_widths, header_row=True)


def _summary_table(doc_type: str, totals: dict):
    rows = [[Paragraph('<font color="#ffffff"><b>Description</b></font>', HEADER_TEXT_STYLE), Paragraph('<font color="#ffffff"><b>Amount</b></font>', HEADER_TEXT_STYLE)]]
    if doc_type in ("tax_invoice", "purchase_order", "purchase_bill") and ("subtotal" in totals or "grand_total" in totals):
        rows.append([Paragraph("Subtotal", SMALL), Paragraph(_format_currency(totals.get("subtotal", 0)), SMALL)])
        if totals.get("cgst_amount"):
            rows.append([Paragraph(f"CGST ({totals.get('cgst_rate', 0)}%)", SMALL), Paragraph(_format_currency(totals.get("cgst_amount", 0)), SMALL)])
        if totals.get("sgst_amount"):
            rows.append([Paragraph(f"SGST ({totals.get('sgst_rate', 0)}%)", SMALL), Paragraph(_format_currency(totals.get("sgst_amount", 0)), SMALL)])
        if totals.get("igst_amount"):
            rows.append([Paragraph(f"IGST ({totals.get('igst_rate', 0)}%)", SMALL), Paragraph(_format_currency(totals.get("igst_amount", 0)), SMALL)])
        if totals.get("freight"):
            rows.append([Paragraph("Freight / S&H", SMALL), Paragraph(_format_currency(totals.get("freight", 0)), SMALL)])
        rows.append([Paragraph("<b>Grand Total</b>", BOLD_SMALL), Paragraph(_format_currency(totals.get("grand_total", totals.get("subtotal", 0))), BOLD_SMALL)])
    else:
        rows.extend([
            [Paragraph("<b>Total</b>", BOLD_SMALL), Paragraph(_format_currency(totals.get("total", 0)), BOLD_SMALL)],
        ])
    return _table(rows, col_widths=[13 * cm, 5 * cm], header_row=True)



def _dedupe(value: str) -> str:
    return (value or "").strip()


def _normalize_po_document_data(data: dict, company: dict) -> dict:
    if not isinstance(data, dict):
        data = {}
    if not isinstance(company, dict):
        company = {}

    comp_name = (company.get("company_name") or company.get("name") or company.get("legal_business_name") or "SOLARIX").strip()
    comp_gst = (company.get("gst_number") or company.get("gstin") or company.get("gst") or "").strip()
    comp_addr = (company.get("address") or company.get("address_line_1") or company.get("office_address") or "").strip()
    comp_phone = (company.get("mobile") or company.get("mobile_number") or company.get("phone") or company.get("phone_number") or "").strip()
    comp_email = (company.get("email") or "").strip()
    comp_tagline = (company.get("tagline") or company.get("subtitle") or "").strip()
    logo_bytes = company.get("logo_bytes")

    vendor_raw = data.get("vendor") or {}
    if not isinstance(vendor_raw, dict):
        vendor_raw = {}
        
    v_name = (data.get("vendor_name") or vendor_raw.get("name") or vendor_raw.get("vendor_name") or "").strip()
    v_addr = (data.get("vendor_address") or vendor_raw.get("address") or vendor_raw.get("vendor_address") or "").strip()
    v_phone = (data.get("vendor_phone") or vendor_raw.get("phone") or vendor_raw.get("vendor_phone") or "").strip()
    v_email = (data.get("vendor_email") or vendor_raw.get("email") or vendor_raw.get("vendor_email") or "").strip()
    v_gstin = (data.get("vendor_gstin") or vendor_raw.get("gstin") or vendor_raw.get("vendor_gstin") or vendor_raw.get("gst_number") or "").strip()
    v_id = (data.get("vendor_id") or vendor_raw.get("id") or "").strip()

    po_num = (data.get("po_number") or data.get("document_number") or data.get("id") or "").strip()
    po_date = (data.get("po_date") or data.get("document_date") or data.get("date") or datetime.now().strftime("%Y-%m-%d")).strip()
    del_date = (data.get("delivery_date") or "").strip()

    ship_via = (data.get("ship_via") or "FOR").strip()
    shipping_method = (data.get("shipping_method") or "PAID").strip()
    shipping_term = (data.get("shipping_term") or "DOOR DELIVERY").strip()

    ship_to_raw = data.get("ship_to") or {}
    if isinstance(ship_to_raw, dict) and ship_to_raw.get("name"):
        ship_name = (ship_to_raw.get("name") or comp_name).strip()
        ship_addr = (ship_to_raw.get("address") or comp_addr).strip()
        ship_phone = (ship_to_raw.get("phone") or comp_phone).strip()
        ship_email = (ship_to_raw.get("email") or comp_email).strip()
        ship_gstin = (ship_to_raw.get("gstin") or ship_to_raw.get("gst_number") or comp_gst).strip()
    else:
        ship_name = comp_name
        ship_addr = comp_addr
        ship_phone = comp_phone
        ship_email = comp_email
        ship_gstin = comp_gst

    items_raw = data.get("items") or []
    line_items = []
    subtotal_calc = 0.0

    for idx, item in enumerate(items_raw, 1):
        if not isinstance(item, dict):
            continue
        p_name = (item.get("product_name") or item.get("product") or "").strip()
        if not p_name:
            continue
        size = (item.get("size") or "").strip()
        qty = float(item.get("quantity") or item.get("qty") or 0.0)
        unit = (item.get("unit") or "Nos").strip()
        price = float(item.get("unit_price") or item.get("rate") or 0.0)
        amount = float(item.get("amount") or (qty * price))
        subtotal_calc += amount

        line_items.append({
            "code": str(idx),
            "product_name": p_name,
            "size": size,
            "quantity": qty,
            "unit": unit,
            "unit_price": price,
            "amount": amount
        })

    subtotal = float(data.get("subtotal") or subtotal_calc)
    discount = float(data.get("discount") or 0.0)

    try:
        cgst_rate = float(data.get("cgst_rate") if data.get("cgst_rate") is not None else 2.5)
    except (ValueError, TypeError):
        cgst_rate = 2.5

    try:
        sgst_rate = float(data.get("sgst_rate") if data.get("sgst_rate") is not None else 2.5)
    except (ValueError, TypeError):
        sgst_rate = 2.5

    try:
        igst_rate = float(data.get("igst_rate") if data.get("igst_rate") is not None else 0.0)
    except (ValueError, TypeError):
        igst_rate = 0.0

    try:
        freight = float(data.get("freight") or data.get("sh_freight") or 0.0)
    except (ValueError, TypeError):
        freight = 0.0

    cgst_amount = round(subtotal * (cgst_rate / 100.0), 2)
    sgst_amount = round(subtotal * (sgst_rate / 100.0), 2)
    igst_amount = round(subtotal * (igst_rate / 100.0), 2)

    grand_total_calc = round(subtotal - discount + cgst_amount + sgst_amount + igst_amount + freight, 2)
    try:
        saved_grand_total = float(data.get("grand_total") or 0.0)
    except (ValueError, TypeError):
        saved_grand_total = 0.0

    grand_total = saved_grand_total if saved_grand_total > 0 else grand_total_calc

    notes = (data.get("notes") or "DELIVERY WILL BE F.O.R. ON-SITE\nLOCATION OF SITE WILL BE PROVIDED AT THE TIME OF DISPATCH").strip()

    return {
        "company": {
            "name": comp_name,
            "tagline": comp_tagline,
            "address": comp_addr,
            "phone": comp_phone,
            "email": comp_email,
            "gstin": comp_gst,
            "logo_bytes": logo_bytes
        },
        "poDetails": {
            "po_number": po_num,
            "po_date": po_date,
            "vendor_id": v_id,
            "delivery_date": del_date
        },
        "vendor": {
            "name": v_name,
            "address": v_addr,
            "phone": v_phone,
            "email": v_email,
            "gstin": v_gstin
        },
        "shipTo": {
            "name": ship_name,
            "address": ship_addr,
            "phone": ship_phone,
            "email": ship_email,
            "gstin": ship_gstin
        },
        "shipping": {
            "ship_via": ship_via,
            "shipping_method": shipping_method,
            "shipping_term": shipping_term,
            "delivery_date": del_date
        },
        "lineItems": line_items,
        "financials": {
            "subtotal": subtotal,
            "discount": discount,
            "cgst_rate": cgst_rate,
            "cgst_amount": cgst_amount,
            "sgst_rate": sgst_rate,
            "sgst_amount": sgst_amount,
            "igst_rate": igst_rate,
            "igst_amount": igst_amount,
            "freight": freight,
            "grand_total": grand_total
        },
        "notes": notes,
        "footer_text": "THIS IS NOT TAX INVOICE !"
    }


def generate_po_pdf(data: dict, company: dict) -> bytes:
    """Generate professional PDF for Purchase Orders matching GVP reference layout."""
    norm = _normalize_po_document_data(data, company)
    comp = norm["company"]
    po = norm["poDetails"]
    vendor = norm["vendor"]
    ship_to = norm["shipTo"]
    shipping = norm["shipping"]
    items = norm["lineItems"]
    fin = norm["financials"]
    notes = norm["notes"]

    buf = BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm
    )
    story = []

    # 1. HEADER (Logo/Company Left, Title/Meta Right)
    logo_bytes = comp.get("logo_bytes")
    logo_d = None
    if logo_bytes:
        try:
            from PIL import Image as PILImage
            img = PILImage.open(BytesIO(logo_bytes))
            orig_w, orig_h = img.size
            if orig_w > 0 and orig_h > 0:
                max_w, max_h = 11.0 * cm, 3.0 * cm
                aspect = orig_w / float(orig_h or 1)
                if orig_w > orig_h:
                    target_w = max_w
                    target_h = max_w / aspect
                    if target_h > max_h:
                        target_h = max_h
                        target_w = max_h * aspect
                else:
                    target_h = max_h
                    target_w = max_h * aspect
                    if target_w > max_w:
                        target_w = max_w
                        target_h = max_w / aspect
                res_buf = BytesIO()
                img.save(res_buf, format="PNG")
                logo_d = RLImage(BytesIO(res_buf.getvalue()), width=target_w, height=target_h)
        except Exception:
            logo_d = None

    hdr_left_flow = []
    if logo_d:
        hdr_left_flow.append(logo_d)
        hdr_left_flow.append(Spacer(1, 0.1 * cm))

    comp_title_style = ParagraphStyle('po_comp_title', parent=styles['Normal'], fontSize=13, leading=15, textColor=colors.HexColor('#1e3a8a'), fontName='Helvetica-Bold')
    comp_sub_style = ParagraphStyle('po_comp_sub', parent=styles['Normal'], fontSize=8, leading=10, textColor=colors.HexColor('#475569'), fontName='Helvetica-Bold')
    comp_gst_style = ParagraphStyle('po_comp_gst', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#0f172a'), fontName='Helvetica-Bold')

    hdr_left_flow.append(Paragraph(comp["name"], comp_title_style))
    if comp["tagline"]:
        hdr_left_flow.append(Paragraph(comp["tagline"], comp_sub_style))
    if comp["gstin"]:
        hdr_left_flow.append(Paragraph(f"GST: {comp['gstin']}", comp_gst_style))

    po_title_style = ParagraphStyle('po_title', parent=styles['Normal'], fontSize=16, leading=18, textColor=colors.HexColor('#1e3a8a'), fontName='Helvetica-Bold', alignment=2)
    hdr_right_flow: list[Any] = [Paragraph("PURCHASE ORDER", po_title_style), Spacer(1, 0.2 * cm)]

    meta_table_data = [
        [Paragraph("<b>DATE</b>", BOLD_SMALL), Paragraph(po["po_date"], SMALL)],
        [Paragraph("<b>P.O. NUMBER</b>", BOLD_SMALL), Paragraph(po["po_number"], SMALL)],
        [Paragraph("<b>VENDOR ID</b>", BOLD_SMALL), Paragraph(po["vendor_id"] or "—", SMALL)],
    ]
    meta_table = Table(meta_table_data, colWidths=[2.8 * cm, 4.5 * cm])
    meta_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#f8fafc')),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ('LEFTPADDING', (0,0), (-1,-1), 5),
        ('RIGHTPADDING', (0,0), (-1,-1), 5),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    hdr_right_flow.append(meta_table)

    header_table = Table([[hdr_left_flow, hdr_right_flow]], colWidths=[10.8 * cm, 7.8 * cm])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ('TOPPADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 0.4 * cm))

    # 2. TWO-COLUMN PARTY SECTION (VENDOR | SHIP TO)
    v_body = f"<b>{vendor['name']}</b><br/>"
    if vendor['address']:
        v_body += f"{vendor['address']}<br/>"
    if vendor['phone']:
        v_body += f"Phone: {vendor['phone']}<br/>"
    if vendor['email']:
        v_body += f"Email: {vendor['email']}<br/>"
    if vendor['gstin']:
        v_body += f"GSTIN: {vendor['gstin']}"

    s_body = f"<b>{ship_to['name']}</b><br/>"
    if ship_to['address']:
        s_body += f"{ship_to['address']}<br/>"
    if ship_to['phone']:
        s_body += f"Phone: {ship_to['phone']}<br/>"
    if ship_to['email']:
        s_body += f"Email: {ship_to['email']}<br/>"
    if ship_to['gstin']:
        s_body += f"GSTIN: {ship_to['gstin']}"

    party_header_style = ParagraphStyle('p_hdr', parent=styles['Normal'], fontSize=9, leading=11, textColor=colors.HexColor('#ffffff'), fontName='Helvetica-Bold')
    party_body_style = ParagraphStyle('p_bdy', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#0f172a'))

    party_table_data = [
        [Paragraph("VENDOR", party_header_style), Paragraph("SHIP TO", party_header_style)],
        [Paragraph(v_body, party_body_style), Paragraph(s_body, party_body_style)]
    ]
    party_table = Table(party_table_data, colWidths=[9.1 * cm, 9.1 * cm])
    party_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e3a8a')),
        ('BACKGROUND', (0,1), (-1,1), colors.HexColor('#ffffff')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(party_table)
    story.append(Spacer(1, 0.3 * cm))

    # 3. SHIPPING TERMS BAR
    ship_hdr_style = ParagraphStyle('s_hdr', parent=styles['Normal'], fontSize=8, leading=10, textColor=colors.HexColor('#ffffff'), fontName='Helvetica-Bold', alignment=1)
    ship_val_style = ParagraphStyle('s_val', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#0f172a'), fontName='Helvetica-Bold', alignment=1)

    ship_bar_data = [
        [
            Paragraph("SHIP VIA", ship_hdr_style),
            Paragraph("SHIPPING METHOD", ship_hdr_style),
            Paragraph("SHIPPING TERM", ship_hdr_style),
            Paragraph("DELIVERY DATE", ship_hdr_style)
        ],
        [
            Paragraph(shipping["ship_via"] or "FOR", ship_val_style),
            Paragraph(shipping["shipping_method"] or "PAID", ship_val_style),
            Paragraph(shipping["shipping_term"] or "DOOR DELIVERY", ship_val_style),
            Paragraph(shipping["delivery_date"] or "—", ship_val_style)
        ]
    ]
    ship_table = Table(ship_bar_data, colWidths=[4.55 * cm, 4.55 * cm, 4.55 * cm, 4.55 * cm])
    ship_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e3a8a')),
        ('BACKGROUND', (0,1), (-1,1), colors.HexColor('#f8fafc')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(ship_table)
    story.append(Spacer(1, 0.3 * cm))

    # 4. LINE ITEMS TABLE
    th_style = ParagraphStyle('th_style', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#ffffff'), fontName='Helvetica-Bold')
    th_center = ParagraphStyle('th_c', parent=th_style, alignment=1)
    th_right = ParagraphStyle('th_r', parent=th_style, alignment=2)

    item_headers = [
        Paragraph("CODE", th_center),
        Paragraph("PRODUCT NAME / DESCRIPTION", th_style),
        Paragraph("QTY", th_center),
        Paragraph("UNIT", th_center),
        Paragraph("UNIT PRICE", th_right),
        Paragraph("TOTAL", th_right),
    ]
    item_rows = [item_headers]

    tb_style = ParagraphStyle('tb_style', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#0f172a'))
    tb_center = ParagraphStyle('tb_c', parent=tb_style, alignment=1)
    tb_right = ParagraphStyle('tb_r', parent=tb_style, alignment=2)

    for it in items:
        p_name = it["product_name"]
        if it["size"] and it["size"] not in p_name:
            p_name += f" ({it['size']})"

        item_rows.append([
            Paragraph(it["code"], tb_center),
            Paragraph(p_name, tb_style),
            Paragraph(f"{it['quantity']:g}", tb_center),
            Paragraph(it["unit"], tb_center),
            Paragraph(_format_currency(it["unit_price"]), tb_right),
            Paragraph(_format_currency(it["amount"]), tb_right),
        ])

    items_table = Table(item_rows, colWidths=[1.2 * cm, 8.0 * cm, 1.8 * cm, 1.6 * cm, 3.0 * cm, 3.0 * cm])
    items_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e3a8a')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 4),
        ('RIGHTPADDING', (0,0), (-1,-1), 4),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(items_table)
    story.append(Spacer(1, 0.3 * cm))

    # 5. BOTTOM SECTION (NOTES & TOTALS)
    notes_p = Paragraph(f"<b>NOTES AND INSTRUCTION</b><br/><br/>{notes.replace(chr(10), '<br/>')}", ParagraphStyle('notes_p', parent=styles['Normal'], fontSize=8, leading=11, textColor=colors.HexColor('#1e293b')))

    tot_label = ParagraphStyle('tot_lbl', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#334155'), fontName='Helvetica-Bold')
    tot_val = ParagraphStyle('tot_v', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#0f172a'), alignment=2)
    tot_grand_lbl = ParagraphStyle('tot_g_lbl', parent=styles['Normal'], fontSize=9.5, leading=12, textColor=colors.HexColor('#1e3a8a'), fontName='Helvetica-Bold')
    tot_grand_val = ParagraphStyle('tot_g_v', parent=styles['Normal'], fontSize=10, leading=12, textColor=colors.HexColor('#1e3a8a'), fontName='Helvetica-Bold', alignment=2)

    totals_rows = [
        [Paragraph("SUBTOTAL", tot_label), Paragraph(_format_currency(fin["subtotal"]), tot_val)]
    ]
    if fin["discount"] > 0:
        totals_rows.append([Paragraph("DISCOUNT", tot_label), Paragraph(_format_currency(fin["discount"]), tot_val)])
    if fin["cgst_amount"] > 0 or fin["cgst_rate"] > 0:
        totals_rows.append([Paragraph(f"CGST ({fin['cgst_rate']}%)", tot_label), Paragraph(_format_currency(fin["cgst_amount"]), tot_val)])
    if fin["sgst_amount"] > 0 or fin["sgst_rate"] > 0:
        totals_rows.append([Paragraph(f"SGST ({fin['sgst_rate']}%)", tot_label), Paragraph(_format_currency(fin["sgst_amount"]), tot_val)])
    if fin["igst_amount"] > 0 or fin["igst_rate"] > 0:
        totals_rows.append([Paragraph(f"IGST ({fin['igst_rate']}%)", tot_label), Paragraph(_format_currency(fin["igst_amount"]), tot_val)])
    if fin["freight"] > 0:
        totals_rows.append([Paragraph("S&H / FREIGHT", tot_label), Paragraph(_format_currency(fin["freight"]), tot_val)])

    totals_rows.append([Paragraph("GRAND TOTAL", tot_grand_lbl), Paragraph(_format_currency(fin["grand_total"]), tot_grand_val)])

    totals_table = Table(totals_rows, colWidths=[4.4 * cm, 3.7 * cm])
    totals_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#eff6ff')),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING', (0,0), (-1,-1), 5),
        ('RIGHTPADDING', (0,0), (-1,-1), 5),
    ]))

    bottom_table = Table([[notes_p, totals_table]], colWidths=[10.3 * cm, 8.3 * cm])
    bottom_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ('TOPPADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(bottom_table)
    story.append(Spacer(1, 0.4 * cm))

    # Amount in Words
    story.append(Paragraph("<b>Amount in Words</b>", ParagraphStyle('po_words_hdr', parent=styles['Normal'], fontSize=9, leading=11, fontName='Helvetica-Bold', textColor=colors.HexColor('#1e3a8a'))))
    story.append(Paragraph(_amount_to_words(fin.get("grand_total") or 0), ParagraphStyle('po_words_body', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#334155'))))
    story.append(Spacer(1, 0.5 * cm))

    # Authorized Signature Block
    SIG_LEFT = ParagraphStyle('po_sig_l', parent=styles['Normal'], fontSize=8.5, leading=12, textColor=colors.HexColor('#1f2937'), alignment=0)
    SIG_RIGHT = ParagraphStyle('po_sig_r', parent=styles['Normal'], fontSize=8.5, leading=12, textColor=colors.HexColor('#1f2937'), alignment=2)
    sig_table = Table([
        [Paragraph("<b>Vendor Acceptance Signature</b><br/><br/>_______________________", SIG_LEFT), Paragraph(f"<b>For {comp['name']}</b><br/><br/>_______________________<br/>Authorized Signatory", SIG_RIGHT)]
    ], colWidths=[9.3 * cm, 9.3 * cm])
    sig_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(KeepTogether([sig_table]))
    story.append(Spacer(1, 0.4 * cm))

    # 6. FOOTER & NOT TAX INVOICE NOTICE
    footer_p = Paragraph(f"Enquiries: {comp['email']} | Contact: {comp['phone']} | Office: {comp['address']}", ParagraphStyle('po_ftr', parent=styles['Normal'], fontSize=7.5, leading=10, textColor=colors.HexColor('#64748b'), alignment=1))
    story.append(footer_p)
    story.append(Spacer(1, 0.2 * cm))

    not_tax_p = Paragraph("<font size='10' color='#dc2626'><b>THIS IS NOT TAX INVOICE !</b></font>", ParagraphStyle('po_not_tax', parent=styles['Normal'], alignment=1))
    story.append(not_tax_p)

    pdf.build(story)
    return buf.getvalue()


def _normalize_invoice_document_data(data: dict, company: dict) -> dict:
    if not isinstance(data, dict):
        data = {}
    if not isinstance(company, dict):
        company = {}

    comp_name = (company.get("company_name") or company.get("name") or company.get("legal_business_name") or "GVP SOLAR ENERGY").strip()
    comp_gst = (company.get("gst_number") or company.get("gstin") or company.get("gst") or "").strip()
    comp_addr = (company.get("address") or company.get("address_line_1") or company.get("office_address") or "").strip()
    comp_phone = (company.get("mobile") or company.get("mobile_number") or company.get("phone") or company.get("phone_number") or "").strip()
    comp_email = (company.get("email") or "").strip()
    comp_website = (company.get("website") or "").strip()
    comp_state = (company.get("state") or "Maharashtra").strip()
    comp_state_code = (company.get("state_code") or "27").strip()
    comp_bank_name = (company.get("bank_name") or "").strip()
    comp_bank_acc = (company.get("bank_account_number") or company.get("account_number") or "").strip()
    comp_ifsc = (company.get("bank_ifsc") or company.get("ifsc_code") or "").strip()
    comp_branch = (company.get("bank_branch") or "").strip()
    logo_bytes = company.get("logo_bytes")

    client_raw = data.get("client") or data.get("buyer") or {}
    if not isinstance(client_raw, dict):
        client_raw = {}

    c_name = (data.get("client_name") or client_raw.get("full_name") or client_raw.get("name") or "Customer").strip()
    c_addr = (client_raw.get("site_address") or client_raw.get("address") or client_raw.get("billing_address") or "").strip()
    c_phone = (client_raw.get("mobile") or client_raw.get("phone") or "").strip()
    c_email = (client_raw.get("email") or "").strip()
    c_gstin = (data.get("buyer_gstin") or client_raw.get("gstin") or client_raw.get("gst_number") or "").strip()
    c_state = (data.get("supply_state") or client_raw.get("state") or comp_state).strip()
    c_sol_id = (client_raw.get("sol_id") or client_raw.get("sol_number") or "").strip()
    c_consumer = (client_raw.get("consumer_number") or "").strip()

    doc_type = (data.get("doc_type") or data.get("invoice_type") or "tax_invoice").lower().strip()
    title_map = {
        "tax_invoice": "TAX INVOICE",
        "proforma": "PROFORMA INVOICE",
        "payment_receipt": "PAYMENT RECEIPT",
        "credit_note": "CREDIT NOTE",
        "debit_note": "DEBIT NOTE"
    }
    title_text = title_map.get(doc_type, "TAX INVOICE")

    inv_num = (data.get("invoice_number") or data.get("receipt_number") or data.get("document_number") or data.get("id") or "").strip()
    inv_date = (data.get("invoice_date") or data.get("receipt_date") or data.get("date") or datetime.now().strftime("%Y-%m-%d")).strip()
    due_date = (data.get("due_date") or "").strip()
    place_of_supply = (data.get("place_of_supply") or c_state).strip()
    reverse_charge = (data.get("reverse_charge") or "No").strip()
    orig_inv_num = (data.get("original_invoice_number") or "").strip()
    reason = (data.get("reason") or "").strip()
    payment_terms = (data.get("payment_terms") or "Due on Receipt").strip()

    items_raw = data.get("items") or []
    line_items = []
    subtotal_calc = 0.0

    for idx, item in enumerate(items_raw, 1):
        if not isinstance(item, dict):
            continue
        p_name = (item.get("product_name") or item.get("product") or "").strip()
        if not p_name:
            continue
        hsn = (item.get("hsn_sac") or item.get("hsn") or "").strip()
        spec = (item.get("size") or item.get("spec") or "").strip()
        qty = float(item.get("quantity") or item.get("qty") or 0.0)
        unit = (item.get("unit") or "Nos").strip()
        rate = float(item.get("rate") if item.get("rate") is not None else item.get("unit_price") or 0.0)
        disc = float(item.get("discount") or 0.0)
        taxable = max(0.0, (qty * rate) - disc)
        subtotal_calc += taxable

        gst_r = float(item.get("gst_rate") if item.get("gst_rate") is not None else item.get("gst") or 18.0)
        line_cgst = float(item.get("cgst") or (taxable * (gst_r / 2.0 / 100.0) if place_of_supply.lower() == comp_state.lower() else 0.0))
        line_sgst = float(item.get("sgst") or (taxable * (gst_r / 2.0 / 100.0) if place_of_supply.lower() == comp_state.lower() else 0.0))
        line_igst = float(item.get("igst") or (taxable * (gst_r / 100.0) if place_of_supply.lower() != comp_state.lower() else 0.0))
        amt = float(item.get("amount") or (taxable + line_cgst + line_sgst + line_igst))

        line_items.append({
            "code": str(idx),
            "product_name": p_name,
            "hsn_sac": hsn,
            "size": spec,
            "quantity": qty,
            "unit": unit,
            "rate": rate,
            "discount": disc,
            "taxable": taxable,
            "gst_rate": gst_r,
            "cgst": line_cgst,
            "sgst": line_sgst,
            "igst": line_igst,
            "amount": amt
        })

    subtotal = float(data.get("subtotal") or subtotal_calc)
    discount = float(data.get("discount") or 0.0)
    taxable_amount = max(0.0, subtotal - discount)

    is_intra_state = (place_of_supply.lower() == comp_state.lower())

    cgst_rate = float(data.get("cgst_rate") if data.get("cgst_rate") is not None else (9.0 if is_intra_state else 0.0))
    sgst_rate = float(data.get("sgst_rate") if data.get("sgst_rate") is not None else (9.0 if is_intra_state else 0.0))
    igst_rate = float(data.get("igst_rate") if data.get("igst_rate") is not None else (0.0 if is_intra_state else 18.0))

    cgst_amount = round(taxable_amount * (cgst_rate / 100.0), 2) if is_intra_state else 0.0
    sgst_amount = round(taxable_amount * (sgst_rate / 100.0), 2) if is_intra_state else 0.0
    igst_amount = round(taxable_amount * (igst_rate / 100.0), 2) if not is_intra_state else 0.0

    freight = float(data.get("freight") or data.get("sh_freight") or 0.0)
    round_off = float(data.get("round_off") or 0.0)

    grand_total_calc = round(taxable_amount + cgst_amount + sgst_amount + igst_amount + freight + round_off, 2)
    saved_grand_total = float(data.get("grand_total") or 0.0)
    grand_total = saved_grand_total if saved_grand_total > 0 else grand_total_calc

    receipt_meta = {}
    if doc_type == "payment_receipt":
        receipt_meta = {
            "payment_mode": data.get("payment_mode") or "Bank Transfer",
            "ref_number": data.get("ref_number") or data.get("transaction_ref") or "—",
            "amount_received": float(data.get("amount_received") or grand_total),
            "payment_date": data.get("payment_date") or inv_date
        }

    notes = (data.get("notes") or "").strip()
    terms = (data.get("terms") or "").strip()

    return {
        "doc_type": doc_type,
        "title_text": title_text,
        "seller": {
            "name": comp_name,
            "address": comp_addr,
            "phone": comp_phone,
            "email": comp_email,
            "website": comp_website,
            "gstin": comp_gst,
            "state": comp_state,
            "state_code": comp_state_code,
            "bank_name": comp_bank_name,
            "bank_account": comp_bank_acc,
            "ifsc": comp_ifsc,
            "branch": comp_branch,
            "logo_bytes": logo_bytes
        },
        "buyer": {
            "name": c_name,
            "address": c_addr,
            "phone": c_phone,
            "email": c_email,
            "gstin": c_gstin,
            "state": c_state,
            "sol_id": c_sol_id,
            "consumer": c_consumer
        },
        "invoiceDetails": {
            "number": inv_num,
            "date": inv_date,
            "due_date": due_date,
            "place_of_supply": place_of_supply,
            "reverse_charge": reverse_charge,
            "original_invoice_number": orig_inv_num,
            "reason": reason,
            "payment_terms": payment_terms
        },
        "shipping": {
            "is_intra_state": is_intra_state
        },
        "lineItems": line_items,
        "financials": {
            "subtotal": subtotal,
            "discount": discount,
            "taxable_amount": taxable_amount,
            "cgst_rate": cgst_rate,
            "cgst_amount": cgst_amount,
            "sgst_rate": sgst_rate,
            "sgst_amount": sgst_amount,
            "igst_rate": igst_rate,
            "igst_amount": igst_amount,
            "freight": freight,
            "round_off": round_off,
            "grand_total": grand_total
        },
        "receipt_meta": receipt_meta,
        "notes": notes,
        "terms": terms
    }


def make_sales_doc_canvas(company: dict, doc_type: str = "tax_invoice"):
    comp_name = (company.get("company_name") or company.get("name") or company.get("legal_business_name") or "").strip()
    mobile = (company.get("mobile") or company.get("mobile_number") or company.get("phone") or company.get("phone_number") or "").strip()
    email = (company.get("email") or "").strip()
    website = (company.get("website") or "").strip()
    gstin = (company.get("gst_number") or company.get("gstin") or company.get("gst") or "").strip()
    addr = (company.get("address") or company.get("address_line_1") or company.get("office_address") or "").strip()
    city = (company.get("city") or "").strip()
    state = (company.get("state") or "").strip()
    pincode = (company.get("pincode") or "").strip()

    full_addr_parts = [addr]
    if city: full_addr_parts.append(city)
    if state: full_addr_parts.append(state)
    if pincode: full_addr_parts.append(f"- {pincode}")
    full_addr = ", ".join(p for p in full_addr_parts if p).replace(", -", " -")

    footer_items = []
    if comp_name: footer_items.append(comp_name)
    if mobile: footer_items.append(f"Ph: {mobile}")
    if email: footer_items.append(f"Email: {email}")
    if website: footer_items.append(f"Web: {website}")
    if gstin: footer_items.append(f"GSTIN: {gstin}")

    line1 = " | ".join(footer_items)
    line2 = f"Office: {full_addr}" if full_addr else ""

    class SalesDocCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            getattr(self, '_startPage')()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            # Divider line
            self.setStrokeColor(colors.HexColor('#1e3a8a'))
            self.setLineWidth(0.8)
            self.line(1.2 * cm, 1.35 * cm, 21.0 * cm - 1.2 * cm, 1.35 * cm)

            self.setFont("Helvetica-Bold", 7.5)
            self.setFillColor(colors.HexColor('#1e3a8a'))
            if line1:
                self.drawString(1.2 * cm, 0.95 * cm, line1[:115])

            self.setFont("Helvetica", 7.0)
            self.setFillColor(colors.HexColor('#475569'))
            if line2:
                self.drawString(1.2 * cm, 0.60 * cm, line2[:115])

            self.setFont("Helvetica-Bold", 7.5)
            self.setFillColor(colors.HexColor('#334155'))
            page_num = getattr(self, '_pageNumber', 1)
            self.drawRightString(21.0 * cm - 1.2 * cm, 0.60 * cm, f"Page {page_num} of {page_count}")
            self.restoreState()

    return SalesDocCanvas


def generate_invoice_pdf(data: dict, company: dict) -> bytes:
    """Generate professional PDF for Tax Invoice, Proforma, Payment Receipt, Credit Note, Debit Note."""
    norm = _normalize_invoice_document_data(data, company)
    doc_type = norm["doc_type"]
    title_text = norm["title_text"]
    seller = norm["seller"]
    buyer = norm["buyer"]
    inv = norm["invoiceDetails"]
    shipping = norm["shipping"]
    items = norm["lineItems"]
    fin = norm["financials"]
    receipt_meta = norm["receipt_meta"]
    notes = norm["notes"]
    terms = norm["terms"]

    buf = BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.8 * cm
    )
    story = []

    # 1. HEADER (Logo/Company Left, Title/Meta Right)
    logo_bytes = seller.get("logo_bytes")
    logo_d = None
    if logo_bytes:
        try:
            from PIL import Image as PILImage
            img = PILImage.open(BytesIO(logo_bytes))
            orig_w, orig_h = img.size
            if orig_w > 0 and orig_h > 0:
                max_w, max_h = 11.0 * cm, 3.65 * cm
                aspect = orig_w / float(orig_h or 1)
                if orig_w > orig_h:
                    target_w = max_w
                    target_h = max_w / aspect
                    if target_h > max_h:
                        target_h = max_h
                        target_w = max_h * aspect
                else:
                    target_h = max_h
                    target_w = max_h * aspect
                    if target_w > max_w:
                        target_w = max_w
                        target_h = max_w / aspect
                res_buf = BytesIO()
                img.save(res_buf, format="PNG")
                logo_d = RLImage(BytesIO(res_buf.getvalue()), width=target_w, height=target_h)
        except Exception:
            logo_d = None

    hdr_left_flow = []
    if logo_d:
        hdr_left_flow.append(logo_d)
        hdr_left_flow.append(Spacer(1, 0.1 * cm))

    comp_title_style = ParagraphStyle('inv_comp_title', parent=styles['Normal'], fontSize=13, leading=15, textColor=colors.HexColor('#1e3a8a'), fontName='Helvetica-Bold')
    comp_gst_style = ParagraphStyle('inv_comp_gst', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#0f172a'), fontName='Helvetica-Bold')
    comp_sub_style = ParagraphStyle('inv_comp_sub', parent=styles['Normal'], fontSize=8, leading=10, textColor=colors.HexColor('#475569'))

    hdr_left_flow.append(Paragraph(seller["name"], comp_title_style))
    if seller["gstin"]:
        hdr_left_flow.append(Paragraph(f"GSTIN: {seller['gstin']}", comp_gst_style))
    if seller["address"]:
        hdr_left_flow.append(Paragraph(seller["address"], comp_sub_style))

    inv_title_style = ParagraphStyle('inv_title', parent=styles['Normal'], fontSize=16, leading=18, textColor=colors.HexColor('#1e3a8a'), fontName='Helvetica-Bold', alignment=2)
    hdr_right_flow: list[Any] = [Paragraph(title_text, inv_title_style), Spacer(1, 0.2 * cm)]

    meta_table_data = [
        [Paragraph(f"<b>{title_text} NO.</b>", BOLD_SMALL), Paragraph(inv["number"], SMALL)],
        [Paragraph("<b>DATE</b>", BOLD_SMALL), Paragraph(inv["date"], SMALL)],
    ]
    if inv["due_date"]:
        meta_table_data.append([Paragraph("<b>DUE DATE</b>", BOLD_SMALL), Paragraph(inv["due_date"], SMALL)])
    if inv["place_of_supply"]:
        meta_table_data.append([Paragraph("<b>PLACE OF SUPPLY</b>", BOLD_SMALL), Paragraph(inv["place_of_supply"], SMALL)])
    if inv["original_invoice_number"]:
        meta_table_data.append([Paragraph("<b>ORIGINAL INVOICE</b>", BOLD_SMALL), Paragraph(inv["original_invoice_number"], SMALL)])

    meta_table = Table(meta_table_data, colWidths=[3.2 * cm, 4.2 * cm])
    meta_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#f8fafc')),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ('LEFTPADDING', (0,0), (-1,-1), 5),
        ('RIGHTPADDING', (0,0), (-1,-1), 5),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    hdr_right_flow.append(meta_table)

    header_table = Table([[hdr_left_flow, hdr_right_flow]], colWidths=[11.0 * cm, 7.6 * cm])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ('TOPPADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 0.4 * cm))

    # 2. SELLER & BUYER DETAILS (2 COLUMNS)
    s_body = f"<b>{seller['name']}</b><br/>"
    if seller['address']: s_body += f"{seller['address']}<br/>"
    if seller['phone']: s_body += f"Phone: {seller['phone']}<br/>"
    if seller['email']: s_body += f"Email: {seller['email']}<br/>"
    if seller['gstin']: s_body += f"GSTIN: {seller['gstin']} (State Code: {seller['state_code']})"

    b_body = f"<b>{buyer['name']}</b><br/>"
    if buyer['address']: b_body += f"{buyer['address']}<br/>"
    if buyer['phone']: b_body += f"Phone: {buyer['phone']}<br/>"
    if buyer['email']: b_body += f"Email: {buyer['email']}<br/>"
    if buyer['gstin']: b_body += f"GSTIN: {buyer['gstin']}<br/>"
    if buyer['sol_id']: b_body += f"SOL ID: {buyer['sol_id']}"

    party_header_style = ParagraphStyle('inv_p_hdr', parent=styles['Normal'], fontSize=9, leading=11, textColor=colors.HexColor('#ffffff'), fontName='Helvetica-Bold')
    party_body_style = ParagraphStyle('inv_p_bdy', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#0f172a'))

    party_table_data = [
        [Paragraph("SUPPLIER / SELLER DETAILS", party_header_style), Paragraph("BUYER / BILL TO", party_header_style)],
        [Paragraph(s_body, party_body_style), Paragraph(b_body, party_body_style)]
    ]
    party_table = Table(party_table_data, colWidths=[9.1 * cm, 9.1 * cm])
    party_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e3a8a')),
        ('BACKGROUND', (0,1), (-1,1), colors.HexColor('#ffffff')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(party_table)
    story.append(Spacer(1, 0.3 * cm))

    # 3. LINE ITEMS TABLE
    th_style = ParagraphStyle('inv_th', parent=styles['Normal'], fontSize=8, leading=10, textColor=colors.HexColor('#ffffff'), fontName='Helvetica-Bold')
    th_c = ParagraphStyle('inv_th_c', parent=th_style, alignment=1)
    th_r = ParagraphStyle('inv_th_r', parent=th_style, alignment=2)

    item_headers = [
        Paragraph("S.NO", th_c),
        Paragraph("DESCRIPTION OF GOODS / SERVICES", th_style),
        Paragraph("HSN/SAC", th_c),
        Paragraph("QTY", th_c),
        Paragraph("UNIT", th_c),
        Paragraph("RATE", th_r),
        Paragraph("TAXABLE", th_r),
        Paragraph("GST %", th_c),
        Paragraph("AMOUNT", th_r),
    ]
    item_rows = [item_headers]

    tb_style = ParagraphStyle('inv_tb', parent=styles['Normal'], fontSize=8, leading=10, textColor=colors.HexColor('#0f172a'))
    tb_c = ParagraphStyle('inv_tb_c', parent=tb_style, alignment=1)
    tb_r = ParagraphStyle('inv_tb_r', parent=tb_style, alignment=2)

    for it in items:
        p_name = it["product_name"]
        if it["size"] and it["size"] not in p_name:
            p_name += f" ({it['size']})"

        item_rows.append([
            Paragraph(it["code"], tb_c),
            Paragraph(p_name, tb_style),
            Paragraph(it["hsn_sac"] or "—", tb_c),
            Paragraph(f"{it['quantity']:g}", tb_c),
            Paragraph(it["unit"], tb_c),
            Paragraph(_format_currency(it["rate"]), tb_r),
            Paragraph(_format_currency(it["taxable"]), tb_r),
            Paragraph(f"{it['gst_rate']:g}%", tb_c),
            Paragraph(_format_currency(it["amount"]), tb_r),
        ])

    items_table = Table(item_rows, colWidths=[0.9 * cm, 5.5 * cm, 1.6 * cm, 1.1 * cm, 1.1 * cm, 2.0 * cm, 2.2 * cm, 1.4 * cm, 2.8 * cm], repeatRows=1)
    items_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e3a8a')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING', (0,0), (-1,-1), 3),
        ('RIGHTPADDING', (0,0), (-1,-1), 3),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(items_table)
    story.append(Spacer(1, 0.3 * cm))

    # 4. BOTTOM SECTION (BANK & TERMS LEFT, TOTALS RIGHT)
    left_notes_text = ""
    if seller["bank_name"] and seller["bank_account"]:
        left_notes_text += f"<b>BANK DETAILS FOR PAYMENT</b><br/>Bank: {seller['bank_name']}<br/>A/C No: {seller['bank_account']}<br/>IFSC: {seller['ifsc']}<br/><br/>"
    if notes:
        left_notes_text += f"<b>NOTES & INSTRUCTION</b><br/>{notes.replace(chr(10), '<br/>')}<br/><br/>"
    if terms:
        left_notes_text += f"<b>TERMS & CONDITIONS</b><br/>{terms.replace(chr(10), '<br/>')}"

    notes_p = Paragraph(left_notes_text or "Thank you for your business!", ParagraphStyle('inv_notes_p', parent=styles['Normal'], fontSize=8, leading=11, textColor=colors.HexColor('#1e293b')))

    tot_lbl = ParagraphStyle('inv_tot_lbl', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#334155'), fontName='Helvetica-Bold')
    tot_val = ParagraphStyle('inv_tot_v', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#0f172a'), alignment=2)
    tot_grand_lbl = ParagraphStyle('inv_tot_g_lbl', parent=styles['Normal'], fontSize=9.5, leading=12, textColor=colors.HexColor('#1e3a8a'), fontName='Helvetica-Bold')
    tot_grand_val = ParagraphStyle('inv_tot_g_v', parent=styles['Normal'], fontSize=10, leading=12, textColor=colors.HexColor('#1e3a8a'), fontName='Helvetica-Bold', alignment=2)

    totals_rows = [
        [Paragraph("SUBTOTAL", tot_lbl), Paragraph(_format_currency(fin["subtotal"]), tot_val)]
    ]
    if fin["discount"] > 0:
        totals_rows.append([Paragraph("DISCOUNT", tot_lbl), Paragraph(_format_currency(fin["discount"]), tot_val)])
    totals_rows.append([Paragraph("TAXABLE AMOUNT", tot_lbl), Paragraph(_format_currency(fin["taxable_amount"]), tot_val)])

    if shipping["is_intra_state"]:
        totals_rows.append([Paragraph(f"CGST ({fin['cgst_rate']}%)", tot_lbl), Paragraph(_format_currency(fin["cgst_amount"]), tot_val)])
        totals_rows.append([Paragraph(f"SGST ({fin['sgst_rate']}%)", tot_lbl), Paragraph(_format_currency(fin["sgst_amount"]), tot_val)])
    else:
        totals_rows.append([Paragraph(f"IGST ({fin['igst_rate']}%)", tot_lbl), Paragraph(_format_currency(fin["igst_amount"]), tot_val)])

    if fin["freight"] > 0:
        totals_rows.append([Paragraph("FREIGHT / S&H", tot_lbl), Paragraph(_format_currency(fin["freight"]), tot_val)])
    if fin["round_off"] != 0:
        totals_rows.append([Paragraph("ROUND OFF", tot_lbl), Paragraph(_format_currency(fin["round_off"]), tot_val)])

    totals_rows.append([Paragraph("GRAND TOTAL", tot_grand_lbl), Paragraph(_format_currency(fin["grand_total"]), tot_grand_val)])

    totals_table = Table(totals_rows, colWidths=[4.4 * cm, 3.7 * cm])
    totals_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#eff6ff')),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING', (0,0), (-1,-1), 5),
        ('RIGHTPADDING', (0,0), (-1,-1), 5),
    ]))

    bottom_table = Table([[notes_p, totals_table]], colWidths=[10.3 * cm, 8.3 * cm])
    bottom_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ('TOPPADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(bottom_table)
    story.append(Spacer(1, 0.4 * cm))

    # 5. AMOUNT IN WORDS & SIGNATURE BLOCK
    words_p = Paragraph(f"<b>Amount in Words:</b> {_amount_to_words(fin['grand_total'])}", ParagraphStyle('inv_words', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#0f172a')))
    story.append(words_p)
    story.append(Spacer(1, 0.4 * cm))

    sig_l = Paragraph("<b>Customer Signature</b><br/><br/><br/>_______________________", ParagraphStyle('sig_l', parent=styles['Normal'], fontSize=8.5, alignment=0))
    sig_r = Paragraph(f"<b>For {seller['name']}</b><br/><br/><br/>Authorized Signatory", ParagraphStyle('sig_r', parent=styles['Normal'], fontSize=8.5, alignment=2))
    sig_table = Table([[sig_l, sig_r]], colWidths=[9.1 * cm, 9.1 * cm])
    sig_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(KeepTogether([sig_table]))

    if doc_type == "proforma":
        story.append(Spacer(1, 0.2 * cm))
        story.append(Paragraph("<font size='9' color='#d97706'><b>THIS IS A PROFORMA INVOICE — NOT A TAX INVOICE</b></font>", ParagraphStyle('pf_not_tax', parent=styles['Normal'], alignment=1)))

    pdf.build(story, canvasmaker=make_sales_doc_canvas(company, doc_type))
    return buf.getvalue()


def generate_invoice_docx(data: dict, company: dict) -> bytes:
    """Generate faithful Word (.docx) document for Tax Invoice, Proforma, Receipt, Credit/Debit Note."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    norm = _normalize_invoice_document_data(data, company)
    doc_type = norm["doc_type"]
    title_text = norm["title_text"]
    seller = norm["seller"]
    buyer = norm["buyer"]
    inv = norm["invoiceDetails"]
    shipping = norm["shipping"]
    items = norm["lineItems"]
    fin = norm["financials"]
    notes = norm["notes"]
    terms = norm["terms"]

    doc = _build_docx_document(left_cm=1.2, right_cm=1.2, top_cm=1.0, bottom_cm=1.5)

    # 1. HEADER (Logo/Company Left, Title/Meta Right)
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.style = 'Table Grid'
    _remove_tbl_borders(hdr_table)

    cell_l = hdr_table.rows[0].cells[0]
    cell_r = hdr_table.rows[0].cells[1]

    logo_bytes = seller.get("logo_bytes")
    if logo_bytes:
        try:
            p_logo = cell_l.paragraphs[0]
            p_logo.add_run().add_picture(BytesIO(logo_bytes), width=Inches(1.8))
        except Exception:
            pass

    p_comp = cell_l.add_paragraph() if logo_bytes else cell_l.paragraphs[0]
    r_comp = p_comp.add_run(seller["name"])
    r_comp.bold = True
    r_comp.font.size = Pt(13)
    r_comp.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)

    if seller["gstin"]:
        p_gst = cell_l.add_paragraph()
        r_gst = p_gst.add_run(f"GSTIN: {seller['gstin']}")
        r_gst.bold = True
        r_gst.font.size = Pt(9)

    if seller["address"]:
        p_addr = cell_l.add_paragraph()
        r_addr = p_addr.add_run(seller["address"])
        r_addr.font.size = Pt(8.5)
        r_addr.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    p_title = cell_r.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_title = p_title.add_run(f"{title_text}\n")
    r_title.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)

    meta_str = f"NO: {inv['number']}\nDATE: {inv['date']}"
    if inv["due_date"]: meta_str += f"\nDUE DATE: {inv['due_date']}"
    if inv["place_of_supply"]: meta_str += f"\nPLACE OF SUPPLY: {inv['place_of_supply']}"
    r_meta = p_title.add_run(meta_str)
    r_meta.font.size = Pt(9)
    r_meta.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph()

    # 2. PARTY SECTION (SELLER | BUYER)
    party_tbl = doc.add_table(rows=2, cols=2)
    party_tbl.style = 'Table Grid'

    c_sh = party_tbl.rows[0].cells[0]
    c_bh = party_tbl.rows[0].cells[1]
    c_sh.text = "SUPPLIER / SELLER"
    c_bh.text = "BUYER / BILL TO"
    _docx_set_cell_bg(c_sh, "1e3a8a")
    _docx_set_cell_bg(c_bh, "1e3a8a")
    for c in (c_sh, c_bh):
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    c_sb = party_tbl.rows[1].cells[0]
    c_bb = party_tbl.rows[1].cells[1]

    p_s = c_sb.paragraphs[0]
    rs_name = p_s.add_run(f"{seller['name']}\n")
    rs_name.bold = True
    rs_name.font.size = Pt(9.5)
    s_details = []
    if seller['address']: s_details.append(seller['address'])
    if seller['phone']: s_details.append(f"Phone: {seller['phone']}")
    if seller['email']: s_details.append(f"Email: {seller['email']}")
    if seller['gstin']: s_details.append(f"GSTIN: {seller['gstin']} (State Code: {seller['state_code']})")
    rs_body = p_s.add_run("\n".join(s_details))
    rs_body.font.size = Pt(8.5)

    p_b = c_bb.paragraphs[0]
    rb_name = p_b.add_run(f"{buyer['name']}\n")
    rb_name.bold = True
    rb_name.font.size = Pt(9.5)
    b_details = []
    if buyer['address']: b_details.append(buyer['address'])
    if buyer['phone']: b_details.append(f"Phone: {buyer['phone']}")
    if buyer['email']: b_details.append(f"Email: {buyer['email']}")
    if buyer['gstin']: b_details.append(f"GSTIN: {buyer['gstin']}")
    if buyer['sol_id']: b_details.append(f"SOL ID: {buyer['sol_id']}")
    rb_body = p_b.add_run("\n".join(b_details))
    rb_body.font.size = Pt(8.5)

    doc.add_paragraph()

    # 3. LINE ITEMS TABLE
    items_tbl = doc.add_table(rows=len(items) + 1, cols=9)
    items_tbl.style = 'Table Grid'

    i_headers = ["S.NO", "DESCRIPTION OF GOODS / SERVICES", "HSN/SAC", "QTY", "UNIT", "RATE (Rs.)", "TAXABLE (Rs.)", "GST %", "AMOUNT (Rs.)"]
    for i, h in enumerate(i_headers):
        cell = items_tbl.rows[0].cells[i]
        cell.text = h
        _docx_set_cell_bg(cell, "1e3a8a")
        for p in cell.paragraphs:
            if i in (0, 2, 3, 4, 7):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i in (5, 6, 8):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    for idx, item in enumerate(items, 1):
        row_cells = items_tbl.rows[idx].cells
        p_name = item["product_name"]
        if item["size"] and item["size"] not in p_name:
            p_name += f" ({item['size']})"

        row_cells[0].text = item["code"]
        row_cells[1].text = p_name
        row_cells[2].text = item["hsn_sac"] or "—"
        row_cells[3].text = f"{item['quantity']:g}"
        row_cells[4].text = item["unit"]
        row_cells[5].text = f"{item['rate']:,.2f}"
        row_cells[6].text = f"{item['taxable']:,.2f}"
        row_cells[7].text = f"{item['gst_rate']:g}%"
        row_cells[8].text = f"{item['amount']:,.2f}"

        for i, c in enumerate(row_cells):
            for p in c.paragraphs:
                if i in (0, 2, 3, 4, 7):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif i in (5, 6, 8):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.add_paragraph()

    # 4. BOTTOM SECTION (NOTES & TOTALS)
    bot_tbl = doc.add_table(rows=1, cols=2)
    bot_tbl.style = 'Table Grid'
    c_n = bot_tbl.rows[0].cells[0]
    c_t = bot_tbl.rows[0].cells[1]

    p_n = c_n.paragraphs[0]
    n_text = ""
    if seller["bank_name"] and seller["bank_account"]:
        n_text += f"BANK DETAILS FOR PAYMENT\nBank: {seller['bank_name']}\nA/C No: {seller['bank_account']}\nIFSC: {seller['ifsc']}\n\n"
    if notes:
        n_text += f"NOTES & INSTRUCTION\n{notes}\n\n"
    if terms:
        n_text += f"TERMS & CONDITIONS\n{terms}"
    rn_bdy = p_n.add_run(n_text or "Thank you for your business!")
    rn_bdy.font.size = Pt(8)

    totals_list = [
        ("SUBTOTAL", f"Rs. {fin['subtotal']:,.2f}"),
    ]
    if fin["discount"] > 0:
        totals_list.append(("DISCOUNT", f"Rs. {fin['discount']:,.2f}"))
    totals_list.append(("TAXABLE AMOUNT", f"Rs. {fin['taxable_amount']:,.2f}"))

    if shipping["is_intra_state"]:
        totals_list.append((f"CGST ({fin['cgst_rate']}%)", f"Rs. {fin['cgst_amount']:,.2f}"))
        totals_list.append((f"SGST ({fin['sgst_rate']}%)", f"Rs. {fin['sgst_amount']:,.2f}"))
    else:
        totals_list.append((f"IGST ({fin['igst_rate']}%)", f"Rs. {fin['igst_amount']:,.2f}"))

    if fin["freight"] > 0:
        totals_list.append(("FREIGHT / S&H", f"Rs. {fin['freight']:,.2f}"))
    if fin["round_off"] != 0:
        totals_list.append(("ROUND OFF", f"Rs. {fin['round_off']:,.2f}"))

    totals_list.append(("GRAND TOTAL", f"Rs. {fin['grand_total']:,.2f}"))

    for lbl, val in totals_list:
        p_row = c_t.add_paragraph()
        r_lbl = p_row.add_run(f"{lbl}: ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(8.5)
        r_val = p_row.add_run(val)
        r_val.font.size = Pt(8.5)
        if lbl == "GRAND TOTAL":
            r_lbl.font.size = Pt(9.5)
            r_lbl.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
            r_val.bold = True
            r_val.font.size = Pt(10)
            r_val.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)

    doc.add_paragraph()

    # 5. FOOTER & NOT TAX INVOICE
    p_ftr = doc.add_paragraph()
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ftr = p_ftr.add_run(f"Enquiries: {seller['email']} | Contact: {seller['phone']} | Registered Office: {seller['address']}")
    r_ftr.font.size = Pt(8)
    r_ftr.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    if doc_type == "proforma":
        p_pf = doc.add_paragraph()
        p_pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_pf = p_pf.add_run("THIS IS A PROFORMA INVOICE — NOT A TAX INVOICE")
        r_pf.bold = True
        r_pf.font.size = Pt(9)
        r_pf.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_document(doc_type: str, data: dict, company: dict) -> bytes:
    if doc_type in ("tax_invoice", "proforma", "payment_receipt", "credit_note", "debit_note"):
        return generate_invoice_pdf(data, company)
    if doc_type == "purchase_order":
        return generate_po_pdf(data, company)
    buf = BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=A4, leftMargin=1.2 * cm, rightMargin=1.2 * cm, topMargin=1.2 * cm, bottomMargin=1.8 * cm)
    
    prepared_by = data.get("prepared_by", "")
    show_owner = data.get("show_owner") is not False and str(data.get("show_owner")).lower() != "false"
    story: list = _header(company, prepared_by, show_owner)
    
    titles = {
        "quotation": "QUOTATION",
        "sales_order": "SALES ORDER",
        "tax_invoice": "TAX INVOICE",
        "delivery_bill": "DELIVERY BILL",
        "purchase_order": "PURCHASE ORDER",
        "purchase_bill": "PURCHASE BILL",
    }
    custom_title = data.get("custom_title")
    if doc_type == "tax_invoice" and custom_title:
        title_text = str(custom_title).strip().upper()
    else:
        title_text = titles.get(doc_type, doc_type.replace("_", " ").upper())
    story.append(Paragraph(title_text, H1))
    story.append(Spacer(1, 0.2 * cm))

    client = data.get("client") or {}
    vendor = data.get("vendor") or {}
    doc_num = data.get("document_number") or data.get("quote_number") or data.get("invoice_number") or data.get("challan_number") or data.get("po_number") or data.get("bill_number") or "DOC-001"
    doc_date = data.get("document_date") or data.get("quote_date") or data.get("invoice_date") or data.get("date") or data.get("bill_date") or datetime.now().strftime("%Y-%m-%d")

    details = [
        [Paragraph(f"<b>{title_text} No.</b>", BOLD_SMALL), Paragraph(str(doc_num), SMALL)],
        [Paragraph("<b>Date</b>", BOLD_SMALL), Paragraph(str(doc_date), SMALL)],
    ]
    if data.get("valid_till"):
        details.append([Paragraph("<b>Valid Till</b>", BOLD_SMALL), Paragraph(str(data.get("valid_till")), SMALL)])
    if data.get("delivery_date"):
        details.append([Paragraph("<b>Delivery Date</b>", BOLD_SMALL), Paragraph(str(data.get("delivery_date")), SMALL)])

    story.append(_table(details, col_widths=[5 * cm, 13 * cm]))
    story.append(Spacer(1, 0.3 * cm))

    # Party details header
    is_vendor_doc = doc_type in ("purchase_order", "purchase_bill")
    party_label = "Vendor Details" if is_vendor_doc else "Customer Details"
    party_obj = vendor if is_vendor_doc else client

    story.append(Paragraph(party_label, H2))
    if is_vendor_doc:
        story.append(_vendor_table(party_obj))
    else:
        story.append(_client_table(party_obj))
    story.append(Spacer(1, 0.3 * cm))

    # Optional Ship To details
    ship_to = data.get("ship_to")
    if ship_to and isinstance(ship_to, dict) and ship_to.get("address"):
        story.append(Paragraph("Ship To", H2))
        story.append(_client_table(ship_to))
        story.append(Spacer(1, 0.3 * cm))

    items = data.get("items") or []
    apply_gst = data.get("apply_gst", True)
    if doc_type == "delivery_bill":
        apply_gst = False
        
    story.append(_render_items_table(doc_type, items, data, apply_gst))
    story.append(Spacer(1, 0.3 * cm))

    subtotal = sum(max(0.0, float(item.get("quantity") or 0) * float(item.get("rate") or 0) - float(item.get("discount") or 0)) for item in items)
    if apply_gst:
        gst_total = sum(float(item.get("cgst") or 0) + float(item.get("sgst") or 0) + float(item.get("igst") or 0) for item in items)
        if gst_total == 0:
            gst_total = sum(max(0.0, float(item.get("quantity") or 0) * float(item.get("rate") or 0) - float(item.get("discount") or 0)) * float(item.get("gst") or 0) / 100 for item in items)
    else:
        gst_total = 0.0
    grand_total = subtotal + gst_total
    totals = {"subtotal": subtotal, "gst_total": gst_total, "grand_total": grand_total, "total": grand_total}

    show_amount = data.get("show_amount", True) if doc_type == "delivery_bill" else True

    if show_amount:
        story.append(_summary_table(doc_type, totals))
        story.append(Spacer(1, 0.3 * cm))
        story.append(Paragraph("Amount in Words", H2))
        story.append(Paragraph(_amount_to_words(totals.get("grand_total") or totals.get("total") or 0), BODY))
        story.append(Spacer(1, 0.3 * cm))

    product_details = (data.get("product_details") or "").strip()
    if doc_type == "quotation" and product_details:
        product_details_heading = (data.get("product_details_heading") or "").strip() or "Product Details"
        story.append(Paragraph(product_details_heading, H2))
        story.append(Paragraph(product_details.replace("\n", "<br/>"), BODY))
        story.append(Spacer(1, 0.3 * cm))

    notes = data.get("notes") or ""
    if notes:
        story.append(Paragraph("Notes & Instructions", H2))
        story.append(Paragraph(notes.replace("\n", "<br/>"), BODY))
        story.append(Spacer(1, 0.3 * cm))
        
    terms = data.get("terms") or ""
    if terms:
        story.append(Paragraph("Terms & Conditions", H2))
        story.append(Paragraph(terms.replace("\n", "<br/>"), BODY))
        story.append(Spacer(1, 0.3 * cm))

    SIG_LEFT = ParagraphStyle('sig_left', parent=styles['BodyText'], fontSize=9, leading=13, textColor=colors.HexColor('#1f2937'), alignment=0)
    SIG_RIGHT = ParagraphStyle('sig_right', parent=styles['BodyText'], fontSize=9, leading=13, textColor=colors.HexColor('#1f2937'), alignment=2)

    sig_table = Table([
        [Paragraph("<b>Receiver Signature</b><br/><br/>_______________________", SIG_LEFT), Paragraph("<b>Authorized Signature</b><br/><br/>_______________________", SIG_RIGHT)]
    ], colWidths=[9.0 * cm, 9.0 * cm])
    sig_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
    ]))
    story.append(KeepTogether([sig_table]))

    # Legal notice at bottom: THIS IS NOT TAX INVOICE ! (Omitted for Quotation, Tax Invoice, Delivery Challan, and Purchase Order)
    if data.get("show_not_tax_invoice", False) and doc_type not in ("quotation", "tax_invoice", "delivery_bill", "purchase_order"):
        story.append(Spacer(1, 0.4 * cm))
        not_tax_p = Paragraph("<font size='10' color='#dc2626'><b>THIS IS NOT TAX INVOICE !</b></font>", ParagraphStyle('not_tax', parent=styles['BodyText'], alignment=1))
        story.append(not_tax_p)

    pdf.build(story, canvasmaker=make_sales_doc_canvas(company, doc_type))
    return buf.getvalue()



def make_wcr_canvas(company: dict):
    company_name = (company.get("company_name") or company.get("name") or company.get("legal_business_name") or "").strip()
    owner_name = (company.get("owner_name") or company.get("proprietor_name") or company.get("authorized_signatory") or company.get("manager_name") or "").strip()
    mobile = (company.get("mobile") or company.get("mobile_number") or company.get("phone") or company.get("phone_number") or "").strip()
    addr = (company.get("address") or company.get("address_line_1") or company.get("office_address") or company.get("registered_address") or "").strip()
    city = (company.get("city") or "").strip()
    state = (company.get("state") or "").strip()
    pincode = (company.get("pincode") or "").strip()

    full_addr_parts = [addr]
    if city:
        full_addr_parts.append(city)
    if state:
        full_addr_parts.append(state)
    if pincode:
        full_addr_parts.append(f"- {pincode}")
    full_addr = ", ".join(p for p in full_addr_parts if p).replace(", -", " -")

    line1 = f"OFFICE :- {full_addr}" if full_addr else ""
    line2_parts = []
    if mobile:
        line2_parts.append(f"PHONE : {mobile}")
    if owner_name:
        line2_parts.append(owner_name)
    line2 = " ".join(line2_parts)

    class WCRCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            getattr(self, '_startPage')()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            # Bottom Divider line
            self.setStrokeColor(colors.HexColor('#9333ea'))
            self.setLineWidth(1.2)
            self.line(1.2 * cm, 1.4 * cm, 21.0 * cm - 1.2 * cm, 1.4 * cm)
            
            self.setFont("Helvetica-Bold", 7.5)
            self.setFillColor(colors.HexColor('#2563eb'))
            
            if line1:
                self.drawString(1.2 * cm, 1.0 * cm, line1)
            if line2:
                self.drawString(1.2 * cm, 0.65 * cm, line2)
            
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor('#475569'))
            page_num = getattr(self, '_pageNumber', 1)
            self.drawRightString(21.0 * cm - 1.2 * cm, 0.65 * cm, f"Page {page_num} of {page_count}")
            self.restoreState()

    return WCRCanvas



def generate_wcr_pdf(client: dict, company: dict) -> bytes:
    # 1. Validation Check
    missing_fields = []
    if not (client.get("full_name") or client.get("name")):
        missing_fields.append("Consumer Name")
    if not client.get("consumer_number"):
        missing_fields.append("Consumer Number")
    if not (client.get("system_kw") or client.get("capacity")):
        missing_fields.append("Solar System Capacity")
    if not company.get("company_name"):
        missing_fields.append("Company Name")
    if not (company.get("gst_number") or company.get("gst")):
        missing_fields.append("Company GST Number")
        
    if missing_fields:
        from fastapi import HTTPException
        raise HTTPException(
            status_code=400,
            detail=f"Missing required data for WCR: {', '.join(missing_fields)}. Please update client/company details before generating."
        )

    buf = BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.8 * cm
    )
    story = []

    company_name = (company.get('company_name') or company.get('name') or company.get('legal_business_name') or '').strip()
    gst_no = (company.get('gst_number') or company.get('gstin') or company.get('gst') or '').strip()

    # Styles
    STYLE_TITLE = ParagraphStyle('wcr_title', parent=styles['Normal'], fontSize=13, fontName='Helvetica-Bold', textColor=colors.HexColor('#0f172a'), alignment=1, spaceBefore=4, spaceAfter=8)
    STYLE_BODY_JUSTIFY = ParagraphStyle('wcr_body_j', parent=styles['Normal'], fontSize=9.5, fontName='Helvetica', textColor=colors.HexColor('#1e293b'), leading=14, alignment=4, spaceAfter=8)
    STYLE_VAL = ParagraphStyle('c_val', parent=styles['Normal'], fontSize=8.5, fontName='Helvetica', textColor=colors.HexColor('#0f172a'))

    # Header Builder
    def _build_header():
        logo_bytes = company.get("logo_bytes")
        logo_d = None
        if logo_bytes:
            try:
                from PIL import Image as PILImage
                img = PILImage.open(BytesIO(logo_bytes))
                img_w, img_h = img.size
                if img_w > 0 and img_h > 0:
                    aspect = img_h / float(img_w)
                    max_w = 6.5 * cm
                    max_h = 3.2 * cm
                    target_w = max_w
                    target_h = target_w * aspect
                    if target_h > max_h:
                        target_h = max_h
                        target_w = target_h / aspect
                    resample_filter = getattr(getattr(PILImage, "Resampling", PILImage), "LANCZOS", getattr(PILImage, "LANCZOS", 1))
                    resampled = img.resize((round(target_w * 4), round(target_h * 4)), resample_filter)
                    res_buf = BytesIO()
                    resampled.save(res_buf, format='PNG')
                    logo_d = RLImage(BytesIO(res_buf.getvalue()), width=target_w, height=target_h)
            except Exception:
                logo_d = None
        if not logo_d:
            logo_d = Spacer(6.5 * cm, 1.2 * cm)

        name_len = len(company_name)
        if name_len > 30:
            title_font_size = 12
            title_leading = 14
        elif name_len > 22:
            title_font_size = 13
            title_leading = 15
        elif name_len > 16:
            title_font_size = 14
            title_leading = 16
        else:
            title_font_size = 15
            title_leading = 17

        p_title = Paragraph(f"<b><font size='{title_font_size}' color='#1d4ed8'>{company_name.upper()}</font></b>", ParagraphStyle('wcr_hdr_title', parent=styles['Normal'], fontName='Helvetica-Bold', leading=title_leading))
        p_gst = Paragraph(f"<b><font size='9' color='#1d4ed8'>GST NO – {gst_no}</font></b>" if gst_no else "", ParagraphStyle('wcr_hdr_gst', parent=styles['Normal'], fontName='Helvetica-Bold', alignment=2, leading=14))
        
        t_hdr = Table([[logo_d, p_title, p_gst]], colWidths=[6.5 * cm, 7.5 * cm, 4.6 * cm])
        t_hdr.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ]))
        
        # Solid Blue Divider Line
        t_div = Table([[""]], colWidths=[18.6 * cm])
        t_div.setStyle(TableStyle([
            ('LINEABOVE', (0, 0), (-1, -1), 1.5, colors.HexColor('#1d4ed8')),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        return [t_hdr, Spacer(1, 0.1 * cm), t_div, Spacer(1, 0.2 * cm)]

    # Real-Time Data Extraction (No Hardcoded Fallbacks or Placeholders)
    stages_dict = dict(client.get("stages") or {})
    ob_dict = dict(stages_dict.get("onboarding_data") or {})

    client_name = (client.get('full_name') or client.get('name') or '').strip()
    consumer_num = str(client.get('consumer_number') or '').strip()
    client_addr = (client.get('address') or '').strip()
    city = (client.get('city') or '').strip()
    pincode = str(client.get('pincode') or '').strip()
    site_addr = f"{client_addr}{', ' + city if city else ''}{' - ' + pincode if pincode else ''}".strip(', -')
    
    category = (client.get('consumer_type') or client.get('consumer_category') or client.get('category') or ob_dict.get('consumer_type') or ob_dict.get('consumer_category') or ob_dict.get('category') or '').strip()
    section_no = str(client.get('section_number') or client.get('section_no') or ob_dict.get('section_number') or ob_dict.get('section_no') or '').strip()

    sol_kw = str(client.get('system_kw') or client.get('capacity') or '').strip()
    sol_kw_str = f"{sol_kw} KW" if sol_kw else ""
    sol_wp = str(client.get('panel_wattage') or client.get('panel_wp') or ob_dict.get('panel_wattage') or '').strip()
    sol_wp_str = f"{sol_wp} WP" if sol_wp else ""
    num_panels = str(client.get('num_panels') or client.get('panel_quantity') or ob_dict.get('num_panels') or '').strip()
    num_panels_str = f"{num_panels} NOS" if num_panels else ""
    panel_make = (client.get('panel_brand') or client.get('panel_make') or ob_dict.get('panel_brand') or ob_dict.get('panel_make') or '').strip()
    panel_tech = (client.get('panel_technology') or client.get('panel_tech') or ob_dict.get('panel_technology') or ob_dict.get('panel_tech') or '').strip()

    if sol_wp_str and panel_tech:
        sol_wp_tech_str = f"{sol_wp_str} / {panel_tech}"
    elif sol_wp_str:
        sol_wp_tech_str = sol_wp_str
    else:
        sol_wp_tech_str = panel_tech

    almm_model = str(client.get('almm_model_number') or sol_wp_tech_str).strip()

    inverter_list = _get_inverters_list(client)

    # 1. Unique Brand Names (Row 8 of WCR) - deduplicated, no "Multiple (See Project Details)"
    brands = []
    if inverter_list:
        for inv in inverter_list:
            b = str(inv.get("brand") or "").strip()
            if b and b not in brands:
                brands.append(b)
    if not brands:
        fallback_b = str(client.get('inverter_make') or client.get('inverter_brand') or ob_dict.get('inverter_make') or '').strip()
        if fallback_b and fallback_b != "Multiple (See Project Details)":
            brands = [fallback_b]

    inverter_make = ", ".join(brands) if brands else ""

    # 2. Serial Numbers (Row 21 of WCR) - all saved serial numbers, no placeholder
    all_serials = []
    if inverter_list:
        for inv in inverter_list:
            inv_s = inv.get("serials")
            if isinstance(inv_s, list) and inv_s:
                for s in inv_s:
                    s_str = str(s).strip()
                    if s_str and s_str not in all_serials:
                        all_serials.append(s_str)
            elif inv.get("serial"):
                s_raw = str(inv.get("serial")).strip()
                if s_raw and s_raw != "Multiple (See Project Details)":
                    for part in s_raw.split(","):
                        p_str = part.strip()
                        if p_str and p_str not in all_serials:
                            all_serials.append(p_str)

    if not all_serials:
        raw_sr = str(client.get('inverter_serial') or client.get('inverter_sn') or ob_dict.get('inverter_serial') or '').strip()
        if raw_sr and raw_sr != "Multiple (See Project Details)":
            for part in raw_sr.split(","):
                p_str = part.strip()
                if p_str and p_str not in all_serials:
                    all_serials.append(p_str)

    inverter_sr = ", ".join(all_serials) if all_serials else ""

    # 3. Total Inverter Capacity (Row 20 of WCR) - manual capacity field
    inverter_kw = str(client.get('inverter_capacity') or ob_dict.get('inverter_capacity') or '').strip()
    inverter_kw_str = f"{inverter_kw} KW" if (inverter_kw and "KW" not in inverter_kw.upper()) else inverter_kw
    inverter_year = str(client.get('inverter_year') or client.get('manufacturing_year') or client.get('year_of_manufacture') or ob_dict.get('inverter_year') or ob_dict.get('manufacturing_year') or ob_dict.get('year_of_manufacture') or '').strip()

    # --- PAGE 1: 3-Column Inspection Table ---
    for item in _build_header():
        story.append(item)

    story.append(Paragraph("<b>Work Completion Report for Solar Power Plant</b>", STYLE_TITLE))
    story.append(Spacer(1, 0.1 * cm))

    cell_hdr = lambda txt: Paragraph(f"<b><font size='9' color='#0f172a'>{txt}</font></b>", ParagraphStyle('c_hdr', parent=styles['Normal'], fontName='Helvetica-Bold', alignment=1))
    cell_lbl = lambda txt: Paragraph(f"<b><font size='8.5' color='#1e293b'>{txt}</font></b>", ParagraphStyle('c_lbl', parent=styles['Normal'], fontName='Helvetica-Bold'))
    cell_obs = lambda txt: Paragraph(f"<font size='8.5' color='#1e293b'>{txt}</font>", ParagraphStyle('c_obs', parent=styles['Normal']))
    cell_val = lambda txt: Paragraph(f"<b><font size='8.5' color='#0f172a'>{txt}</font></b>", ParagraphStyle('c_val', parent=styles['Normal'], alignment=1))
    cell_subhdr = lambda txt: Paragraph(f"<b><font size='9' color='#0f172a'>{txt}</font></b>", ParagraphStyle('c_subhdr', parent=styles['Normal'], fontName='Helvetica-Bold', alignment=1))

    table_data = [
        [cell_hdr("Sr.No"), cell_hdr("Component"), cell_hdr("Observation")],
        [cell_obs("1"), cell_lbl("Name"), cell_val(client_name)],
        [cell_obs("2"), cell_lbl("Consumer number"), cell_val(consumer_num)],
        [cell_obs("3"), cell_lbl("Site/Location with Complete Address"), cell_val(site_addr)],
        [cell_obs("4"), cell_lbl("Category: Govt/Private Sector"), cell_val(category)],
        [cell_obs("5"), cell_lbl("Section number"), cell_val(section_no)],
        [cell_obs("6"), cell_lbl("Sanctioned Capacity of solar PV system (KW) Installed"), cell_val(sol_kw_str)],
        ["", cell_lbl("Capacity of solar PV system (KW)"), cell_val(sol_kw_str)],
        [cell_subhdr("Specification of the Modules"), "", ""],
        [cell_obs("7"), cell_lbl("Make & Type of modules"), cell_val(panel_make)],
        ["", cell_lbl("ALMM Model Number"), cell_val(almm_model)],
        ["", cell_lbl("Wattage per module"), cell_val(sol_wp_str)],
        ["", cell_lbl("No. of Module"), cell_val(num_panels_str)],
        ["", cell_lbl("Total Capacity (KWP)"), cell_val(sol_kw_str)],
        ["", cell_lbl("Warrantee Details (Product + Performance)"), cell_val("12+15 YEARS" if sol_kw else "")],
        [cell_subhdr("PCU"), "", ""],
        [cell_obs("8"), cell_lbl("Make & Model number of Inverter"), cell_val(inverter_make)],
        ["", cell_lbl("Rating"), cell_val(inverter_kw_str)],
        ["", cell_lbl("Type of charge controller/ MPPT"), cell_val("MPPT" if inverter_make else "")],
        ["", cell_lbl("Capacity of Inverter"), cell_val(inverter_kw_str)],
        ["", cell_lbl("SR Number"), cell_val(inverter_sr)],
        ["", cell_lbl("Year of manufacturing"), cell_val(inverter_year)],
        [cell_subhdr("EARTHING & PROTACTION"), "", ""],
        [cell_obs("9"), cell_lbl("No of Separate Earthings with earth Resistance"), cell_val("NON_TRACKING" if sol_kw else "")],
        ["", cell_lbl("It is certified that the Earth Resistance measure in presence of Licensed Electrical Contractor/Supervisor and found in order i.e. < 5 Ohms as per MNRE OM Dtd. 07.06.24 for CFA Component."), cell_val("")],
        ["", cell_lbl("Lightening Arrester"), cell_val("Yes" if sol_kw else "")],
    ]

    t1 = Table(table_data, colWidths=[0.7 * cm, 8.8 * cm, 9.1 * cm])
    t1.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.6, colors.HexColor('#64748b')),
        ('SPAN', (0, 8), (2, 8)),
        ('BACKGROUND', (0, 8), (2, 8), colors.HexColor('#f1f5f9')),
        ('SPAN', (0, 15), (2, 15)),
        ('BACKGROUND', (0, 15), (2, 15), colors.HexColor('#f1f5f9')),
        ('SPAN', (0, 22), (2, 22)),
        ('BACKGROUND', (0, 22), (2, 22), colors.HexColor('#f1f5f9')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 2.5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
    ]))
    story.append(t1)
    story.append(Spacer(1, 0.4 * cm))

    # Clean Signature Page 1
    sign1 = Table([
        [
            Paragraph(f"<b>Authorized Signature [Vendor]</b><br/><br/><br/>For <b>{company_name.upper()}</b>", STYLE_VAL),
            Paragraph(f"<b>Consumer Signature</b><br/><br/><br/><b>{client_name}</b>", ParagraphStyle('sig_c1', parent=styles['Normal'], alignment=2))
        ]
    ], colWidths=[9.3 * cm, 9.3 * cm])
    sign1.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(sign1)
    story.append(PageBreak())

    # --- PAGE 2: DECLARATION & UNDERTAKING ---
    discom_code = client.get("discom_code") or client.get("discom") or "MSEDCL"
    for item in _build_header():
        story.append(item)

    story.append(Spacer(1, 0.3 * cm))
    p1_text = (
        f"We <b>{company_name}</b> [Vendor] & <b>{client_name}</b> [Consumer] bearing Consumer Number "
        f"<b>{consumer_num}</b> Ensured structural stability of installed solar power plant and obtained "
        f"requisite permissions from the concerned authority. If in future, by virtue of any means "
        f"due to collapsing or damage to the installed solar power plant, {discom_code} will not be held "
        f"responsible for any loss to property or human life, if any."
    )
    story.append(Paragraph(p1_text, STYLE_BODY_JUSTIFY))
    story.append(Spacer(1, 0.2 * cm))

    p2_text = (
        "This is to Certify above Installed Solar PV System is working properly with electrical safety & "
        "Islanding switch in case of any presence of backup inverter an arrangement should be made in "
        "such way the backup inverter supply should never be synchronized with solar inverter to avoid "
        "any electrical accident due to back feeding. We will be held responsible for non-working of "
        "islanding mechanism and back feed to the de-energized grid."
    )
    story.append(Paragraph(p2_text, STYLE_BODY_JUSTIFY))
    story.append(Spacer(1, 1.5 * cm))

    sign2 = Table([
        [
            Paragraph(f"<b>Authorized Signature [Vendor]</b><br/><br/><br/>For <b>{company_name.upper()}</b>", STYLE_VAL),
            Paragraph(f"<b>Consumer Signature</b><br/><br/><br/><b>{client_name}</b>", ParagraphStyle('sig_c2', parent=styles['Normal'], alignment=2))
        ]
    ], colWidths=[9.3 * cm, 9.3 * cm])
    sign2.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(sign2)
    story.append(PageBreak())

    # --- PAGE 3: GUARANTEE CERTIFICATE UNDERTAKING ---
    for item in _build_header():
        story.append(item)

    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("<b>Guarantee Certificate Undertaking to be submitted by VENDOR</b>", STYLE_TITLE))
    story.append(Spacer(1, 0.2 * cm))

    body3_text = (
        "The undersigned will provide services to the consumers for repairs/maintenance of the "
        "RTS plant free of cost for 5 years of the comprehensive Maintenance Contract (CMC) period "
        "from the date of commissioning of the plant. Nonperforming/under-performing system "
        "components will be replaced/repaired free of cost in the CMC period"
    )
    story.append(Paragraph(body3_text, STYLE_BODY_JUSTIFY))
    story.append(Spacer(1, 0.3 * cm))

    pan_num = str(client.get('pan_number') or client.get('pan_card_number') or ob_dict.get('pan_number') or ob_dict.get('pan_card_number') or '').strip()
    aadhaar_num = str(client.get('aadhaar') or client.get('aadhaar_number') or ob_dict.get('aadhaar') or ob_dict.get('aadhaar_number') or '').strip()

    if pan_num:
        id_title = "[ CONSUMER PAN CARD / IDENTITY VERIFICATION ]"
        id_label = "PAN CARD"
        id_detail = f"PAN Number: {pan_num}"
    elif aadhaar_num:
        id_title = "[ CONSUMER AADHAAR CARD / IDENTITY VERIFICATION ]"
        id_label = "ADHAR CARD"
        id_detail = f"Aadhar Number: {aadhaar_num}"
    else:
        id_title = "[ CONSUMER IDENTITY VERIFICATION ]"
        id_label = "IDENTITY CARD"
        id_detail = ""

    id_body_text = f"<b>Stamp & Seal</b><br/><br/><b>Identity Details of Consumer: - {id_label}</b>"
    if id_detail:
        id_body_text += f"<br/><b>{id_detail}</b>"

    id_box_data = [
        [Paragraph(f"<b>{id_title}</b>", ParagraphStyle('a_hdr', parent=styles['Normal'], alignment=1, fontSize=9, fontName='Helvetica-Bold', textColor=colors.HexColor('#1e3a8a')))],
        [Paragraph(id_body_text, ParagraphStyle('a_body', parent=styles['Normal'], fontSize=8.5, leading=13, textColor=colors.HexColor('#1e293b')))]
    ]
    id_table = Table(id_box_data, colWidths=[14 * cm])
    id_table.setStyle(TableStyle([
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#3b82f6')),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#eff6ff')),
        ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#f8fafc')),
        ('PADDING', (0, 0), (-1, -1), 8),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
    ]))
    story.append(id_table)
    story.append(Spacer(1, 1.0 * cm))

    sign3 = Table([
        [
            Paragraph(f"<b>Authorized Signature [Vendor]</b><br/><br/><br/>For <b>{company_name.upper()}</b>", STYLE_VAL),
            Paragraph(f"<b>Consumer Signature</b><br/><br/><br/><b>{client_name}</b>", ParagraphStyle('sig_c3', parent=styles['Normal'], alignment=2))
        ]
    ], colWidths=[9.3 * cm, 9.3 * cm])
    sign3.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(sign3)

    pdf.build(story, canvasmaker=make_wcr_canvas(company))
    return buf.getvalue()


def _build_sldr_drawing(sol_kw="5", sol_wp="540", num_panels="10", panel_make="GVP SOLAR", inverter_make="GROWATT", inverter_kw="5"):
    d = Drawing(490, 370)
    d.add(Rect(0, 0, 490, 370, fillColor=colors.HexColor("#ffffff"), strokeColor=colors.HexColor("#334155"), strokeWidth=1))  # type: ignore
    
    d.add(String(245, 355, "Grid Tied Solar Inverter System Electrical Single Line Diagram", fontName="Helvetica-Bold", fontSize=9.5, textAnchor="middle", fillColor=colors.HexColor("#1e293b")))

    # 1. PV Modules Array (Left Side: 2 Vertical Strings of Small Panel Rectangles with Junction Boxes)
    # Column 1 (x=22) & Column 2 (x=62)
    for col in range(2):
        cx = 22 + col * 40
        for row in range(4):
            ry = 285 - row * 43
            # Individual Solar Panel Rectangle
            d.add(Rect(cx, ry, 30, 38, fillColor=colors.HexColor("#f8fafc"), strokeColor=colors.HexColor("#0284c7"), strokeWidth=1))  # type: ignore
            # Internal Solar Cell Grid Lines (3x2 grid)
            d.add(Line(cx+10, ry, cx+10, ry+38, strokeColor=colors.HexColor("#cbd5e1"), strokeWidth=0.5))
            d.add(Line(cx+20, ry, cx+20, ry+38, strokeColor=colors.HexColor("#cbd5e1"), strokeWidth=0.5))
            d.add(Line(cx, ry+19, cx+30, ry+19, strokeColor=colors.HexColor("#cbd5e1"), strokeWidth=0.5))
            
            # Embedded Junction Box with + / - Terminals
            d.add(Rect(cx+7, ry+13, 16, 12, rx=1, ry=1, fillColor=colors.HexColor("#ffffff"), strokeColor=colors.HexColor("#0369a1"), strokeWidth=0.6))  # type: ignore
            d.add(String(cx+15, ry+19, "Junction", fontName="Helvetica-Bold", fontSize=4.5, textAnchor="middle", fillColor=colors.HexColor("#0369a1")))
            d.add(String(cx+15, ry+14, "box", fontName="Helvetica-Bold", fontSize=4.5, textAnchor="middle", fillColor=colors.HexColor("#0369a1")))
            d.add(String(cx+3, ry+31, "+", fontName="Helvetica-Bold", fontSize=5.5, fillColor=colors.HexColor("#dc2626")))
            d.add(String(cx+22, ry+31, "-", fontName="Helvetica-Bold", fontSize=5.5, fillColor=colors.HexColor("#1e293b")))

    # Interconnecting String Wiring
    d.add(Line(37, 323, 37, 156, strokeColor=colors.HexColor("#0284c7"), strokeWidth=1))
    d.add(Line(77, 323, 77, 156, strokeColor=colors.HexColor("#0284c7"), strokeWidth=1))

    d.add(String(57, 142, f"PV Array: {num_panels} x {sol_wp}Wp", fontName="Helvetica-Bold", fontSize=7.5, textAnchor="middle", fillColor=colors.HexColor("#0f172a")))

    # PV Ground Earth Wire (Green)
    d.add(Line(37, 140, 37, 50, strokeColor=colors.HexColor("#16a34a"), strokeWidth=1.2))
    d.add(Line(27, 50, 47, 50, strokeColor=colors.HexColor("#16a34a"), strokeWidth=1.2))
    d.add(Line(30, 46, 44, 46, strokeColor=colors.HexColor("#16a34a"), strokeWidth=1.2))
    d.add(Line(33, 42, 41, 42, strokeColor=colors.HexColor("#16a34a"), strokeWidth=1.2))
    d.add(String(37, 30, "Ground Earth", fontName="Helvetica", fontSize=7, textAnchor="middle", fillColor=colors.HexColor("#15803d")))

    # DC Output Conductors from PV Array (Red)
    d.add(Line(77, 156, 77, 75, strokeColor=colors.HexColor("#dc2626"), strokeWidth=1.5))
    d.add(Line(77, 75, 140, 75, strokeColor=colors.HexColor("#dc2626"), strokeWidth=1.5))
    d.add(String(108, 83, "DC240V", fontName="Helvetica-Bold", fontSize=7, textAnchor="middle", fillColor=colors.HexColor("#b91c1c")))
    d.add(String(108, 65, "8.31A", fontName="Helvetica-Bold", fontSize=6.5, textAnchor="middle", fillColor=colors.HexColor("#b91c1c")))

    # 2. DC Isolator Box
    d.add(Rect(140, 52, 55, 45, rx=3, ry=3, fillColor=colors.HexColor("#fef2f2"), strokeColor=colors.HexColor("#ef4444"), strokeWidth=1))  # type: ignore
    d.add(Line(150, 75, 180, 75, strokeColor=colors.HexColor("#dc2626"), strokeWidth=1.2))
    d.add(Circle(155, 75, 2.5, fillColor=colors.HexColor("#dc2626"), strokeColor=colors.HexColor("#dc2626")))
    d.add(Circle(175, 75, 2.5, fillColor=colors.HexColor("#dc2626"), strokeColor=colors.HexColor("#dc2626")))
    d.add(Line(155, 75, 172, 83, strokeColor=colors.HexColor("#dc2626"), strokeWidth=1.5))
    d.add(String(167, 38, "DC Isolator", fontName="Helvetica-Bold", fontSize=7.5, textAnchor="middle", fillColor=colors.HexColor("#991b1b")))
    d.add(String(167, 28, "B 31A", fontName="Helvetica", fontSize=6.5, textAnchor="middle", fillColor=colors.HexColor("#7f1d1d")))

    d.add(Line(195, 75, 225, 75, strokeColor=colors.HexColor("#dc2626"), strokeWidth=1.5))
    d.add(Line(225, 75, 225, 140, strokeColor=colors.HexColor("#dc2626"), strokeWidth=1.5))

    # 3. Grid Tied Solar Inverter Box (Enlarged Height & Spacing)
    d.add(Rect(200, 140, 130, 175, rx=8, ry=8, fillColor=colors.HexColor("#f8fafc"), strokeColor=colors.HexColor("#1e293b"), strokeWidth=1.5))  # type: ignore
    d.add(String(265, 302, "Grid Tied Solar Inverter", fontName="Helvetica-Bold", fontSize=8.5, textAnchor="middle", fillColor=colors.HexColor("#0f172a")))
    
    # Wi-Fi Monitor Plug
    d.add(Rect(286, 276, 36, 18, rx=3, ry=3, fillColor=colors.HexColor("#e0f2fe"), strokeColor=colors.HexColor("#0284c7"), strokeWidth=1))  # type: ignore
    d.add(String(304, 285, "Wi-Fi Plug", fontName="Helvetica-Bold", fontSize=5.5, textAnchor="middle", fillColor=colors.HexColor("#0369a1")))
    d.add(String(304, 279, "(Monitor)", fontName="Helvetica", fontSize=5, textAnchor="middle", fillColor=colors.HexColor("#0284c7")))
    d.add(String(304, 297, "((( Wi-Fi )))", fontName="Helvetica", fontSize=5.5, textAnchor="middle", fillColor=colors.HexColor("#0284c7")))

    # Converter Compartments (DC & AC)
    d.add(Rect(218, 165, 44, 44, fillColor=colors.HexColor("#ffffff"), strokeColor=colors.HexColor("#475569"), strokeWidth=1))  # type: ignore
    d.add(String(240, 183, "DC", fontName="Helvetica-Bold", fontSize=9.5, textAnchor="middle", fillColor=colors.HexColor("#334155")))
    
    d.add(Rect(218, 225, 44, 44, fillColor=colors.HexColor("#ffffff"), strokeColor=colors.HexColor("#475569"), strokeWidth=1))  # type: ignore
    d.add(String(240, 243, "AC", fontName="Helvetica-Bold", fontSize=9.5, textAnchor="middle", fillColor=colors.HexColor("#334155")))

    d.add(Line(240, 209, 240, 225, strokeColor=colors.HexColor("#2563eb"), strokeWidth=1.5))
    d.add(String(215, 146, "DC In", fontName="Helvetica-Bold", fontSize=6.5, fillColor=colors.HexColor("#64748b")))
    d.add(String(290, 146, "AC Out", fontName="Helvetica-Bold", fontSize=6.5, fillColor=colors.HexColor("#64748b")))
    d.add(String(265, 126, f"Make: {inverter_make} ({inverter_kw})", fontName="Helvetica-Bold", fontSize=7.5, textAnchor="middle", fillColor=colors.HexColor("#1e293b")))

    d.add(Line(330, 175, 355, 175, strokeColor=colors.HexColor("#2563eb"), strokeWidth=1.5))
    d.add(String(342, 182, "AC 230V", fontName="Helvetica-Bold", fontSize=6.5, textAnchor="middle", fillColor=colors.HexColor("#1d4ed8")))
    d.add(String(342, 164, "8.7A", fontName="Helvetica-Bold", fontSize=6.5, textAnchor="middle", fillColor=colors.HexColor("#1d4ed8")))

    # 4. AC Breaker Box
    d.add(Rect(355, 153, 44, 44, rx=3, ry=3, fillColor=colors.HexColor("#eff6ff"), strokeColor=colors.HexColor("#3b82f6"), strokeWidth=1))  # type: ignore
    d.add(Line(365, 175, 389, 175, strokeColor=colors.HexColor("#1d4ed8"), strokeWidth=1.2))
    d.add(Circle(368, 175, 2, fillColor=colors.HexColor("#1d4ed8"), strokeColor=colors.HexColor("#1d4ed8")))
    d.add(Circle(386, 175, 2, fillColor=colors.HexColor("#1d4ed8"), strokeColor=colors.HexColor("#1d4ed8")))
    d.add(Line(368, 175, 384, 182, strokeColor=colors.HexColor("#1d4ed8"), strokeWidth=1.5))
    d.add(String(377, 140, "AC Breaker", fontName="Helvetica-Bold", fontSize=7.5, textAnchor="middle", fillColor=colors.HexColor("#1e40af")))

    d.add(Line(399, 175, 420, 175, strokeColor=colors.HexColor("#2563eb"), strokeWidth=1.5))

    # 5. Main Distribution Panel & Utility Grid Meter (Taller Box)
    d.add(Rect(420, 90, 62, 235, rx=5, ry=5, fillColor=colors.HexColor("#f8fafc"), strokeColor=colors.HexColor("#0f172a"), strokeWidth=1.5))  # type: ignore
    d.add(String(451, 312, "Main Distribution", fontName="Helvetica-Bold", fontSize=6.5, textAnchor="middle", fillColor=colors.HexColor("#0f172a")))
    d.add(String(451, 303, "Panel", fontName="Helvetica-Bold", fontSize=6.5, textAnchor="middle", fillColor=colors.HexColor("#0f172a")))

    # Meter Box
    d.add(Rect(427, 245, 48, 48, fillColor=colors.HexColor("#ffffff"), strokeColor=colors.HexColor("#0284c7"), strokeWidth=1))  # type: ignore
    d.add(Circle(451, 273, 12, fillColor=colors.HexColor("#f0f9ff"), strokeColor=colors.HexColor("#0284c7"), strokeWidth=0.8))
    d.add(String(451, 270, "Meter", fontName="Helvetica-Bold", fontSize=8, textAnchor="middle", fillColor=colors.HexColor("#0369a1")))
    d.add(String(451, 252, "[ P14 ]", fontName="Helvetica", fontSize=6.5, textAnchor="middle", fillColor=colors.HexColor("#0284c7")))

    # Main Switch / PSE Box
    d.add(Rect(427, 155, 48, 38, fillColor=colors.HexColor("#ffffff"), strokeColor=colors.HexColor("#475569"), strokeWidth=1))  # type: ignore
    d.add(String(451, 174, "Main Switch", fontName="Helvetica-Bold", fontSize=6.5, textAnchor="middle", fillColor=colors.HexColor("#334155")))
    d.add(String(451, 163, "[ PSE ]", fontName="Helvetica-Bold", fontSize=6, textAnchor="middle", fillColor=colors.HexColor("#475569")))

    # Utility Grid Connection Line
    d.add(Line(475, 269, 488, 269, strokeColor=colors.HexColor("#16a34a"), strokeWidth=1.5))
    d.add(PolyLine([484, 273, 489, 269, 484, 265], strokeColor=colors.HexColor("#16a34a"), strokeWidth=1.5, fillColor=colors.HexColor("#16a34a")))
    d.add(String(451, 296, "To Utility Grid", fontName="Helvetica-Bold", fontSize=6.5, textAnchor="middle", fillColor=colors.HexColor("#15803d")))

    # Local Load Connection Line
    d.add(Line(475, 174, 488, 174, strokeColor=colors.HexColor("#d97706"), strokeWidth=1.5))
    d.add(PolyLine([484, 178, 489, 174, 484, 170], strokeColor=colors.HexColor("#d97706"), strokeWidth=1.5, fillColor=colors.HexColor("#d97706")))
    d.add(String(451, 144, "To Local Load", fontName="Helvetica-Bold", fontSize=6.5, textAnchor="middle", fillColor=colors.HexColor("#b45309")))

    # Main Panel Ground Earth Wire
    d.add(Line(451, 90, 451, 50, strokeColor=colors.HexColor("#16a34a"), strokeWidth=1.2))
    d.add(Line(441, 50, 461, 50, strokeColor=colors.HexColor("#16a34a"), strokeWidth=1.2))
    d.add(Line(444, 46, 458, 46, strokeColor=colors.HexColor("#16a34a"), strokeWidth=1.2))
    d.add(Line(447, 42, 455, 42, strokeColor=colors.HexColor("#16a34a"), strokeWidth=1.2))
    d.add(String(451, 30, "Ground Earth", fontName="Helvetica", fontSize=7, textAnchor="middle", fillColor=colors.HexColor("#15803d")))

    return d


def _get_inverters_list(client: dict) -> list[dict]:
    stages_dict = dict(client.get("stages") or {})
    ob_dict = dict(stages_dict.get("onboarding_data") or {})

    raw_inverters = client.get("inverters") or ob_dict.get("inverters")
    if isinstance(raw_inverters, list) and len(raw_inverters) > 0:
        cleaned = []
        for inv in raw_inverters:
            if isinstance(inv, dict):
                brand = str(inv.get("brand") or inv.get("make") or "").strip()
                model = str(inv.get("model") or "").strip()
                cap = str(inv.get("capacity") or "").strip()
                qty = str(inv.get("quantity") or inv.get("qty") or "1").strip()
                serials_raw = inv.get("serials")
                serials_list = []
                if isinstance(serials_raw, list):
                    serials_list = [str(s).strip() for s in serials_raw if str(s).strip()]
                single_serial = str(inv.get("serial") or inv.get("serial_number") or "").strip()
                if not serials_list and single_serial:
                    serials_list = [s.strip() for s in single_serial.split(",") if s.strip()]
                
                serial_display = ", ".join(serials_list) if serials_list else single_serial

                if brand or model or cap or serials_list or single_serial:
                    cleaned.append({
                        "brand": brand,
                        "model": model,
                        "capacity": cap,
                        "quantity": qty or "1",
                        "serials": serials_list,
                        "serial": serial_display
                    })
        if cleaned:
            return cleaned

    brand = str(client.get("inverter_make") or client.get("inverter_brand") or ob_dict.get("inverter_make") or "").strip()
    model = str(client.get("inverter_model") or ob_dict.get("inverter_model") or "").strip()
    cap = str(client.get("inverter_capacity") or ob_dict.get("inverter_capacity") or "").strip()
    raw_serial = str(client.get("inverter_serial") or client.get("inverter_sn") or ob_dict.get("inverter_serial") or "").strip()
    serials_list = [s.strip() for s in raw_serial.split(",") if s.strip()] if raw_serial else []

    if brand or cap or raw_serial or model:
        return [{
            "brand": brand,
            "model": model,
            "capacity": cap,
            "quantity": "1",
            "serials": serials_list,
            "serial": raw_serial
        }]
    return []


def _get_vendor_inverter_clause_text(client: dict) -> str:
    """
    Returns '<Inverter Brand/Make>, <Model/Capacity> model' string for Vendor Agreement.
    Example: 'UTL make, 3 kW model'
    """
    stages_dict = dict(client.get("stages") or {})
    ob_dict = dict(stages_dict.get("onboarding_data") or {})

    model = str(client.get("inverter_model") or ob_dict.get("inverter_model") or "").strip()
    brand = str(client.get("inverter_make") or client.get("inverter_brand") or ob_dict.get("inverter_make") or ob_dict.get("inverter_brand") or "").strip()
    cap   = str(client.get("inverter_capacity") or client.get("system_kw") or client.get("capacity") or ob_dict.get("inverter_capacity") or ob_dict.get("system_kw") or "").strip()

    if cap and not ("KW" in cap.upper() or "KWATT" in cap.upper()):
        cap = f"{cap} kW"

    brand_str = brand if ("make" in brand.lower() or not brand) else f"{brand} make"
    if model and cap:
        mod_cap_str = f"{model} {cap} model" if "model" not in model.lower() else f"{model} {cap}"
    elif model:
        mod_cap_str = f"{model} model" if "model" not in model.lower() else model
    elif cap:
        mod_cap_str = f"{cap} model"
    else:
        mod_cap_str = ""

    if brand_str and mod_cap_str:
        return f"{brand_str}, {mod_cap_str}"
    return brand_str or mod_cap_str




def generate_sldr_pdf(client: dict, company: dict) -> bytes:
    buf = BytesIO()

    def _draw_sldr_frame(canvas, doc):
        canvas.saveState()
        canvas.setLineWidth(1.5)
        canvas.setStrokeColor(colors.HexColor("#0f172a"))
        canvas.rect(1.0 * cm, 1.0 * cm, 19.0 * cm, 27.7 * cm)
        canvas.restoreState()

    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.8 * cm
    )
    story = []

    stages_dict = dict(client.get("stages") or {})
    ob_dict = dict(stages_dict.get("onboarding_data") or {})

    client_name = str(client.get('full_name') or client.get('name') or '').upper()
    consumer_num = str(client.get('consumer_number') or client.get('consumer_no') or '').upper()
    bu_num = str(client.get('bu_number') or client.get('bu') or '').upper()
    sol_kw = str(client.get('system_kw') or client.get('capacity') or '')
    sol_wp = str(client.get('panel_wattage') or '')
    num_panels = str(client.get('num_panels') or '')
    panel_make = (client.get('panel_brand') or client.get('panel_make') or '').upper()
    
    inverter_list = _get_inverters_list(client)
    first_inv = inverter_list[0] if inverter_list else {}
    inverter_make = (first_inv.get("brand") or client.get('inverter_make') or '').upper()
    
    raw_ob_cap = str(client.get('inverter_capacity') or ob_dict.get('inverter_capacity') or '').strip()
    if raw_ob_cap:
        inverter_kw = raw_ob_cap.upper() if "KW" in raw_ob_cap.upper() else f"{raw_ob_cap} KW"
    else:
        inverter_kw = f"{sol_kw} KW" if sol_kw else ""

    company_name = (company.get('company_name') or '').upper()

    STYLE_SLDR_TITLE = ParagraphStyle('sldr_t', parent=styles['Normal'], fontSize=15, fontName='Helvetica-Bold', alignment=1, spaceAfter=8, textColor=colors.HexColor('#0f172a'))
    STYLE_SLDR_META = ParagraphStyle('sldr_m', parent=styles['Normal'], fontSize=10, fontName='Helvetica-Bold', leading=15, textColor=colors.HexColor('#0f172a'))
    STYLE_TBL_HDR = ParagraphStyle('sldr_th', parent=styles['Normal'], fontSize=9, fontName='Helvetica-Bold', alignment=1, textColor=colors.HexColor('#0f172a'))
    STYLE_TBL_CELL = ParagraphStyle('sldr_tc', parent=styles['Normal'], fontSize=8.5, fontName='Helvetica-Bold', alignment=1, textColor=colors.HexColor('#1e293b'))
    STYLE_FTR = ParagraphStyle('sldr_ftr', parent=styles['Normal'], fontSize=8.5, fontName='Helvetica-Bold', textColor=colors.HexColor('#0f172a'))

    story.append(Paragraph("<u><b>SINGLE LINE DIAGRAM</b></u>", STYLE_SLDR_TITLE))
    story.append(Spacer(1, 0.15 * cm))

    meta_text = (
        f"<b>CONSUMER NAME :-</b> {client_name}<br/>"
        f"<b>CONSUMER NO.:-</b>{consumer_num} <b>B.U.:-</b>{bu_num}<br/>"
        f"<b>PROJECT:-</b> GCRT OF {sol_kw} KW"
    )
    story.append(Paragraph(meta_text, STYLE_SLDR_META))
    story.append(Spacer(1, 0.2 * cm))

    story.append(_build_sldr_drawing(sol_kw, sol_wp, num_panels, panel_make, inverter_make, inverter_kw))
    story.append(Spacer(1, 0.25 * cm))

    story.append(Paragraph("<b>TECHNICAL SPECIFICATIONS</b>", ParagraphStyle('tech_title', parent=styles['Normal'], fontSize=9.5, fontName='Helvetica-Bold', spaceAfter=4, textColor=colors.HexColor('#0f172a'))))

    tech_table_data = [
        [Paragraph("<b>PARAMETER</b>", STYLE_TBL_HDR), Paragraph("<b>SPECIFICATIONS</b>", STYLE_TBL_HDR), Paragraph("<b>MAKE</b>", STYLE_TBL_HDR), Paragraph("<b>KWP</b>", STYLE_TBL_HDR)],
        [Paragraph("PV MODULES", STYLE_TBL_CELL), Paragraph(f"{sol_wp} Wp X {num_panels} Nos", STYLE_TBL_CELL), Paragraph(panel_make, STYLE_TBL_CELL), Paragraph(f"{sol_kw} KW", STYLE_TBL_CELL)],
    ]

    if inverter_list:
        for idx, inv in enumerate(inverter_list):
            inv_label = f"INVERTER #{idx+1}" if len(inverter_list) > 1 else "INVERTER"
            c_val = inv.get("capacity") or ""
            cap_str = f"{c_val}" if "KW" in c_val.upper() else (f"{c_val} kW" if c_val else "")
            q_val = inv.get("quantity") or "1"
            spec_str = f"{cap_str} × {q_val} Nos" if cap_str else f"{q_val} Nos"
            make_str = f"{inv.get('brand','')} {inv.get('model','')}".strip().upper()
            kw_str = inverter_kw
            tech_table_data.append([
                Paragraph(inv_label, STYLE_TBL_CELL),
                Paragraph(spec_str, STYLE_TBL_CELL),
                Paragraph(make_str, STYLE_TBL_CELL),
                Paragraph(kw_str, STYLE_TBL_CELL)
            ])
    else:
        inverter_kw_display = f"{inverter_kw}" if "KW" in inverter_kw.upper() else (f"{inverter_kw} kW" if inverter_kw else "")
        tech_table_data.append([
            Paragraph("INVERTER", STYLE_TBL_CELL), Paragraph(f"{inverter_kw_display} × 1 Nos", STYLE_TBL_CELL), Paragraph(inverter_make, STYLE_TBL_CELL), Paragraph(f"{inverter_kw}", STYLE_TBL_CELL)
        ])

    t_tech = Table(tech_table_data, colWidths=[4.5 * cm, 5.5 * cm, 4.6 * cm, 4.0 * cm])
    t_tech.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.8, colors.HexColor('#000000')),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f8fafc')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(t_tech)
    story.append(Spacer(1, 0.4 * cm))

    footer_table = Table([
        [
            Paragraph(f"<b>{company_name}</b>", STYLE_FTR),
            Paragraph("___________________________<br/><br/><b>Consumer / Authorized Signature</b>", ParagraphStyle('sig_r', parent=STYLE_FTR, alignment=2))
        ]
    ], colWidths=[9.3 * cm, 9.3 * cm])
    footer_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(footer_table)

    pdf.build(story, onFirstPage=_draw_sldr_frame)
    return buf.getvalue()


def make_net_meter_canvas(company: dict):
    company_name = (company.get("company_name") or company.get("name") or "").strip()

    class NetMeterCanvas(canvas.Canvas):
        _startPage: Any
        _pageNumber: Any

        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            # Footer line
            self.setStrokeColor(colors.HexColor('#cbd5e1'))
            self.setLineWidth(0.5)
            self.line(1.5 * cm, 1.2 * cm, 21.0 * cm - 1.5 * cm, 1.2 * cm)
            
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor('#475569'))
            if company_name:
                self.drawString(1.5 * cm, 0.8 * cm, company_name)
            self.drawCentredString(10.5 * cm, 0.8 * cm, "Net Metering Connection Agreement")
            self.drawRightString(21.0 * cm - 1.5 * cm, 0.8 * cm, f"Page {self._pageNumber} of {page_count}")
            self.restoreState()

    return NetMeterCanvas


def generate_net_meter_agreement_pdf(client: dict, company: dict) -> bytes:
    # 1. Validation check
    missing_fields = []
    if not client.get("full_name"):
        missing_fields.append("Client Name")
    if not client.get("consumer_number"):
        missing_fields.append("Consumer Number")
    if not client.get("system_kw") and not client.get("capacity"):
        missing_fields.append("Solar System Capacity (kW)")
    
    if missing_fields:
        from fastapi import HTTPException
        raise HTTPException(
            status_code=400,
            detail=f"Missing required client data for Net Meter Agreement: {', '.join(missing_fields)}. Please update client details before generating."
        )

    buf = BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.6 * cm,
        rightMargin=1.6 * cm,
        topMargin=1.3 * cm,
        bottomMargin=1.3 * cm
    )
    story = []

    # Dynamic Data Extraction
    client_name = (client.get("full_name") or client.get("name") or "").strip()
    consumer_no = str(client.get("consumer_number") or "").strip()
    client_addr = (client.get("address") or "").strip()
    city = (client.get("city") or "").strip()
    pincode = str(client.get("pincode") or "").strip()
    full_address = f"{client_addr}{', ' + city if city else ''}{' - ' + pincode if pincode else ''}".strip(", -")
    
    system_kw = str(client.get("system_kw") or client.get("capacity") or "").strip()
    
    date_str = client.get("installation_date") or client.get("created_at") or datetime.now().strftime("%d/%m/%Y")
    if len(date_str) > 10:
        date_str = date_str[:10]
        
    company_name = company.get("company_name") or ""
    # BU Number and BU Text from onboarding (enriched via _enrich_client_doc before this call)
    raw_bu = str(client.get("bu_number") or client.get("bu_no") or client.get("bu") or "").strip()
    if raw_bu:
        bu_no = raw_bu if raw_bu.upper().startswith("BU-") else f"BU-{raw_bu}"
    else:
        bu_no = ""
    bu_text = client.get("bu_text") or ""
    sub_div = bu_text or client.get("sub_division") or ""
    division = client.get("division") or ""

    if sub_div and bu_no:
        sub_bu_str = bu_no if sub_div.strip().upper() in ("BU", "BU-") else f"{sub_div} {bu_no}"
    else:
        sub_bu_str = sub_div or bu_no
    licensee_sub = f", {sub_bu_str}" if sub_bu_str else ""


    # Define Styles (Refined font sizing & spacing for exact 5-page layout with compact vertical gaps)
    style_h1 = ParagraphStyle('NMA_H1', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=14, leading=17, alignment=1, spaceBefore=4, spaceAfter=3)
    style_h2 = ParagraphStyle('NMA_H2', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=12, leading=15, alignment=1, spaceAfter=8)
    style_clause_h = ParagraphStyle('NMA_ClauseH', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=9.5, leading=12.5, spaceBefore=4, spaceAfter=1.5)
    style_body = ParagraphStyle('NMA_Body', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, leading=11.0, alignment=4, spaceBefore=0, spaceAfter=2)
    style_body_bold = ParagraphStyle('NMA_BodyBold', parent=style_body, fontName='Helvetica-Bold')

    # ==================== PAGE 1 (AGREEMENT PREAMBLE) ====================
    story.append(Paragraph("<b>ANNEXURE – 3</b>", style_h1))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("<b>Net Metering Connection Agreement</b>", style_h2))
    story.append(Spacer(1, 0.4 * cm))

    discom_code = client.get("discom_code") or client.get("discom") or "MSEDCL"
    discom_name = client.get("discom_name") or "Maharashtra State Electricity Distribution Company Limited"
    licensee_title = client.get("distribution_licensee") or f"Additional Executive Engineer{licensee_sub}, {discom_code}"

    preamble_p1 = (
        f"This Agreement is made and entered into at (location) <b>{city}</b> on this "
        f"<b>(date {date_str})</b> between the Eligible Consumer <b>{client_name}</b> "
        f"having premises at <b>{full_address}</b> and Consumer No <b>{consumer_no}</b> "
        f"as the first Party<br/>"
        f"AND<br/>"
        f"The Distribution Licensee <b>{licensee_title}</b>, "
        f"(hereinafter referred to as 'the Licensee') and having its Registered Office at <b>{division}</b> as second Party of this Agreement;"
    )
    story.append(Paragraph(preamble_p1, style_body))
    story.append(Spacer(1, 0.3 * cm))

    preamble_p2 = (
        "Whereas, the Eligible Consumer has applied to the Licensee for approval of a Net Metering Arrangement "
        "under the provisions of the State Electricity Regulatory Commission (Net Metering for Roof-top Solar Photo Voltaic Systems) Regulations "
        "('the Net Metering Regulations') and sought its connectivity to the Licensee's Distribution Network;"
    )
    story.append(Paragraph(preamble_p2, style_body))
    story.append(Spacer(1, 0.3 * cm))

    preamble_p3 = (
        f"And whereas, the Licensee has agreed to provide Network connectivity to the Eligible Consumer for injection "
        f"of electricity generated from its Roof-top Solar PV System of <b>{system_kw} kilowatt</b>;"
    )
    story.append(Paragraph(preamble_p3, style_body))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("<b>Both Parties hereby agree as follows</b>", style_body_bold))

    # ==================== CONTINUOUS DYNAMIC FLOW (PAGES 2 - END) ====================
    story.append(Paragraph("<b>1. Eligibility:</b>", style_clause_h))
    story.append(Paragraph(
        "The Roof-top Solar PV System meets the applicable norms for being integrated into the Distribution Network, "
        "and that the Eligible Consumer shall maintain the System accordingly for the duration of this Agreement.",
        style_body
    ))

    story.append(Paragraph("<b>2. Technical and Inter-connection Requirements:</b>", style_clause_h))
    story.append(Paragraph(
        "2.1. The metering arrangement and the inter-connection of the Roof-top Solar PV System with the Network of the Licensee "
        "shall be as per the provisions of the Net Metering Regulations and the technical standards and norms specified by the "
        "Central Electricity Authority for connectivity of distributed generation resources and for the installation and operation of meters.",
        style_body
    ))
    story.append(Paragraph(
        "2.2. The Eligible Consumer agrees, that he shall install, prior to connection of the Roof-top Solar PV System to the Network of the Licensee, "
        "an isolation device (both automatic and in built within inverter and external manual relays); and the Licensee shall have access to it "
        "if required for the repair and maintenance of the Distribution Network.",
        style_body
    ))
    story.append(Paragraph(
        "2.3. The Licensee shall specify the interface/inter-connection point and metering point.",
        style_body
    ))
    story.append(Paragraph(
        "2.4. The Eligible Consumer shall specify relevant data, such as voltage, frequency, circuit breaker, isolator position in his System, "
        "as and when required by the Licensee.",
        style_body
    ))

    story.append(Paragraph("<b>3. Safety:</b>", style_clause_h))
    story.append(Paragraph(
        "3.1 The equipment connected to the Licensee's Distribution System shall be compliant with relevant International (IEEE/IEC) "
        "or Indian Standards (BIS), as the case may be, and the installation of electrical equipment shall comply with the requirements "
        "specified by the Electricity Authority regarding safety and electricity supply.",
        style_body
    ))
    story.append(Paragraph(
        "3.2 The design, installation, maintenance and operation of the Roof-top Solar PV System shall be undertaken in a manner "
        "conducive to the safety of the Roof-top Solar PV System as well as the Licensee's Network.",
        style_body
    ))
    story.append(Paragraph(
        "3.3 If, at any time, the Licensee determines that the Eligible Consumer's Roof-top Solar PV System is causing or may cause damage "
        "to and/or results in the Licensee's other consumers or its assets, the Eligible Consumer shall disconnect the Roof-top Solar PV System "
        "from the distribution Network upon direction from the Licensee, and shall undertake corrective measures at his own expense prior to re-connection.",
        style_body
    ))
    story.append(Paragraph(
        "3.4 The Licensee shall not be responsible for any accident resulting in injury to human beings or animals or damage to property "
        "that may occur due to back- feeding from the Roof-top Solar PV System when the grid supply is off. The Licensee may disconnect "
        "the installation at any time in the event of such exigencies to prevent such accident.",
        style_body
    ))

    story.append(Paragraph("<b>Other Clearances and Approvals:</b>", style_clause_h))
    story.append(Paragraph(
        "The Eligible Consumer shall obtain any statutory approvals and clearances that may be required, such as from the Electrical Inspector "
        "or the municipal or other authorities, before connecting the Roof-top Solar PV System to the distribution Network.",
        style_body
    ))

    story.append(Paragraph("<b>4. Period of Agreement, and Termination:</b>", style_clause_h))
    story.append(Paragraph("This Agreement shall be for a period of 20 years, but may be terminated prematurely", style_body))
    story.append(Paragraph("(a) By mutual consent; or", style_body))
    story.append(Paragraph("(b) By the Eligible Consumer, by giving 30 days' notice to the Licensee ;", style_body))
    story.append(Paragraph(
        "(c) By the Licensee, by giving 30 days' notice, if the Eligible Consumer breaches any terms of this Agreement or the provisions "
        "of the Net Metering Regulations and does not remedy such breach within 30 days, or such other reasonable period as may be provided, "
        "of receiving notice of such breach, or for any other valid reason communicated by the Licensee in writing.",
        style_body
    ))

    story.append(Paragraph("<b>6. Access and Disconnection:</b>", style_clause_h))
    story.append(Paragraph(
        "6.1. The Eligible Consumer shall provide access to the Licensee to the metering equipment and disconnecting devices "
        "of Roof-top Solar PV System, both automatic and manual, by the Eligible Consumer.",
        style_body
    ))
    story.append(Paragraph(
        "6.2. If, in an emergent or outage situation, the Licensee cannot access the disconnecting devices of the Roof-top Solar PV System, "
        "both automatic and manual, it may disconnect power supply to the premises.",
        style_body
    ))
    story.append(Paragraph(
        "6.3 Upon termination of this Agreement under Clause 5, the Eligible Consumer shall disconnect the Roof-top Solar PV System "
        "forthwith from the Network of the Licensee.",
        style_body
    ))

    story.append(Paragraph("<b>7. Liabilities:</b>", style_clause_h))
    story.append(Paragraph(
        "7.1. The Parties shall indemnify each other for damages or adverse effects of either Party's negligence or misconduct "
        "during the installation of the Roof-top Solar PV System, connectivity with the distribution Network and operation of the System.",
        style_body
    ))
    story.append(Paragraph(
        "7.2. The Parties shall not be liable to each other for any loss of profits or revenues, business interruption losses, "
        "loss of contract or goodwill, or for indirect, consequential, incidental or special damages including, but not limited to, "
        "punitive or exemplary damages, whether any of these liabilities, losses or damages arise in contract, or otherwise.",
        style_body
    ))

    story.append(Paragraph("<b>8. Commercial Settlement:</b>", style_clause_h))
    story.append(Paragraph(
        "8.1. The commercial settlements under this Agreement shall be in accordance with the Net Metering Regulations.",
        style_body
    ))
    story.append(Paragraph(
        "8.2. The Licensee shall not be liable to compensate the Eligible Consumer if his Rooftop Solar PV System is unable to inject surplus power "
        "generated into the Licensee's Network on account of failure of power supply in the grid/Network.",
        style_body
    ))
    story.append(Paragraph(
        "8.3. The existing metering System, if not in accordance with the Net Metering Regulations, shall be replaced by a bi-directional meter "
        "(whole current/CT operated) or a pair of meters (as per the definition of 'Net Meter' in the Regulations), and a separate generation meter "
        "may be provided to measure Solar power generation. The bi-directional meter (whole current/CT operated) or pair of meters shall be installed "
        "at the inter-connection point to the Licensee's Network for recording export and import of energy.",
        style_body
    ))
    story.append(Paragraph(
        "8.4. The uni-directional and bi-directional or pair of meters shall be fixed in separate meter boxes in the same proximity.",
        style_body
    ))
    story.append(Paragraph(
        "8.5. The Licensee shall issue monthly electricity bill for the net metered energy on the scheduled date of meter reading. "
        "If the exported energy exceeds the imported energy, the Licensee shall show the net energy exported as credited Units of electricity "
        "as specified in the Net Metering Regulations, 2015. If the exported energy is less than the imported energy, the Eligible Consumer "
        "shall pay the Distribution Licensee for the net energy imported at the prevailing tariff approved by the Commission for the consumer "
        "category to which he belongs.",
        style_body
    ))

    story.append(Paragraph("<b>9. Connection Costs:</b>", style_clause_h))
    story.append(Paragraph(
        "The Eligible Consumer shall bear all costs related to the setting up of the Roof-top Solar PV System, excluding the Net Metering Arrangement costs.",
        style_body
    ))

    story.append(Paragraph("<b>10. Dispute Resolution:</b>", style_clause_h))
    story.append(Paragraph(
        "10.1 Any dispute arising under this Agreement shall be resolved promptly, in good faith and in an equitable manner by both the Parties.",
        style_body
    ))
    story.append(Paragraph(
        "10.2 The Eligible Consumer shall have recourse to the concerned Consumer Grievance Redressal Forum constituted under the relevant Regulations "
        "in respect of any grievance regarding billing which has not been redressed by the Licensee.",
        style_body
    ))

    witness_intro = Paragraph(
        f"In the witness where of <b>{client_name}</b> for and on behalf of Eligible Consumer and Shri. "
        f"<b>{licensee_title}</b>, for and on behalf of {discom_code} agree to this agreement.",
        style_body
    )

    # Signature Table (Aligned cleanly with consumer left, DISCOM right, Witnesses aligned)
    sig_table_data = [
        [
            Paragraph(f"<b>Signature of Eligible Consumer</b><br/><br/><br/>___________________________<br/><b>{client_name}</b><br/>Eligible Consumer", style_body),
            Paragraph(f"<b>Signature of Licensee</b><br/><br/><br/>Shri. ___________________________<br/>{licensee_title}<br/>for and on behalf of {discom_code}", style_body)
        ],
        [
            Paragraph("<br/><b>Witness 1:</b> ___________________________", style_body),
            Paragraph("<br/><b>Witness 1:</b> ___________________________", style_body)
        ],
        [
            Paragraph("<b>Witness 2:</b> ___________________________", style_body),
            Paragraph("<b>Witness 2:</b> ___________________________", style_body)
        ],
        [
            Paragraph(f"<br/><br/><b>{company_name}</b><br/>Proprietor / Authorized Manager", style_body),
            Paragraph("<br/><br/><b>Official Stamp / Seal</b>", style_body)
        ]
    ]
    t_sig = Table(sig_table_data, colWidths=[9.0 * cm, 9.0 * cm])
    t_sig.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))

    sig_block = KeepTogether([
        Spacer(1, 0.4 * cm),
        witness_intro,
        Spacer(1, 0.6 * cm),
        t_sig
    ])
    story.append(sig_block)

    pdf.build(story, canvasmaker=make_net_meter_canvas(company))
    return buf.getvalue()


def make_vendor_canvas(company: dict):
    company_name = (company.get("company_name") or company.get("name") or "").strip()

    class VendorCanvas(canvas.Canvas):
        _startPage: Any
        _pageNumber: Any

        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            # Footer line
            self.setStrokeColor(colors.HexColor('#cbd5e1'))
            self.setLineWidth(0.5)
            self.line(1.5 * cm, 1.2 * cm, 21.0 * cm - 1.5 * cm, 1.2 * cm)
            
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor('#475569'))
            if company_name:
                self.drawString(1.5 * cm, 0.8 * cm, company_name)
            self.drawCentredString(10.5 * cm, 0.8 * cm, "Rooftop Solar Vendor Agreement")
            self.drawRightString(21.0 * cm - 1.5 * cm, 0.8 * cm, f"Page {self._pageNumber} of {page_count}")
            self.restoreState()

    return VendorCanvas


def generate_vendor_agreement_pdf(client: dict, company: dict) -> bytes:
    buf = BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.1 * cm,
        bottomMargin=1.1 * cm
    )
    story = []

    # Dynamic Data Extraction
    client_name = (client.get("full_name") or client.get("name") or "").strip()
    consumer_no = str(client.get("consumer_number") or "").strip()
    client_addr = (client.get("address") or "").strip()
    city = (client.get("city") or "").strip()
    pincode = str(client.get("pincode") or "").strip()
    full_address = f"{client_addr}{', ' + city if city else ''}{' - ' + pincode if pincode else ''}".strip(", -")
    
    system_kw = str(client.get("system_kw") or client.get("capacity") or "").strip()
    panel_make = (client.get("panel_brand") or client.get("panel_make") or "").strip()
    panel_wattage = str(client.get("panel_wattage") or "").strip()
    inverter_make = (client.get("inverter_make") or "").strip()
    inverter_kw = str(client.get("inverter_capacity") or client.get("system_kw") or "").strip()
    total_cost = str(client.get("total_cost") or client.get("quotation_amount") or "").strip()
    
    date_obj = datetime.now()
    day_str = date_obj.strftime("%d")
    month_str = date_obj.strftime("%m")
    year_str = date_obj.strftime("%Y")

    company_name = company.get("company_name") or ""
    company_address = company.get("address") or ""
    company_pincode = company.get("pincode") or ""

    # Define Styles (Optimized spacing to ensure 3-page layout with signature on Page 3)
    style_h1 = ParagraphStyle('VA_H1', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=14, leading=17, alignment=1, spaceBefore=3, spaceAfter=2)
    style_sub = ParagraphStyle('VA_Sub', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=9, leading=12, alignment=1, spaceAfter=4)
    style_center_b = ParagraphStyle('VA_CenterB', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=10, leading=13, alignment=1, spaceBefore=3, spaceAfter=3)
    style_clause_h = ParagraphStyle('VA_ClauseH', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=9.5, leading=12.0, spaceBefore=4, spaceAfter=1.5)
    style_body = ParagraphStyle('VA_Body', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, leading=11.0, alignment=4, spaceBefore=0, spaceAfter=1.5)
    style_body_bold = ParagraphStyle('VA_BodyBold', parent=style_body, fontName='Helvetica-Bold')

    # ==================== PREAMBLE & TITLES ====================
    story.append(Paragraph("<u><b>Agreement Between</b></u>", style_h1))
    story.append(Paragraph("<b>Applicant and the registered/empaneled Vendor for installation of rooftop solar system in residential house of the Applicant under simplified procedure of Rooftop Solar Program Ph-II</b>", style_sub))
    story.append(Spacer(1, 0.2 * cm))

    exec_p = (
        f"This agreement is executed on (Day) <b>{day_str}</b> , (Month) <b>{month_str}</b> , (Year) <b>{year_str}</b> "
        f"for design, installation, commissioning and five years comprehensive maintenance of rooftop solar system "
        f"to be installed undersimplified procedure of Rooftop Solar Program Ph-II."
    )
    story.append(Paragraph(exec_p, style_body))
    story.append(Spacer(1, 0.2 * cm))

    discom_code = client.get("discom_code") or client.get("discom") or "MSEDCL"

    story.append(Paragraph("<b>Between</b>", style_center_b))
    applicant_p = (
        f"<b>{client_name}</b> has residential electricity connection with consumer number <b>{consumer_no}</b> "
        f"from {discom_code} (DISCOM) at <b>{full_address} PIN code : {pincode}</b> (Hereinafter referred to as Applicant)."
    )
    story.append(Paragraph(applicant_p, style_body))
    story.append(Spacer(1, 0.2 * cm))

    story.append(Paragraph("<b>And</b>", style_center_b))
    vendor_p = (
        f"<b>{company_name}</b> (Name of Vendor) is registered/ empaneled with the {discom_code} (hereinafter referred as DISCOM) "
        f"and is having registered/functional office at <b>{company_address} . PIN CODE- {company_pincode}</b> . "
        f"Both Applicant and the Vendor are jointly referred as Parties."
    )
    story.append(Paragraph(vendor_p, style_body))
    story.append(Spacer(1, 0.2 * cm))

    story.append(Paragraph("<b>Whereas</b>", style_body_bold))
    story.append(Paragraph("- The Applicant intends to install rooftop solar system under simplified procedure of Rooftop Solar Programmed Ph-II of the MNRE.", style_body))
    story.append(Paragraph("- The Vendor is registered/empaneled vendor with DISCOM for installation of rooftop solar under MNRE Schemes. The Vendor satisfies all the existing regulation pertaining to electrical safety and license in the respective state and it is not debarred or blacklisted from undertaking any such installations by any state/central Government agency.", style_body))
    story.append(Paragraph("- Both the parties are mutually agreed and understand their roles and responsibilities and have no liability to any other agency/firm/stakeholder especially to DISCOM and MNRE.", style_body))
    story.append(Spacer(1, 0.3 * cm))

    # ==================== SECTIONS 1 - 15 ====================
    story.append(Paragraph("<b>1. GENERAL TERMS:</b>", style_clause_h))
    story.append(Paragraph("1.1. The Applicant hereby represents and warrants that the Applicant has the sole legal capacity to enter into this Agreement and authorize the construction, installation and commissioning of the Rooftop Solar System (“RTS System”) which is inclusive of Balance of System (“BoS”) on the Applicant's premises (“Applicant Site”). The Vendor reserves its right to verify ownership of the Applicant Site and Applicant covenants to co-operate and provide all information and documentation required by the Vendor for the same.", style_body))
    story.append(Paragraph("1.2. Vendor may propose changes to the scope, nature and or schedule of the services being performed under this Agreement. All proposed changes must be mutually agreed between the Parties. If Parties fail to agree on the variation proposed, either Party may terminate this Agreement by serving notice as per Clause 13.", style_body))
    story.append(Paragraph("1.3. The Applicant understands and agrees that future changes in load, electricity usage patterns and/or electrical grid issues may affect the performance of the RTS System and these factors have not been and cannot be considered in any analysis or quotation provided by Vendor or its Authorized Persons (defined below).", style_body))

    story.append(Paragraph("<b>2. RTS System:</b>", style_clause_h))
    story.append(Paragraph(f"2.1. Total capacity of RTS System will be minimum <b>{system_kw} KWatt</b>.", style_body))
    story.append(Paragraph("2.2. The Solar modules, inverters and BoS will confirm to minimum specifications and DCR requirement of MNRE.", style_body))
    story.append(Paragraph(f"2.3. Solar modules of <b>{panel_make}</b> make model, <b>{panel_wattage} Wp</b> capacity each and <b>21.13%</b> efficiency will be procured and installed by the Vendor", style_body))
    inv_mid = _get_vendor_inverter_clause_text(client)
    inv_str = f" <b>{inv_mid}</b>" if inv_mid else ""
    inv_clause_p = f"2.4. Solar inverter of{inv_str} rated output capacity will be procured and installed by the Vendor"
    story.append(Paragraph(inv_clause_p, style_body))
    story.append(Paragraph("2.5. The module mounting structure must withstand minimum wind load pressure as specified by MNRE.", style_body))
    story.append(Paragraph("2.6. Other BoS installations shall be as per best industry practice with all safety and protection gears installed by the vendor.", style_body))

    story.append(Paragraph("<b>3. PRICE AND PAYMENT TERMS:</b>", style_clause_h))
    story.append(Paragraph(f"3.1. The cost of an RTS System will be <b>Rs. {total_cost}/-</b> (to be decided mutually). The Applicant shall pay the total cost to the Vendor as under:", style_body))
    story.append(Paragraph("(i) 50 % as an advance on confirmation of the order.", style_body))
    story.append(Paragraph("(ii) 40 % against Proforma Invoice (PI) before dispatch of solar panels, inverters and other BoS items to be delivered.", style_body))
    story.append(Paragraph("(iii) 10 % after installation and commissioning of the RTS System. The order value and payment terms are fixed and will not be subject to any adjustment except as approved in writing by Vendor. The payment shall be made only through bankers' cheque / NEFT / RTGS / online payment portal as intimated by Vendor. No cash payments shall be accepted by Vendor or its Authorized Person.", style_body))

    story.append(Paragraph("<b>4. REPRESENTATIONS MADE BY THE APPLICANT:</b>", style_clause_h))
    story.append(Paragraph("The Applicant acknowledges and agrees that:", style_body))
    story.append(Paragraph("4.1. any timeline or schedule shared by Vendor for the provision of services and delivery of the RTS System is only an estimate and Vendor will not be liable for any delay that is not attributable to Vendor.", style_body))
    story.append(Paragraph("4.2. all information disclosed by the Applicant to Vendor in connection with the supply of the RTS System (or any part thereof), services and generation estimation (including, without limitation, the load profile and power bill) are true and accurate and acknowledges that Vendor has relied on the information produced by the Applicant to customize the RTS System layout and BoS design for the purposes of this Agreement.", style_body))
    story.append(Paragraph("4.3. all descriptive specifications, illustrations, drawings, data, dimensions, quotation, fact sheets, price lists and any advertising material circulated/published/provided by Vendor are approximate only.", style_body))
    story.append(Paragraph("4.4. any drawings, pre-feasibility report, specifications and plans composed by Vendor shall require the Applicant's approval within 5 (five) days of its receipt by electronic mail to Vendor and if the Applicant does not respond within this period, the drawings, specifications or plans shall be final and deemed to have been approved by the Applicant.", style_body))
    story.append(Paragraph("4.5. the Applicant shall not use the RTS System or any part thereof, other than in accordance with the product manufacturer's specifications, and covenants that any risk arising from misuse or/and inappropriate use shall be to the account of the Applicant alone.", style_body))
    story.append(Paragraph("4.6. The Applicant represents, warrants and covenants that:", style_body))
    story.append(Paragraph("(i) All electrical and plumbing infrastructure at the Applicant Site are in conformity with applicable laws.", style_body))
    story.append(Paragraph("(ii) the Applicant has the legal capacity to permit unfettered access to Vendor and its Authorized Persons for the purposes of execution and performance of this Agreement.", style_body))
    story.append(Paragraph("(iii) the Applicant has and will provide requisite power, water and other requisite resources and storage facilities for construction, installation, operation and maintenance of the RTS System.", style_body))
    story.append(Paragraph("(iv) The Applicant will provide support for site fabrication of structure, assembly and fitting of module mounting structure at Applicant Site.", style_body))
    story.append(Paragraph("(v) The Applicant will ensure that the Applicant Site is shadow free and free of all encumbrances during the lifetime of the RTS System.", style_body))
    story.append(Paragraph("(vi) Applicant should ensure that the Applicant regularly cleans and ensures accessibility and safety to the RTS System, as required by Vendor and dusting frequency in the premises.", style_body))
    story.append(Paragraph("(vii) The vendor is entitled to permit geo-tagging of the Applicant Site as a Vendor installation site.", style_body))
    story.append(Paragraph("(viii) Unless otherwise intimated by the Applicant in writing, Vendor is entitled to take photographs, videos and testimonials of the Applicant and the Applicant Site, and to create content which will become the property of Vendor and the same can be freely used by Vendor as part of its promotional and marketing activities across all platforms as it deems fit;", style_body))
    story.append(Paragraph("(ix) The Applicant validates the stability of the Applicant Site for the installation of the RTS System.", style_body))

    story.append(Paragraph("<b>5. MAINTENANCE:</b>", style_clause_h))
    story.append(Paragraph("5.1. Vendor shall provide five-year free workmanship maintenance. Vendor shall visit the Applicant's premises at least once every quarter after commissioning of the RTS System for maintenance purposes.", style_body))
    story.append(Paragraph("5.2. During such maintenance visit, Vendor shall check all nuts and bolts, fuses, earth resistance and other consumables in respect of the RTS System to ensure that it is in good working condition.", style_body))
    story.append(Paragraph("5.3. Cleaning requirement/expectation from the Applicant side – Applicant responsibility, minimum expectation from Applicant that it will be cleaned regularly as per the dusting frequency.", style_body))

    story.append(Paragraph("<b>6. ACCESS AND RIGHT OF ENTRY:</b>", style_clause_h))
    story.append(Paragraph("6.1. The Applicant hereby grants permission to Vendor and its authorized personnel, representatives, associates, officers, employees, financing agents, subcontractors (“Authorized Persons”) to enter the Applicant Site for the purposes of:", style_body))
    story.append(Paragraph("(a) conducting feasibility study.", style_body))
    story.append(Paragraph("(b) storing the RTS System/any part thereof.", style_body))
    story.append(Paragraph("(c) installing the RTS System.", style_body))
    story.append(Paragraph("(d) inspecting the RTS System.", style_body))
    story.append(Paragraph("(e) conducting repairs and maintenance to the RTS System.", style_body))
    story.append(Paragraph("(f) removing the RTS System (or any part thereof), if necessary for any reason whatsoever.", style_body))
    story.append(Paragraph("(g) Such other matters as necessary to execute and perform its rights and obligations under this Agreement.", style_body))
    story.append(Paragraph("6.2. The Applicant shall ensure that third-party consents necessary for the Authorized Persons to access the Applicant Site are obtained prior to commencement of services under this Agreement.", style_body))

    story.append(Paragraph("<b>7. WARRANTIES:</b>", style_clause_h))
    story.append(Paragraph("7.1. Product Warranty: The Applicant shall be entitled to manufacturers' warranty. Any warranty in relation to RTS System supplied to the Applicant by Vendor under this Agreement is limited to the warranty given by the manufacturer of the RTS System (or any part thereof) to Vendor.", style_body))
    story.append(Paragraph("7.2. Installation Warranty: Vendor warrants that all installations shall be free from workmanship defects or BOS defects for a period of five years from the date of installation of the RTS System. The warranty is limited to Vendor rectifying the workmanship or BOS defects at Vendor's expense in respect of those defects reported by the Applicant, in writing. The Applicant is obliged and liable to report such defects within 15 (fifteen) days of occurrence of such defect.", style_body))
    story.append(Paragraph("7.3. Subject to manufacturer warranty, Vendor warrants that the solar modules supplied herein shall have tolerance within a five-percentage range (+/-5%). The peak-power point voltage and the peak-power point current of any supplied solar module and/or any module string (series connected modules) shall not vary by more than 5% (five percent) from the respective arithmetic means for all modules and/or for all module strings, as the case may be, provided The RTS System is properly maintained, and the Applicant Site is free from shadow at the time of operation of the RTS System.", style_body))
    story.append(Paragraph("7.4. Exceptions for warranty:", style_body))
    story.append(Paragraph("(a) Any attempt by any person other than Vendor or its Authorized Persons to adjust, modify, repair or provide maintenance to the RTS System, shall disentitle the Applicant of the warranty provided by Vendor hereunder.", style_body))
    story.append(Paragraph("(b) Vendor shall not be liable for any degeneration or damage to the RTS System due to any action or inaction on the part of the Applicant.", style_body))
    story.append(Paragraph("(c) Vendor shall not be bound or liable to remedy any damage, fault, failure or malfunction of the RTS System owing to external causes, including but not limited to accidents, misuse, neglect, if usage and/or storage and/or installation are non-confirming to product instructions, modifications by the Applicant leading to shading or accessibility issues, failure to perform required maintenance, normal wear and tear, Force Majeure Event, or negligence or default attributable to the Applicant.", style_body))
    story.append(Paragraph("(d) Vendor shall not be liable to repair or remedy any accessories or parts added to the RTS System that were not originally sourced by Vendor to the Applicant.", style_body))

    story.append(Paragraph("<b>8. PERFORMANCE GUARANTEE:</b>", style_clause_h))
    story.append(Paragraph("8.1. Vendor guarantees minimum system performance ratio of 75% as per performance ratio test carried out in adherence to IEC 61724 or equivalent BIS for a period of five years.", style_body))

    story.append(Paragraph("<b>9. INSURANCE:</b>", style_clause_h))
    story.append(Paragraph("9.1. Vendor may, at its sole discretion, obtain insurance covering risks of loss/damage to the RTS System (any part thereof) during transit from Vendor's warehouse until delivery to the Applicant Site and until installation and commissioning.", style_body))
    story.append(Paragraph("9.2. Thereafter, all risk shall pass on to the Applicant and the Applicant may accordingly procure relevant insurances.", style_body))

    story.append(Paragraph("<b>10. CANCELLATION:</b>", style_clause_h))
    story.append(Paragraph("10.1. The Applicant may cancel the order placed on Vendor within 7 (seven) days from the date of remittance of advance money or the date of order acceptance, whichever is earlier (“Order Confirmation”) by serving notice as per Clause 13.", style_body))
    story.append(Paragraph("10.2. If the Applicant cancels the order after the expiry of 7 (seven) days from the date of Order Form, the Applicant shall be liable to pay Vendor, a cancellation fee of 30 % of the total order value plus costs and expenses incurred by Vendor, including, costs for labour, design, return of products, administrative costs, subvention costs.", style_body))
    story.append(Paragraph("10.3. Notwithstanding the aforesaid, the Applicant shall not be entitled to cancel the Order Form after Vendor has dispatched the RTS System (or any part thereof, including BOS) to the Applicant Site. If Applicant chooses to terminate the Order Form after dispatch, the entire amount paid by the Applicant till date, shall be forfeited by Vendor.", style_body))

    story.append(Paragraph("<b>11. LIMITATION OF LIABILITY AND INDEMNITY:</b>", style_clause_h))
    story.append(Paragraph("11.1. To the extent that terms implied by law apply to the RTS System and the services rendered under this Agreement, Vendor's liability for any breach of those terms is limited to:", style_body))
    story.append(Paragraph("(a) repairing or replacing the RTS System/any part thereof, as applicable; or", style_body))
    story.append(Paragraph("(b) Refund of the moneys paid by the Applicant to Vendor, if Vendor cannot fulfil the order.", style_body))

    story.append(Paragraph("<b>12. SUSPENSION AND TERMINATION:</b>", style_clause_h))
    story.append(Paragraph("12.1. If the Applicant fails to pay any sum due under this Agreement on the due date, Vendor may, in addition to its other rights under this Agreement, suspend its obligations under this Agreement until all outstanding amounts (including interest due) are paid.", style_body))

    story.append(Paragraph("<b>13. NOTICES:</b>", style_clause_h))
    story.append(Paragraph("Any notice or other communication under this Agreement to Vendor and or to the Applicant, shall be in writing, in English language and shall be delivered or sent: (a) by electronic mail and/or (b) by hand delivery or registered post/courier, at the registered address of Applicant/Vendor.", style_body))

    story.append(Paragraph("<b>14. FORCE MAJEURE EVENT:</b>", style_clause_h))
    story.append(Paragraph("14.1. Neither Party shall be in default due to any delay or failure to perform its/his/her/their obligations under this Agreement which arises from or is a consequence of occurrence of an event which is beyond the reasonable control of such Party, and which makes performance of its/his/her/their obligations under this Agreement impossible or so impractical as reasonably to be considered impossible in the circumstances, and includes, but is not limited to, war, riot, civil disorder, earthquake, fire, explosion, storm, flood or other adverse weather conditions, pandemic, epidemic, embargo, strikes, lockouts, labour difficulties, other industrial action, acts of government, unavailability of equipment from vendor, changes requested by the Applicant (“Force Majeure Event”).", style_body))

    story.append(Paragraph("<b>15. GOVERNING LAW AND DISPUTE RESOLUTION:</b>", style_clause_h))
    story.append(Paragraph("15.1. The interpretation and enforcement of this Agreement shall be governed by the laws of India.", style_body))
    story.append(Paragraph("15.2. In the event of any dispute, controversy or difference between the Parties arising out of, or relating to this Agreement (“Dispute”), both Parties shall make an effort to resolve the Dispute in good faith, failing which, any Party to the Dispute shall be entitled to refer the Dispute to arbitration to resolve the Dispute in the manner set out in this Clause. The rights and obligations of the Parties under this Agreement shall remain in full force and effect pending the award in such arbitration proceeding.", style_body))
    story.append(Paragraph("15.3. The arbitration proceeding shall be governed by the provisions of the Arbitration and Conciliation Act, 1996 and shall be settled by a sole arbitrator mutually appointed by the Parties.", style_body))

    # Signature Block
    sig_table_data = [
        [
            Paragraph(f"<br/><br/>___________________________<br/>(Applicant)<br/><b>{client_name}</b>", style_body),
            Paragraph(f"<br/><br/>___________________________<br/>(Vendor)<br/><b>{company_name}</b>", style_body)
        ],
        [
            Paragraph("", style_body),
            Paragraph("<br/><br/><b>Official Stamp / Seal</b>", style_body)
        ]
    ]
    t_sig = Table(sig_table_data, colWidths=[9.0 * cm, 9.0 * cm])
    t_sig.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))

    sig_block = KeepTogether([
        Spacer(1, 0.5 * cm),
        t_sig
    ])
    story.append(sig_block)

    pdf.build(story, canvasmaker=make_vendor_canvas(company))
    return buf.getvalue()


def make_meter_testing_canvas(company: dict):
    addr = (company.get("address") or company.get("address_line_1") or company.get("office_address") or "").strip()
    city = (company.get("city") or "").strip()
    state = (company.get("state") or "").strip()
    pincode = (company.get("pincode") or "").strip()
    full_addr_parts = [addr]
    if city:
        full_addr_parts.append(city)
    if state:
        full_addr_parts.append(state)
    if pincode:
        full_addr_parts.append(f"- {pincode}")
    full_addr = ", ".join(p for p in full_addr_parts if p).replace(", -", " -")
    phone = (company.get("mobile") or company.get("phone") or company.get("phone_number") or "").strip()

    class MeterTestingCanvas(canvas.Canvas):
        _startPage: Any
        _pageNumber: Any

        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            self.setStrokeColor(colors.HexColor('#991b1b'))
            self.setLineWidth(1.5)
            self.line(1.2 * cm, 1.4 * cm, 21.0 * cm - 1.2 * cm, 1.4 * cm)

            self.setFont("Helvetica-Bold", 7.5)
            self.setFillColor(colors.HexColor('#2563eb'))

            line1 = f"OFFICE :- {full_addr}" if full_addr else ""
            line2 = f"PHONE : {phone}" if phone else ""

            self.drawString(1.2 * cm, 1.0 * cm, line1)
            self.drawString(1.2 * cm, 0.65 * cm, line2)

            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor('#475569'))
            self.drawRightString(21.0 * cm - 1.2 * cm, 0.65 * cm, f"Page {self._pageNumber} of {page_count}")
            self.restoreState()

    return MeterTestingCanvas


def generate_meter_testing_request_pdf(client: dict, company: dict) -> bytes:
    buf = BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.8 * cm
    )
    story = []

    company_name = (company.get('company_name') or company.get('name') or company.get('legal_business_name') or '').strip()
    gst_no = (company.get('gst_number') or company.get('gst') or '').strip()

    stages_dict = dict(client.get("stages") or {})
    ob_dict = dict(stages_dict.get("onboarding_data") or {})

    client_name = (client.get('full_name') or client.get('name') or client.get('client_name') or ob_dict.get('full_name') or '').strip()
    consumer_num = str(client.get('consumer_number') or client.get('consumer_no') or ob_dict.get('consumer_number') or '').strip()

    client_addr = (client.get('address') or ob_dict.get('address') or '').strip()
    city = (client.get('city') or ob_dict.get('city') or '').strip()
    pincode = str(client.get('pincode') or ob_dict.get('pincode') or '').strip()

    location_parts = []
    if client_addr:
        location_parts.append(client_addr)
    if city and pincode:
        location_parts.append(f"{city} - {pincode}")
    elif city:
        location_parts.append(city)
    elif pincode:
        location_parts.append(pincode)
    location_str = ", ".join(location_parts)

    gen_make = (client.get('gen_meter_make') or client.get('generation_meter_make') or ob_dict.get('gen_meter_make') or ob_dict.get('generation_meter_make') or '').strip()
    gen_serial = (client.get('gen_meter_serial') or client.get('generation_meter_serial') or client.get('gen_meter_sn') or ob_dict.get('gen_meter_serial') or '').strip()
    net_make = (client.get('net_meter_make') or ob_dict.get('net_meter_make') or '').strip()
    net_serial = (client.get('net_meter_serial') or client.get('net_meter_sn') or ob_dict.get('net_meter_serial') or '').strip()

    def _clean_field(val: str) -> str:
        if val.upper() in ("NA", "N/A", "0", "DEFAULT", "GROWATT", "NULL", "NONE"):
            return ""
        return val

    gen_make = _clean_field(gen_make)
    gen_serial = _clean_field(gen_serial)
    net_make = _clean_field(net_make)
    net_serial = _clean_field(net_serial)

    logo_bytes = company.get("logo_bytes")
    logo_d = None
    if logo_bytes:
        try:
            from PIL import Image as PILImage
            img = PILImage.open(BytesIO(logo_bytes))
            img_w, img_h = img.size
            if img_w > 0 and img_h > 0:
                aspect = img_h / float(img_w)
                max_w = 6.2 * cm
                max_h = 2.4 * cm
                target_w = max_w
                target_h = target_w * aspect
                if target_h > max_h:
                    target_h = max_h
                    target_w = target_h / aspect
                resample_filter = getattr(getattr(PILImage, "Resampling", PILImage), "LANCZOS", getattr(PILImage, "LANCZOS", 1))
                resampled = img.resize((round(target_w * 4), round(target_h * 4)), resample_filter)
                res_buf = BytesIO()
                resampled.save(res_buf, format='PNG')
                logo_d = RLImage(BytesIO(res_buf.getvalue()), width=target_w, height=target_h)
        except Exception:
            logo_d = None
    if not logo_d:
        logo_d = Spacer(6.2 * cm, 1.2 * cm)

    p_title = Paragraph(f"<b><font size='18' color='#1d4ed8'>{company_name.upper()}</font></b>", ParagraphStyle('mtr_hdr_title', parent=styles['Normal'], fontName='Helvetica-Bold', leading=20, spaceBefore=5))
    gst_text = f"GST NO – {gst_no}" if gst_no else ""
    p_gst = Paragraph(f"<b><font size='9' color='#1d4ed8'>{gst_text}</font></b>", ParagraphStyle('mtr_hdr_gst', parent=styles['Normal'], fontName='Helvetica-Bold', alignment=2, leading=14))

    t_hdr = Table([[logo_d, p_title, p_gst]], colWidths=[6.2 * cm, 7.2 * cm, 5.2 * cm])
    t_hdr.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))

    t_div = Table([[""]], colWidths=[18.6 * cm])
    t_div.setStyle(TableStyle([
        ('LINEABOVE', (0, 0), (-1, -1), 1.5, colors.HexColor('#1d4ed8')),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))

    story.extend([t_hdr, Spacer(1, 0.1 * cm), t_div, Spacer(1, 0.4 * cm)])

    STYLE_BOLD_SERIF = ParagraphStyle('mtr_bold', parent=styles['Normal'], fontSize=10, fontName='Helvetica-Bold', textColor=colors.HexColor('#000000'), leading=14)
    STYLE_BODY = ParagraphStyle('mtr_body', parent=styles['Normal'], fontSize=10, fontName='Helvetica', textColor=colors.HexColor('#000000'), leading=14.5, alignment=4, spaceAfter=8)

    discom_code = client.get("discom_code") or client.get("discom") or "MSEDCL"
    licensee_title = client.get("distribution_licensee") or "Additional Executive Engineer"
    city_line = f"{discom_code} Meter Lab {city}." if city else f"{discom_code} Meter Lab."
    to_text_parts = ["<b>To,</b>", f"<b>{licensee_title}</b>", f"<b>{city_line}</b>"]
    if pincode:
        to_text_parts.append(f"<b>{pincode}</b>")
    to_text = "<br/>".join(to_text_parts)
    story.append(Paragraph(to_text, STYLE_BOLD_SERIF))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("<b>Sub: Request for Gen-meter Letter.</b>", STYLE_BOLD_SERIF))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("Dear Sir,", STYLE_BODY))
    story.append(Spacer(1, 0.1 * cm))

    p1 = "I hope this letter finds you well. I am writing to request meter testing services for my solar photovoltaic (PV) system. As a responsible solar PV system owner, I understand the importance of accurate and reliable meter readings to ensure the system's optimal performance and compliance with regulatory standards."
    story.append(Paragraph(p1, STYLE_BODY))

    p2 = f"<b>Customer Name:</b> <u>{client_name}</u> <b>C NO.</b> <u>{consumer_num}</u> I currently have a solar PV system installed at the following <b>location:</b> {location_str}"
    story.append(Paragraph(p2, STYLE_BODY))

    p3 = "To ensure the system's efficiency and adherence to industry standards, I am seeking a comprehensive meter testing service for the following meters within the system:"
    story.append(Paragraph(p3, STYLE_BODY))

    meter_rows = [
        [
            Paragraph(f"<b>Generation Meter - Make-</b> {gen_make}", STYLE_BODY),
            Paragraph(f"<b>SERIAL NO-</b> {gen_serial}", STYLE_BODY),
        ],
        [
            Paragraph(f"<b>NET METER – MAKE -</b> {net_make}", STYLE_BODY),
            Paragraph(f"<b>SERIAL NO -</b> {net_serial}", STYLE_BODY),
        ]
    ]
    t_meters = Table(meter_rows, colWidths=[9.3 * cm, 9.3 * cm])
    t_meters.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))
    story.append(t_meters)
    story.append(Spacer(1, 0.3 * cm))

    p4 = "I kindly request that the meter testing service be conducted by a certified and accredited organization, ensuring accurate and unbiased results. The testing should include a thorough assessment of the meters' functionality, calibration, and accuracy, as well as verification of their compliance with relevant industry standards and regulations."
    story.append(Paragraph(p4, STYLE_BODY))

    p5 = "Thank you for your attention to this matter. I look forward to receiving your response and arranging the necessary meter testing for my solar PV system."
    story.append(Paragraph(p5, STYLE_BODY))
    story.append(Spacer(1, 0.4 * cm))

    owner_name = (company.get('owner_name') or company.get('proprietor_name') or company.get('authorized_signatory') or company.get('manager_name') or '').strip()

    story.append(Paragraph("Thanks & Regards,", STYLE_BODY))
    story.append(Spacer(1, 0.6 * cm))
    story.append(Paragraph(f"<b>{company_name.upper()}</b>", STYLE_BOLD_SERIF))
    if owner_name:
        story.append(Paragraph(f"<b>{owner_name}</b>", STYLE_BODY))
    story.append(Spacer(1, 0.8 * cm))

    encl_text = "<b>Encl:</b><br/>1. Gen-meter<br/>2. Test report of meter<br/>3. Electricity Bill<br/>4. Solar PV System Approval Latter Copy."
    story.append(Paragraph(encl_text, STYLE_BODY))

    canvas_cls = make_meter_testing_canvas(company)
    pdf.build(story, canvasmaker=canvas_cls)
    return buf.getvalue()



def generate(doc_type: str, client: dict, company: dict) -> bytes:
    doc_type_clean = (doc_type or "").lower().strip()
    if doc_type_clean == "wcr":
        return generate_wcr_pdf(client, company)
    if doc_type_clean == "sldr":
        return generate_sldr_pdf(client, company)
    if doc_type_clean == "net_meter_agreement":
        return generate_net_meter_agreement_pdf(client, company)
    if doc_type_clean in ("vendor_agreement", "vendor"):
        return generate_vendor_agreement_pdf(client, company)
    if doc_type_clean in ("meter_testing_request", "meter_testing"):
        return generate_meter_testing_request_pdf(client, company)
    if doc_type_clean == "annexure":
        try:
            import annexure_generator
            pdf_bytes, content_type = annexure_generator.generate_annexure(client, company)
            return pdf_bytes
        except Exception as _e:
            import logging as _logging
            _logging.getLogger(__name__).error(f"DOCX-based Annexure generation failed, using ReportLab fallback: {_e}")

    buf = BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=A4, leftMargin=1.5*cm, rightMargin=1.5*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
    story: list = _header(company)

    title_map = {
        "annexure": "ANNEXURE — Material & Site Details",
        "wcr": "WORK COMPLETION REPORT (WCR)",
        "sldr": "SINGLE LINE DIAGRAM REPORT (SLDR)",
        "net_meter_agreement": "NET METER AGREEMENT",
        "vendor_agreement": "VENDOR AGREEMENT",
        "quotation": "SOLAR PV SYSTEM QUOTATION",
        "installation_report": "INSTALLATION & COMMISSIONING REPORT",
        "completion_report": "FINAL SYSTEM COMPLETION REPORT",
    }
    story.append(Paragraph(title_map.get(doc_type_clean, doc_type_clean.upper()), H2))
    story.append(Paragraph(f"Document No.: <b>{doc_type_clean.upper()}</b> &nbsp;&nbsp; Date: <b>{datetime.now(timezone.utc).strftime('%d %b %Y')}</b>", SMALL))
    story.append(Spacer(1, 0.4 * cm))

    pan_num_val = str(client.get("pan_number") or client.get("pan_card_number") or "").strip()
    aadhaar_num_val = str(client.get("aadhaar") or client.get("aadhaar_number") or "").strip()
    if pan_num_val:
        id_kv_row = ["PAN Card", pan_num_val]
    elif aadhaar_num_val:
        id_kv_row = ["Aadhaar (last 4)", (aadhaar_num_val)[-4:] or "—"]
    else:
        id_kv_row = ["Aadhaar (last 4)", "—"]

    story.append(Paragraph("Client Details", H2))
    story.append(_kv_table([
        ["Client Name", client.get("full_name", "")],
        ["Mobile", client.get("mobile", "")],
        ["Consumer Number", client.get("consumer_number", "—")],
        ["Address", f"{client.get('address','')}, {client.get('city','')}, {client.get('state','')} - {client.get('pincode','')}"],
        id_kv_row,
    ]))
    story.append(Spacer(1, 0.4 * cm))

    story.append(Paragraph("System Specifications", H2))
    _ann_inv_list = _get_inverters_list(client)
    if _ann_inv_list:
        _inv_rows = []
        for _i, _inv in enumerate(_ann_inv_list):
            _lbl = "Inverter" if len(_ann_inv_list) == 1 else f"Inverter #{_i+1}"
            _brand = str(_inv.get("brand") or client.get("inverter_make") or "").strip()
            _cap = str(_inv.get("capacity") or "").strip()
            _qty = str(_inv.get("quantity") or "1").strip()
            _inv_rows.append([_lbl, f"{_brand} · {_cap} kW × {_qty}".strip(" ·")])
        story.append(_kv_table([
            ["System Size", f"{client.get('system_kw',0)} kW"],
            ["Phase Type", client.get("phase_type", "")],
            ["Subsidy Eligible", "Yes" if client.get("subsidy_eligible") else "No"],
            ["Panel", f"{client.get('panel_make','')} · {client.get('panel_wattage','')}W × {client.get('num_panels','')}"],
            *_inv_rows,
            ["Inverter Serial", client.get("inverter_serial", "—")],
        ]))
    else:
        story.append(_kv_table([
            ["System Size", f"{client.get('system_kw',0)} kW"],
            ["Phase Type", client.get("phase_type", "")],
            ["Subsidy Eligible", "Yes" if client.get("subsidy_eligible") else "No"],
            ["Panel", f"{client.get('panel_make','')} · {client.get('panel_wattage','')}W × {client.get('num_panels','')}"],
            ["Inverter", f"{client.get('inverter_make','')} · {client.get('inverter_capacity','')}"],
            ["Inverter Serial", client.get("inverter_serial", "—")],
        ]))
    story.append(Spacer(1, 0.5 * cm))

    if doc_type_clean == "annexure":
        story.append(Paragraph("Material Annexure", H2))
        story.append(Paragraph(
            "This annexure certifies that the following materials have been used in the installation as per the agreed BOM. "
            "Quantities and serial numbers reflect the field verification report.", BODY,
        ))
    elif doc_type_clean == "sldr":
        story.append(Paragraph("Single Line Diagram Summary", H2))
        story.append(Paragraph(
            "DC side: Solar panels → DCDB (with surge arrester & DC isolator) → Inverter MPPT input. "
            "AC side: Inverter AC output → ACDB (with MCB + RCBO) → Net Meter → DISCOM grid. "
            "Earthing: Separate earth pits for AC, DC and lightning arrester as per IS 3043.",
            BODY,
        ))
    elif doc_type_clean in ("net_meter_agreement", "vendor_agreement"):
        story.append(Paragraph("Agreement Terms & Undertaking", H2))
        story.append(Paragraph(
            "1. The consumer agrees to install a bi-directional net meter at their premises.<br/>"
            "2. Excess generation will be credited as per the prevailing DISCOM tariff.<br/>"
            "3. Annual settlement will be carried out by the DISCOM as per state regulations.<br/>"
            "4. The vendor agrees to provide 5-year Comprehensive Maintenance Contract (CMC) coverage.",
            BODY,
        ))

    story.append(Spacer(1, 1.2 * cm))
    sign = Table([
        [Paragraph("<b>Customer Signature</b><br/><br/><br/>_____________________<br/>" + client.get("full_name", ""), SMALL),
         Paragraph("<b>Authorized Signatory</b><br/><br/><br/>_____________________<br/>" + company.get("company_name", "") + "<br/>" + company.get("owner_name", ""), SMALL)],
    ], colWidths=[9*cm, 9*cm])
    sign.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(sign)

    pdf.build(story)
    return buf.getvalue()


def generate_ledger_pdf(client: dict, ledger: dict, company: dict) -> bytes:
    buf = BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=A4, leftMargin=1.5 * cm, rightMargin=1.5 * cm, topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    story: list = _header(company)
    
    story.append(Paragraph("<b>CLIENT MATERIAL LEDGER REPORT</b>", H1))
    story.append(Spacer(1, 0.2 * cm))
    
    details = [
        ["Client Name", client.get("full_name", "")],
        ["Project ID", client.get("sol_id") or client.get("client_code") or ""],
        ["Generated Date", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
    ]
    story.append(_table(details, col_widths=[5 * cm, 13 * cm]))
    story.append(Spacer(1, 0.4 * cm))
    
    story.append(Paragraph("Ledger Summary", H2))
    summary = ledger.get("summary") or {}
    summary_data = [
        ["Total Products", str(summary.get("total_products", 0))],
        ["Total Outward Qty", str(summary.get("total_outward_qty", 0))],
        ["Total Returned Qty", str(summary.get("total_returned_qty", 0))],
        ["Current Balance", str(summary.get("current_balance", 0))],
        ["Negative Items", str(summary.get("negative_items", 0))]
    ]
    story.append(_table(summary_data, col_widths=[9 * cm, 9 * cm]))
    story.append(Spacer(1, 0.4 * cm))
    
    story.append(Paragraph("Material Details", H2))
    
    style_normal = ParagraphStyle('normal_cell', parent=styles['Normal'], fontSize=8, textColor=colors.HexColor('#000000'))
    style_red = ParagraphStyle('red_cell', parent=styles['Normal'], fontSize=8, textColor=colors.HexColor('#dc2626'))
    style_gray = ParagraphStyle('gray_cell', parent=styles['Normal'], fontSize=8, textColor=colors.HexColor('#94a3b8'))
    
    headers = [
        Paragraph('<font color="#ffffff"><b>Product</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Size</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Unit</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Outward</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Returned</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Balance</b></font>', HEADER_TEXT_STYLE),
        Paragraph('<font color="#ffffff"><b>Status</b></font>', HEADER_TEXT_STYLE),
    ]
    rows: list = [headers]
    
    for item in ledger.get("items") or []:
        bal = float(item.get("current_balance") or 0)
        cstyle = style_normal
        if bal < 0:
            cstyle = style_red
        elif bal == 0:
            cstyle = style_gray
            
        rows.append([
            Paragraph(str(item.get("product", "")), cstyle),
            Paragraph(str(item.get("size", "") or ""), cstyle),
            Paragraph(str(item.get("unit", "") or "Nos"), cstyle),
            Paragraph(str(item.get("total_outward", 0)), cstyle),
            Paragraph(str(item.get("total_returned", 0)), cstyle),
            Paragraph(str(item.get("current_balance", 0)), cstyle),
            Paragraph(str(item.get("status", "")), cstyle),
        ])
        
    story.append(_table(rows, col_widths=[5.5 * cm, 2.5 * cm, 1.5 * cm, 2 * cm, 2 * cm, 2 * cm, 2.5 * cm], header_row=True))
    
    pdf.build(story)
    return buf.getvalue()


def make_product_master_canvas(company: dict):
    company_name = (company.get("company_name") or company.get("name") or "Solar App").strip()

    class ProductMasterCanvas(canvas.Canvas):
        _startPage: Any
        _pageNumber: Any

        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            self.setStrokeColor(colors.HexColor('#cbd5e1'))
            self.setLineWidth(0.5)
            # A4 Landscape dimensions: width 29.7 cm, height 21.0 cm
            self.line(1.2 * cm, 1.2 * cm, 29.7 * cm - 1.2 * cm, 1.2 * cm)

            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor('#475569'))
            if company_name:
                self.drawString(1.2 * cm, 0.8 * cm, company_name)
            self.drawCentredString(14.85 * cm, 0.8 * cm, "Product Master Inventory Report")
            self.drawRightString(29.7 * cm - 1.2 * cm, 0.8 * cm, f"Page {self._pageNumber} of {page_count}")
            self.restoreState()

    return ProductMasterCanvas


def generate_product_master_pdf(products: list, company: dict) -> bytes:
    buf = BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=landscape(A4),
        leftMargin=1.0 * cm,
        rightMargin=1.0 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.5 * cm
    )
    story = []

    story.extend(_header(company))

    STYLE_TITLE = ParagraphStyle('pm_t', parent=styles['Normal'], fontSize=14, fontName='Helvetica-Bold', leading=17, textColor=colors.HexColor('#0f172a'))
    STYLE_SUBTITLE = ParagraphStyle('pm_st', parent=styles['Normal'], fontSize=8.5, fontName='Helvetica', leading=12, textColor=colors.HexColor('#475569'))

    now_str = datetime.now().strftime("%d %b %Y, %I:%M %p")
    story.append(Paragraph("<b>PRODUCT MASTER INVENTORY REPORT</b>", STYLE_TITLE))
    story.append(Paragraph(f"Exported Date & Time: {now_str} | Total Products Exported: <b>{len(products)}</b>", STYLE_SUBTITLE))
    story.append(Spacer(1, 0.25 * cm))

    style_cell = ParagraphStyle('cell_norm', parent=styles['Normal'], fontSize=7.5, fontName='Helvetica', leading=9.5, textColor=colors.HexColor('#1e293b'))
    style_cell_bold = ParagraphStyle('cell_bold', parent=styles['Normal'], fontSize=7.5, fontName='Helvetica-Bold', leading=9.5, textColor=colors.HexColor('#0f172a'))
    style_cell_right = ParagraphStyle('cell_r', parent=styles['Normal'], fontSize=7.5, fontName='Helvetica', leading=9.5, alignment=2, textColor=colors.HexColor('#1e293b'))
    style_cell_center = ParagraphStyle('cell_c', parent=styles['Normal'], fontSize=7.5, fontName='Helvetica', leading=9.5, alignment=1, textColor=colors.HexColor('#1e293b'))

    def hdr_cell(txt: str, align: Literal[0, 1, 2, 4] = 0) -> Paragraph:
        return Paragraph(f"<b><font color='#ffffff' size='7.5'>{txt}</font></b>", ParagraphStyle('h', parent=styles['Normal'], alignment=align))

    table_data = [
        [
            hdr_cell("Product Name"),
            hdr_cell("Specification"),
            hdr_cell("Unit", 1),
            hdr_cell("HSN", 1),
            hdr_cell("Category"),
            hdr_cell("Opening Stock", 2),
            hdr_cell("Current Stock", 2),
            hdr_cell("High Value", 1),
            hdr_cell("Barcode", 1),
            hdr_cell("Status", 1)
        ]
    ]

    for p in products:
        name = str(p.get("name") or "")
        spec = str(p.get("size") or p.get("specification") or "—")
        unit = str(p.get("unit") or "Nos")
        hsn = str(p.get("hsn") or p.get("hsn_code") or p.get("sku") or "—")
        category = str(p.get("category") or "Solar")
        op_stock = str(p.get("opening_stock") if p.get("opening_stock") is not None else p.get("min_stock") or 0)
        cur_stock = str(p.get("balance") if p.get("balance") is not None else p.get("current_stock") or 0)
        is_hv = "Yes" if (p.get("high_value_goods") or p.get("is_high_value") or p.get("high_value")) else "No"
        barcode = str(p.get("barcode") or p.get("barcode_num") or p.get("code") or "—")
        status = str(p.get("stock_status") or p.get("status") or "Normal")

        table_data.append([
            Paragraph(name, style_cell_bold),
            Paragraph(spec, style_cell),
            Paragraph(unit, style_cell_center),
            Paragraph(hsn, style_cell_center),
            Paragraph(category, style_cell),
            Paragraph(op_stock, style_cell_right),
            Paragraph(cur_stock, style_cell_right),
            Paragraph(is_hv, style_cell_center),
            Paragraph(barcode, style_cell_center),
            Paragraph(status, style_cell_center)
        ])

    col_widths = [5.0 * cm, 3.0 * cm, 1.6 * cm, 2.2 * cm, 2.8 * cm, 2.4 * cm, 2.4 * cm, 2.0 * cm, 2.6 * cm, 2.4 * cm]
    t = Table(table_data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),
    ]))

    story.append(t)
    pdf.build(story, canvasmaker=make_product_master_canvas(company))
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# DOCX generator — mirrors PDF content using python-docx
# Does NOT touch any existing PDF generator functions.
# ─────────────────────────────────────────────────────────────────────────────

def _docx_heading(doc, text: str, level: int = 1):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        run.font.size = Pt(14 if level == 1 else 11)


def _docx_kv_table(doc, rows: list):
    """Add a 2-column label-value table (generic fallback)."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate(rows):
        cells = tbl.rows[i].cells
        cells[0].text = str(lbl)
        cells[1].text = str(val) if val is not None else ""
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    return tbl


def _docx_company_header(doc, company: dict):
    from docx.shared import Pt, RGBColor
    company_name = company.get("company_name") or company.get("name") or ""
    mobile = company.get("mobile") or company.get("phone") or ""
    email = company.get("email") or ""
    gst = company.get("gst_number") or company.get("gst") or ""
    address = company.get("address") or ""
    city = company.get("city") or ""
    state = company.get("state") or ""
    pincode = company.get("pincode") or ""
    full_address = ", ".join(p for p in [address, city, state] if p)
    if pincode:
        full_address += f" - {pincode}"
    p = doc.add_paragraph()
    run = p.add_run(company_name)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    if mobile:
        doc.add_paragraph(f"Mobile: {mobile}")
    if email:
        doc.add_paragraph(f"Email: {email}")
    if gst:
        doc.add_paragraph(f"GSTIN: {gst}")
    if full_address:
        doc.add_paragraph(f"Address: {full_address}")


# ─────────────────────────────────────────────────────────────────────────────
# Faithful DOCX helpers (match PDF layout exactly)
# ─────────────────────────────────────────────────────────────────────────────

def _emu(val: float | int) -> Any:
    from docx.shared import Emu
    return Emu(int(val))


def _build_docx_document(left_cm=1.2, right_cm=1.2, top_cm=1.0, bottom_cm=1.8):
    """Create a python-docx Document with A4 page size and given margins (in cm)."""
    from docx import Document as _Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    doc = _Document()
    for section in doc.sections:
        section.page_width  = _emu(21.0 * 360000)
        section.page_height = _emu(29.7 * 360000)
        section.left_margin   = _emu(left_cm   * 360000)
        section.right_margin  = _emu(right_cm  * 360000)
        section.top_margin    = _emu(top_cm    * 360000)
        section.bottom_margin = _emu(bottom_cm * 360000)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    return doc


def _docx_set_cell_bg(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def _remove_tbl_borders(tbl):
    """Remove all borders from a table."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = tbl._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'), 'none')
        bd.set(qn('w:sz'), '0')
        bd.set(qn('w:space'), '0')
        bd.set(qn('w:color'), 'auto')
        tblBorders.append(bd)
    tblPr.append(tblBorders)


def _docx_header_block(doc, company: dict):
    """
    Build the branded document header matching the PDF:
      [Logo | Company Name (large blue bold) | GST (right-aligned blue)]
      ─── solid blue divider line ───
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    company_name = (company.get('company_name') or company.get('name') or
                    company.get('legal_business_name') or '').strip()
    gst_no = (company.get('gst_number') or company.get('gstin') or
              company.get('gst') or '').strip()
    logo_bytes = company.get('logo_bytes')

    # 3-column header table: [logo | company name | GST]
    hdr_tbl = doc.add_table(rows=1, cols=3)
    hdr_tbl.style = 'Table Grid'
    _remove_tbl_borders(hdr_tbl)
    for idx, w_cm in enumerate([6.5, 7.5, 4.6]):
        hdr_tbl.columns[idx].width = _emu(w_cm * 360000)

    row = hdr_tbl.rows[0]
    logo_cell = row.cells[0]
    logo_cell.width = _emu(6.5 * 360000)
    if logo_bytes:
        try:
            from PIL import Image as _PILImage
            img = _PILImage.open(BytesIO(logo_bytes))
            iw, ih = img.size
            if iw > 0 and ih > 0:
                aspect = ih / float(iw)
                target_w = 6.5
                target_h = target_w * aspect
                if target_h > 3.2:
                    target_h = 3.2
                    target_w = target_h / aspect
                res_buf = BytesIO()
                resized = img.resize((int(target_w * 4 * 28.35), int(target_h * 4 * 28.35)))
                resized.save(res_buf, format='PNG')
                res_buf.seek(0)
                lp = logo_cell.paragraphs[0]
                lrun = lp.add_run()
                lrun.add_picture(res_buf, width=_emu(target_w * 360000))
        except Exception:
            logo_cell.paragraphs[0].add_run('')
    else:
        logo_cell.paragraphs[0].add_run('')

    name_cell = row.cells[1]
    name_cell.width = _emu(7.5 * 360000)
    name_para = name_cell.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_run = name_para.add_run(company_name.upper())
    name_run.bold = True
    name_run.font.size = Pt(15)
    name_run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)

    gst_cell = row.cells[2]
    gst_cell.width = _emu(4.6 * 360000)
    gst_para = gst_cell.paragraphs[0]
    gst_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if gst_no:
        gst_run = gst_para.add_run(f"GST NO - {gst_no}")
        gst_run.bold = True
        gst_run.font.size = Pt(9)
        gst_run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)

    # Blue divider line (thick top border on a minimal-height row)
    div_tbl = doc.add_table(rows=1, cols=1)
    div_tbl.style = 'Table Grid'
    _remove_tbl_borders(div_tbl)
    div_cell = div_tbl.rows[0].cells[0]
    div_cell.width = _emu(18.6 * 360000)
    tc = div_cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    top_bd = OxmlElement('w:top')
    top_bd.set(qn('w:val'), 'single')
    top_bd.set(qn('w:sz'), '12')
    top_bd.set(qn('w:space'), '0')
    top_bd.set(qn('w:color'), '1D4ED8')
    tcBorders.append(top_bd)
    for side in ['bottom', 'left', 'right']:
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'), 'none')
        bd.set(qn('w:sz'), '0')
        bd.set(qn('w:space'), '0')
        bd.set(qn('w:color'), 'auto')
        tcBorders.append(bd)
    tcPr.append(tcBorders)
    trPr = div_tbl.rows[0]._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '80')
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(2)


def _docx_add_spacer(doc, height_pt=6):
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(height_pt)


def _docx_add_body_paragraph(doc, text: str, font_size=9, justify=False):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    return p


def _docx_add_section_title(doc, text, font_size=11, bold=True, center=False):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(3)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    return p


def _docx_signature_block(doc, client_name: str, company_name: str):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _docx_add_spacer(doc, 18)
    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.style = 'Table Grid'
    _remove_tbl_borders(sig_tbl)
    sig_tbl.rows[0].cells[0].width = _emu(9.3 * 360000)
    sig_tbl.rows[0].cells[1].width = _emu(9.3 * 360000)
    lp = sig_tbl.rows[0].cells[0].paragraphs[0]
    lr = lp.add_run(f"Authorized Signature [Vendor]\n\n\nFor {company_name.upper()}")
    lr.bold = True
    lr.font.size = Pt(9)
    rp = sig_tbl.rows[0].cells[1].paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(f"Consumer Signature\n\n\n{client_name}")
    rr.bold = True
    rr.font.size = Pt(9)

    under_tbl = doc.add_table(rows=1, cols=2)
    under_tbl.style = 'Table Grid'
    _remove_tbl_borders(under_tbl)
    under_tbl.rows[0].cells[0].width = _emu(9.3 * 360000)
    under_tbl.rows[0].cells[1].width = _emu(9.3 * 360000)
    under_tbl.rows[0].cells[0].paragraphs[0].add_run("_________________________")
    rp2 = under_tbl.rows[0].cells[1].paragraphs[0]
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.add_run("_________________________")


# ─── Per-document faithful DOCX generators ───────────────────────────────────

def _generate_wcr_docx(client: dict, company: dict) -> bytes:
    """WCR Word document mirroring the 3-page PDF layout."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = _build_docx_document(left_cm=1.2, right_cm=1.2, top_cm=1.0, bottom_cm=1.8)

    stages_dict = dict(client.get("stages") or {})
    ob_dict = dict(stages_dict.get("onboarding_data") or {})
    company_name = (company.get('company_name') or '').strip()
    client_name  = (client.get('full_name') or client.get('name') or '').strip()
    consumer_num = str(client.get('consumer_number') or '').strip()
    client_addr  = (client.get('address') or '').strip()
    city         = (client.get('city') or '').strip()
    pincode      = str(client.get('pincode') or '').strip()
    site_addr    = f"{client_addr}{', ' + city if city else ''}{' - ' + pincode if pincode else ''}".strip(', -')
    category     = (client.get('consumer_type') or client.get('consumer_category') or ob_dict.get('consumer_type') or '').strip()
    section_no   = str(client.get('section_number') or client.get('section_no') or ob_dict.get('section_number') or '').strip()
    sol_kw       = str(client.get('system_kw') or client.get('capacity') or '').strip()
    sol_kw_str   = f"{sol_kw} KW" if sol_kw else ''
    sol_wp       = str(client.get('panel_wattage') or ob_dict.get('panel_wattage') or '').strip()
    sol_wp_str   = f"{sol_wp} WP" if sol_wp else ''
    num_panels   = str(client.get('num_panels') or ob_dict.get('num_panels') or '').strip()
    num_panels_str = f"{num_panels} NOS" if num_panels else ''
    panel_make   = (client.get('panel_brand') or client.get('panel_make') or ob_dict.get('panel_brand') or ob_dict.get('panel_make') or '').strip()
    panel_tech   = (client.get('panel_technology') or ob_dict.get('panel_technology') or '').strip()
    sol_wp_tech_str = f"{sol_wp_str} / {panel_tech}" if (sol_wp_str and panel_tech) else (sol_wp_str or panel_tech)
    almm_model   = str(client.get('almm_model_number') or sol_wp_tech_str).strip()
    inverter_list = _get_inverters_list(client)
    brands = []
    if inverter_list:
        for inv in inverter_list:
            b = str(inv.get("brand") or "").strip()
            if b and b not in brands:
                brands.append(b)
    if not brands:
        fb = str(client.get('inverter_make') or ob_dict.get('inverter_make') or '').strip()
        if fb:
            brands = [fb]
    inverter_make = ", ".join(brands) if brands else ''
    all_serials = []
    if inverter_list:
        for inv in inverter_list:
            inv_s = inv.get("serials")
            if isinstance(inv_s, list):
                for s in inv_s:
                    s_str = str(s).strip()
                    if s_str and s_str not in all_serials:
                        all_serials.append(s_str)
            elif inv.get("serial"):
                for part in str(inv.get("serial")).split(","):
                    p_str = part.strip()
                    if p_str and p_str not in all_serials:
                        all_serials.append(p_str)
    if not all_serials:
        raw_sr = str(client.get('inverter_serial') or ob_dict.get('inverter_serial') or '').strip()
        if raw_sr:
            all_serials = [p.strip() for p in raw_sr.split(",") if p.strip()]
    inverter_sr = ", ".join(all_serials) if all_serials else ''
    inverter_kw = str(client.get('inverter_capacity') or ob_dict.get('inverter_capacity') or '').strip()
    inverter_kw_str = f"{inverter_kw} KW" if (inverter_kw and "KW" not in inverter_kw.upper()) else inverter_kw
    inverter_year = str(client.get('inverter_year') or ob_dict.get('inverter_year') or '').strip()
    pan_num = str(client.get('pan_number') or client.get('pan_card_number') or ob_dict.get('pan_number') or '').strip()
    aadhaar_num = str(client.get('aadhaar') or client.get('aadhaar_number') or ob_dict.get('aadhaar') or '').strip()

    # PAGE 1
    _docx_header_block(doc, company)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after  = Pt(8)
    tr = title_p.add_run("Work Completion Report for Solar Power Plant")
    tr.bold = True; tr.font.size = Pt(13)
    tr.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    insp_rows = [
        ("Sr.No", "Component", "Observation"),
        ("1", "Name", client_name),
        ("2", "Consumer number", consumer_num),
        ("3", "Site/Location with Complete Address", site_addr),
        ("4", "Category: Govt/Private Sector", category),
        ("5", "Section number", section_no),
        ("6", "Sanctioned Capacity of solar PV system (KW) Installed", sol_kw_str),
        ("",  "Capacity of solar PV system (KW)", sol_kw_str),
        ("SPECIFICATION OF THE MODULES", "", ""),  # sub-header row 8
        ("7", "Make & Type of modules", panel_make),
        ("",  "ALMM Model Number", almm_model),
        ("",  "Wattage per module", sol_wp_str),
        ("",  "No. of Module", num_panels_str),
        ("",  "Total Capacity (KWP)", sol_kw_str),
        ("",  "Warrantee Details (Product + Performance)", "12+15 YEARS" if sol_kw else ""),
        ("PCU", "", ""),  # sub-header row 15
        ("8", "Make & Model number of Inverter", inverter_make),
        ("",  "Rating", inverter_kw_str),
        ("",  "Type of charge controller/ MPPT", "MPPT" if inverter_make else ""),
        ("",  "Capacity of Inverter", inverter_kw_str),
        ("",  "SR Number", inverter_sr),
        ("",  "Year of manufacturing", inverter_year),
        ("EARTHING & PROTECTION", "", ""),  # sub-header row 22
        ("9", "No of Separate Earthings with earth Resistance", "NON_TRACKING" if sol_kw else ""),
        ("",  "It is certified that the Earth Resistance measure in presence of Licensed Electrical "
              "Contractor/Supervisor and found in order i.e. < 5 Ohms as per MNRE OM Dtd. 07.06.24 for CFA Component.", ""),
        ("",  "Lightening Arrester", "Yes" if sol_kw else ""),
    ]
    SUB_HDR = {8, 15, 22}
    insp_tbl = doc.add_table(rows=len(insp_rows), cols=3)
    insp_tbl.style = 'Table Grid'
    insp_tbl.columns[0].width = _emu(0.7 * 360000)
    insp_tbl.columns[1].width = _emu(8.8 * 360000)
    insp_tbl.columns[2].width = _emu(9.1 * 360000)
    for r_idx, (sr, comp, obs) in enumerate(insp_rows):
        row_cells = insp_tbl.rows[r_idx].cells
        if r_idx in SUB_HDR:
            merged = row_cells[0].merge(row_cells[1]).merge(row_cells[2])
            _docx_set_cell_bg(merged, 'f1f5f9')
            sp = merged.paragraphs[0]
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            srn = sp.add_run(sr)
            srn.bold = True; srn.font.size = Pt(9)
            srn.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
            continue
        for c_idx, text in enumerate([sr, comp, obs]):
            cell = row_cells[c_idx]
            p = cell.paragraphs[0]
            if r_idx == 0:
                _docx_set_cell_bg(cell, 'e2e8f0')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(text))
                run.bold = True; run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
            else:
                run = p.add_run(str(text))
                run.font.size = Pt(8.5)
                if c_idx in (0, 2):
                    run.bold = True
                run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    _docx_add_spacer(doc, 8)
    _docx_signature_block(doc, client_name, company_name)

    # PAGE 2 – Declaration
    doc.add_page_break()
    _docx_header_block(doc, company)
    _docx_add_spacer(doc, 6)
    discom_code = client.get("discom_code") or client.get("discom") or "MSEDCL"
    _docx_add_body_paragraph(doc,
        f"We {company_name} [Vendor] & {client_name} [Consumer] bearing Consumer Number "
        f"{consumer_num} Ensured structural stability of installed solar power plant and obtained "
        "requisite permissions from the concerned authority. If in future, by virtue of any means "
        f"due to collapsing or damage to the installed solar power plant, {discom_code} will not be held "
        "responsible for any loss to property or human life, if any.",
        font_size=9.5, justify=True)
    _docx_add_spacer(doc, 4)
    _docx_add_body_paragraph(doc,
        "This is to Certify above Installed Solar PV System is working properly with electrical safety & "
        "Islanding switch in case of any presence of backup inverter an arrangement should be made in "
        "such way the backup inverter supply should never be synchronized with solar inverter to avoid "
        "any electrical accident due to back feeding. We will be held responsible for non-working of "
        "islanding mechanism and back feed to the de-energized grid.",
        font_size=9.5, justify=True)
    _docx_add_spacer(doc, 18)
    _docx_signature_block(doc, client_name, company_name)

    # PAGE 3 – Guarantee Certificate
    doc.add_page_break()
    _docx_header_block(doc, company)
    _docx_add_spacer(doc, 4)
    cert_title = doc.add_paragraph()
    cert_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ct_run = cert_title.add_run("Guarantee Certificate Undertaking to be submitted by VENDOR")
    ct_run.bold = True; ct_run.font.size = Pt(13)
    ct_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    _docx_add_spacer(doc, 4)
    _docx_add_body_paragraph(doc,
        "The undersigned will provide services to the consumers for repairs/maintenance of the "
        "RTS plant free of cost for 5 years of the comprehensive Maintenance Contract (CMC) period "
        "from the date of commissioning of the plant. Nonperforming/under-performing system "
        "components will be replaced/repaired free of cost in the CMC period",
        font_size=9.5, justify=True)
    _docx_add_spacer(doc, 8)
    if pan_num:
        id_title = "[ CONSUMER PAN CARD / IDENTITY VERIFICATION ]"
        id_label = "PAN CARD"
        id_detail = f"PAN Number: {pan_num}"
    elif aadhaar_num:
        id_title = "[ CONSUMER AADHAAR CARD / IDENTITY VERIFICATION ]"
        id_label = "ADHAR CARD"
        id_detail = f"Aadhar Number: {aadhaar_num}"
    else:
        id_title = "[ CONSUMER IDENTITY VERIFICATION ]"
        id_label = "IDENTITY CARD"
        id_detail = ""
    id_box_tbl = doc.add_table(rows=2, cols=1)
    id_box_tbl.style = 'Table Grid'
    id_box_tbl.columns[0].width = _emu(14.0 * 360000)
    hdr_cell = id_box_tbl.rows[0].cells[0]
    _docx_set_cell_bg(hdr_cell, 'eff6ff')
    hp = hdr_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(id_title)
    hr.bold = True; hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    body_cell = id_box_tbl.rows[1].cells[0]
    _docx_set_cell_bg(body_cell, 'f8fafc')
    bp = body_cell.paragraphs[0]
    brun = bp.add_run(f"Stamp & Seal\n\nIdentity Details of Consumer: - {id_label}")
    brun.bold = True; brun.font.size = Pt(8.5)
    if id_detail:
        drun = bp.add_run(f"\n{id_detail}")
        drun.bold = True; drun.font.size = Pt(8.5)
    _docx_add_spacer(doc, 18)
    _docx_signature_block(doc, client_name, company_name)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_sldr_docx(client: dict, company: dict) -> bytes:
    """SLDR Word document mirroring the PDF layout."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = _build_docx_document(left_cm=1.2, right_cm=1.2, top_cm=1.2, bottom_cm=1.8)
    company_name = (company.get('company_name') or '').strip().upper()
    client_name  = str(client.get('full_name') or client.get('name') or '').upper()
    consumer_num = str(client.get('consumer_number') or '').upper()
    bu_num       = str(client.get('bu_number') or '').upper()
    sol_kw       = str(client.get('system_kw') or '').strip()
    sol_wp       = str(client.get('panel_wattage') or '').strip()
    num_panels   = str(client.get('num_panels') or '').strip()
    panel_make   = (client.get('panel_brand') or client.get('panel_make') or '').upper()
    stages_dict  = dict(client.get("stages") or {})
    ob_dict      = dict(stages_dict.get("onboarding_data") or {})

    inverter_list = _get_inverters_list(client)
    first_inv    = inverter_list[0] if inverter_list else {}
    inverter_make = (first_inv.get("brand") or client.get('inverter_make') or '').upper()

    raw_ob_cap   = str(client.get('inverter_capacity') or ob_dict.get('inverter_capacity') or '').strip()
    if raw_ob_cap:
        inverter_kw = raw_ob_cap.upper() if "KW" in raw_ob_cap.upper() else f"{raw_ob_cap} KW"
    else:
        inverter_kw = f"{sol_kw} KW" if sol_kw else ""

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("SINGLE LINE DIAGRAM")
    title_run.bold = True; title_run.underline = True
    title_run.font.size = Pt(15)
    title_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    _docx_add_spacer(doc, 4)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(4)
    for label, value in [
        ("CONSUMER NAME :-", f" {client_name}  "),
        ("CONSUMER NO.:-",   f" {consumer_num}  "),
        ("B.U.:-",           f" {bu_num}\n"),
        ("PROJECT:-",        f" GCRT OF {sol_kw} KW"),
    ]:
        lr = meta_p.add_run(label)
        lr.bold = True; lr.font.size = Pt(10)
        lr.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        vr = meta_p.add_run(value)
        vr.font.size = Pt(10)
        vr.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    _docx_add_spacer(doc, 4)
    # Bordered diagram area box
    diag_tbl = doc.add_table(rows=1, cols=1)
    diag_tbl.style = 'Table Grid'
    diag_cell = diag_tbl.rows[0].cells[0]
    diag_cell.width = _emu(18.6 * 360000)
    trPr = diag_tbl.rows[0]._tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), '3969')
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)
    note_p = diag_cell.paragraphs[0]
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note_p.add_run(
        f"Grid Tied Solar Inverter System Electrical Single Line Diagram\n\n"
        f"PV Array: {num_panels} x {sol_wp}Wp  |  Inverter: {inverter_make} ({inverter_kw})\n\n"
        f"DC side: PV Array -> DCDB (Surge Arrester + DC Isolator) -> Inverter MPPT\n"
        f"AC side: Inverter -> ACDB (MCB + RCBO) -> Net Meter -> DISCOM Grid\n"
        f"Earthing: Separate earth pits for AC, DC and Lightning Arrester (IS 3043)"
    )
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    _docx_add_spacer(doc, 6)
    _docx_add_section_title(doc, "TECHNICAL SPECIFICATIONS", font_size=9.5)

    tech_rows_data = [["PARAMETER", "SPECIFICATIONS", "MAKE", "KWP"]]
    tech_rows_data.append(["PV MODULES", f"{sol_wp} Wp X {num_panels} Nos", panel_make, f"{sol_kw} KW"])
    if inverter_list:
        for idx, inv in enumerate(inverter_list):
            inv_label = f"INVERTER #{idx+1}" if len(inverter_list) > 1 else "INVERTER"
            c_val = inv.get("capacity") or ""
            cap_str = c_val if "KW" in c_val.upper() else (f"{c_val} kW" if c_val else "")
            q_val = inv.get("quantity") or "1"
            spec_str = f"{cap_str} x {q_val} Nos" if cap_str else f"{q_val} Nos"
            make_str = f"{inv.get('brand','')} {inv.get('model','')}".strip().upper()
            tech_rows_data.append([inv_label, spec_str, make_str, inverter_kw])
    else:
        inv_kw_d = inverter_kw if "KW" in inverter_kw else (f"{inverter_kw} kW" if inverter_kw else "")
        tech_rows_data.append(["INVERTER", f"{inv_kw_d} x 1 Nos", inverter_make, inverter_kw])

    tech_tbl = doc.add_table(rows=len(tech_rows_data), cols=4)
    tech_tbl.style = 'Table Grid'
    for c_idx, w in enumerate([4.5, 5.5, 4.6, 4.0]):
        tech_tbl.columns[c_idx].width = _emu(w * 360000)
    for r_idx, row_data in enumerate(tech_rows_data):
        for c_idx, text in enumerate(row_data):
            cell = tech_tbl.rows[r_idx].cells[c_idx]
            if r_idx == 0:
                _docx_set_cell_bg(cell, 'f8fafc')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True; run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a) if r_idx == 0 else RGBColor(0x1e, 0x29, 0x3b)

    _docx_add_spacer(doc, 10)
    ftr_tbl = doc.add_table(rows=1, cols=2)
    ftr_tbl.style = 'Table Grid'
    _remove_tbl_borders(ftr_tbl)
    ftr_tbl.rows[0].cells[0].width = _emu(9.3 * 360000)
    ftr_tbl.rows[0].cells[1].width = _emu(9.3 * 360000)
    lp = ftr_tbl.rows[0].cells[0].paragraphs[0]
    lr = lp.add_run(company_name)
    lr.bold = True; lr.font.size = Pt(8.5)
    lr.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    rp = ftr_tbl.rows[0].cells[1].paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run("___________________________\n\nConsumer / Authorized Signature")
    rr.bold = True; rr.font.size = Pt(8.5)
    rr.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_net_meter_docx(client: dict, company: dict) -> bytes:
    """Net Meter Agreement Word document mirroring the PDF layout."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = _build_docx_document(left_cm=1.6, right_cm=1.6, top_cm=1.3, bottom_cm=1.3)
    company_name = company.get("company_name") or ""
    client_name  = (client.get("full_name") or client.get("name") or "").strip()
    consumer_no  = str(client.get("consumer_number") or "").strip()
    city         = (client.get("city") or "").strip()
    client_addr  = (client.get("address") or "").strip()
    pincode      = str(client.get("pincode") or "").strip()
    full_address = f"{client_addr}{', ' + city if city else ''}{' - ' + pincode if pincode else ''}".strip(", -")
    system_kw    = str(client.get("system_kw") or client.get("capacity") or "").strip()
    date_str_raw = client.get("installation_date") or datetime.now().strftime("%d/%m/%Y")
    date_str     = date_str_raw[:10] if len(date_str_raw) > 10 else date_str_raw
    raw_bu = str(client.get("bu_number") or client.get("bu_no") or client.get("bu") or "").strip()
    if raw_bu:
        bu_no = raw_bu if raw_bu.upper().startswith("BU-") else f"BU-{raw_bu}"
    else:
        bu_no = ""
    bu_text = client.get("bu_text") or ""
    sub_div = bu_text or client.get("sub_division") or ""
    division = client.get("division") or ""
    if sub_div and bu_no:
        sub_bu_str = bu_no if sub_div.strip().upper() in ("BU", "BU-") else f"{sub_div} {bu_no}"
    else:
        sub_bu_str = sub_div or bu_no
    licensee_sub = f", {sub_bu_str}" if sub_bu_str else ""
    discom_code  = client.get("discom_code") or client.get("discom") or "MSEDCL"
    licensee_title = client.get("distribution_licensee") or f"Additional Executive Engineer{licensee_sub}, {discom_code}"


    def _hdr(text, size=14, center=True):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    def _body(text, size=8.5, justify=True):
        p = doc.add_paragraph()
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    def _clause(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(1.5)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    _hdr("ANNEXURE - 3", size=14)
    _hdr("Net Metering Connection Agreement", size=12)
    _docx_add_spacer(doc, 8)
    _body(f"This Agreement is made and entered into at (location) {city} on this (date {date_str}) "
          f"between the Eligible Consumer {client_name} having premises at {full_address} "
          f"and Consumer No {consumer_no} as the first Party\nAND\n"
          f"The Distribution Licensee {licensee_title}, {discom_code}, "
          f"(hereinafter referred to as 'the Licensee') and having its Registered Office at {division} "
          f"as second Party of this Agreement;")
    _docx_add_spacer(doc, 4)
    _body("Whereas, the Eligible Consumer has applied to the Licensee for approval of a Net Metering Arrangement "
          "under the provisions of the State Electricity Regulatory Commission (Net Metering for Roof-top "
          "Solar Photo Voltaic Systems) Regulations and sought its "
          "connectivity to the Licensee's Distribution Network;")
    _docx_add_spacer(doc, 4)
    _body(f"And whereas, the Licensee has agreed to provide Network connectivity to the Eligible Consumer for "
          f"injection of electricity generated from its Roof-top Solar PV System of {system_kw} kilowatt;")
    _docx_add_spacer(doc, 4)
    _body("Both Parties hereby agree as follows:")

    _clause("1. Eligibility & System Standards:")
    _body("1.1 The Roof-top Solar PV System meets the applicable norms for being integrated into the Distribution Network, "
          "and that the Eligible Consumer shall maintain the System accordingly for the duration of this Agreement.")
    _body("1.2 The Solar PV system shall conform to the applicable State Net Metering Regulations, CEA Regulations, "
          "and relevant safety standards.")

    _clause("2. Technical and Inter-connection Requirements:")
    _body(f"2.1 The Solar PV system capacity is {system_kw} kW. Metering and inter-connection with the Network of the "
          f"Licensee shall be as per {discom_code} standards and norms specified by the Central Electricity Authority.")
    _body("2.2 The Eligible Consumer agrees to install prior to connection an isolation device (both automatic and manual); "
          "and the Licensee shall have access to it at all times for repair and maintenance of the Distribution Network.")
    _body("2.3 The Licensee shall specify the interface/inter-connection point and metering point.")

    _clause("3. Safety & Protection:")
    _body("3.1 The equipment connected to the Licensee's Distribution System shall be compliant with relevant IEEE/IEC or BIS standards.")
    _body("3.2 The design, installation, maintenance and operation of the Solar PV System shall ensure the safety of both "
          "the Solar PV System and the Licensee's Network.")
    _body("3.3 Consumer shall provide islanding mechanism to ensure Solar PV system disconnects automatically during grid failure.")

    _clause("4. Period of Agreement & Termination:")
    _body("4.1 This Agreement shall remain in force for a period of 25 years from commissioning unless terminated prematurely "
          "by mutual consent or 30 days written notice upon uncured breach.")

    _clause("5. Access and Disconnection:")
    _body("5.1 The Eligible Consumer shall provide access to authorized personnel of the Licensee to the metering equipment "
          "and disconnecting devices at all times.")
    _body("5.2 In emergent outage situations, the Licensee may disconnect supply to ensure grid safety.")

    _clause("6. Metering Arrangement & Energy Accounting:")
    _body("6.1 Net metering system shall include a bi-directional meter recording both export and import of electricity, "
          "and a Solar Generation Meter to measure total solar output.")
    _body("6.2 Energy accounting, billing and settlement of excess generation shall be strictly governed by applicable Regulations.")

    _clause("7. Dispute Resolution:")
    _body("7.1 Any dispute arising under this Agreement shall be resolved promptly and amicably, or through the concerned "
          "Consumer Grievance Redressal Forum.")

    _docx_add_spacer(doc, 4)
    _body(f"In the witness where of {client_name} for and on behalf of Eligible Consumer and Shri. "
          f"{licensee_title}, for and on behalf of {discom_code} agree to this agreement.")
    _docx_add_spacer(doc, 12)

    sig_tbl = doc.add_table(rows=4, cols=2)
    sig_tbl.style = 'Table Grid'
    _remove_tbl_borders(sig_tbl)
    sig_data = [
        (f"Signature of Eligible Consumer\n\n\n___________________________\n{client_name}\nEligible Consumer",
         f"Signature of Licensee\n\n\nShri. ___________________________\n{licensee_title}\nfor and on behalf of {discom_code}"),
        ("\nWitness 1: ___________________________", "\nWitness 1: ___________________________"),
        ("Witness 2: ___________________________", "Witness 2: ___________________________"),
        (f"\n\n{company_name}\nProprietor / Authorized Manager", "\n\nOfficial Stamp / Seal"),
    ]
    for r_idx, (lt, rt) in enumerate(sig_data):
        row = sig_tbl.rows[r_idx]
        row.cells[0].width = _emu(9.0 * 360000)
        row.cells[1].width = _emu(9.0 * 360000)
        lp = row.cells[0].paragraphs[0]
        lp.add_run(lt).font.size = Pt(8.5)
        rp = row.cells[1].paragraphs[0]
        rp.add_run(rt).font.size = Pt(8.5)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_vendor_docx(client: dict, company: dict) -> bytes:
    """Vendor Agreement Word document mirroring the PDF layout."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = _build_docx_document(left_cm=1.6, right_cm=1.6, top_cm=1.3, bottom_cm=1.3)
    company_name    = company.get("company_name") or ""
    company_address = company.get("address") or ""
    company_pincode = company.get("pincode") or ""
    client_name  = (client.get("full_name") or client.get("name") or "").strip()
    consumer_no  = str(client.get("consumer_number") or "").strip()
    city         = (client.get("city") or "").strip()
    client_addr  = (client.get("address") or "").strip()
    pincode      = str(client.get("pincode") or "").strip()
    full_address = f"{client_addr}{', ' + city if city else ''}{' - ' + pincode if pincode else ''}".strip(", -")
    system_kw    = str(client.get("system_kw") or client.get("capacity") or "").strip()
    panel_make   = (client.get("panel_brand") or client.get("panel_make") or "").strip()
    panel_wattage = str(client.get("panel_wattage") or "").strip()
    total_cost   = str(client.get("total_cost") or client.get("quotation_amount") or "").strip()
    discom_code  = client.get("discom_code") or client.get("discom") or "MSEDCL"
    date_obj = datetime.now()
    day_str = date_obj.strftime("%d")
    month_str = date_obj.strftime("%m")
    year_str = date_obj.strftime("%Y")

    def _hdr(text, size=14, center=True, underline=False):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(text)
        r.bold = True; r.underline = underline; r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    def _body(text, size=8.5, justify=True):
        p = doc.add_paragraph()
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    def _clause(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    _hdr("Agreement Between", size=14, underline=True)
    _hdr("Applicant and the registered/empaneled Vendor for installation of rooftop solar system in "
         "residential house of the Applicant under simplified procedure of Rooftop Solar Program Ph-II",
         size=9, center=True)
    _docx_add_spacer(doc, 4)
    _body(f"This agreement is executed on (Day) {day_str}, (Month) {month_str}, (Year) {year_str} "
          "for design, installation, commissioning and five years comprehensive maintenance of rooftop "
          "solar system to be installed under simplified procedure of Rooftop Solar Program Ph-II.")
    _docx_add_spacer(doc, 4)
    _hdr("Between", size=10)
    _body(f"{client_name} has residential electricity connection with consumer number {consumer_no} "
          f"from {discom_code} (DISCOM) at {full_address} PIN code : {pincode} (Hereinafter referred to as Applicant).")
    _docx_add_spacer(doc, 4)
    _hdr("And", size=10)
    _body(f"{company_name} (Name of Vendor) is registered/empaneled with the {discom_code} (hereinafter referred "
          f"as DISCOM) and is having registered/functional office at {company_address}. PIN CODE- {company_pincode}. "
          "Both Applicant and the Vendor are jointly referred as Parties.")
    _docx_add_spacer(doc, 4)
    for wh in [
        "- The Applicant intends to install rooftop solar system under simplified procedure of Rooftop Solar Programmed Ph-II of the MNRE.",
        f"- The Vendor is registered/empaneled vendor with {discom_code} for installation of rooftop solar under MNRE Schemes. The Vendor satisfies all the existing regulation pertaining to electrical safety.",
        "- Both the parties are mutually agreed and understand their roles and responsibilities and have no liability to any other agency/firm/stakeholder.",
    ]:
        _body(wh)
    _docx_add_spacer(doc, 6)

    _clause("1. GENERAL TERMS:")
    _body("1.1. The Applicant hereby represents and warrants that the Applicant has the sole legal capacity to enter into "
          "this Agreement and authorize the construction, installation and commissioning of the Rooftop Solar System "
          "(\"RTS System\") which is inclusive of Balance of System (\"BoS\") on the Applicant's premises.")
    _body("1.2. Vendor may propose changes to the scope, nature and or schedule of the services being performed under this Agreement.")
    _body("1.3. The Applicant understands and agrees that future changes in load, electricity usage patterns and/or "
          "electrical grid issues may affect the performance of the RTS System.")

    _clause("2. RTS System:")
    _body(f"2.1. Total capacity of RTS System will be minimum {system_kw} KWatt.")
    _body("2.2. The Solar modules, inverters and BoS will confirm to minimum specifications and DCR requirement of MNRE.")
    _body(f"2.3. Solar modules of {panel_make} make model, {panel_wattage} Wp capacity each and 21.13% efficiency will be procured and installed by the Vendor.")
    inv_mid = _get_vendor_inverter_clause_text(client)
    inv_str = f" {inv_mid}" if inv_mid else ""
    inv_clause_p = f"2.4. Solar inverter of{inv_str} rated output capacity will be procured and installed by the Vendor."
    _body(inv_clause_p)
    _body("2.5. The module mounting structure must withstand minimum wind load pressure as specified by MNRE.")
    _body("2.6. Other BoS installations shall be as per best industry practice with all safety and protection gears installed by the vendor.")

    _clause("3. PRICE AND PAYMENT TERMS:")
    _body(f"3.1. The cost of an RTS System will be Rs. {total_cost}/- (to be decided mutually). The Applicant shall pay the total cost to the Vendor as under:")
    _body("(i) 50 % as an advance on confirmation of the order.")
    _body("(ii) 40 % against Proforma Invoice (PI) before dispatch of solar panels, inverters and other BoS items.")
    _body("(iii) 10 % after installation and commissioning of the RTS System.")

    _clause("4. REPRESENTATIONS MADE BY THE APPLICANT:")
    _body("4.1. any timeline or schedule shared by Vendor for the provision of services and delivery of the RTS System is only an estimate.")
    _body("4.2. all information disclosed by the Applicant to Vendor in connection with the supply of the RTS System are true and accurate.")
    _body("4.3. all descriptive specifications, illustrations, drawings, data, dimensions, quotation, fact sheets, price lists and any advertising material circulated/published/provided by Vendor are approximate only.")

    _clause("5. MAINTENANCE:")
    _body("5.1. Vendor shall provide five-year free workmanship maintenance. Vendor shall visit the Applicant's premises at least once every quarter after commissioning of the RTS System for maintenance purposes.")
    _body("5.2. During such maintenance visit, Vendor shall check all nuts and bolts, fuses, earth resistance and other consumables in respect of the RTS System to ensure that it is in good working condition.")

    _clause("7. WARRANTIES:")
    _body("7.1. Product Warranty: The Applicant shall be entitled to manufacturers' warranty. Any warranty in relation to RTS System supplied to the Applicant by Vendor under this Agreement is limited to the warranty given by the manufacturer.")
    _body("7.2. Installation Warranty: Vendor warrants that all installations shall be free from workmanship defects or BOS defects for a period of five years from the date of installation of the RTS System.")
    _body("7.3. Subject to manufacturer warranty, Vendor warrants that the solar modules supplied herein shall have tolerance within a five-percentage range (+/-5%).")

    _clause("8. PERFORMANCE GUARANTEE:")
    _body("8.1. Vendor guarantees minimum system performance ratio of 75% as per performance ratio test carried out in adherence to IEC 61724 or equivalent BIS for a period of five years.")

    _clause("9. INSURANCE:")
    _body("9.1. Vendor may, at its sole discretion, obtain insurance covering risks of loss/damage to the RTS System during transit from Vendor's warehouse until delivery to the Applicant Site.")
    _body("9.2. Thereafter, all risk shall pass on to the Applicant and the Applicant may accordingly procure relevant insurances.")

    _clause("10. CANCELLATION:")
    _body("10.1. The Applicant may cancel the order placed on Vendor within 7 days from the date of remittance of advance money.")
    _body("10.2. If the Applicant cancels the order after the expiry of 7 days, the Applicant shall be liable to pay Vendor a cancellation fee of 30% of the total order value.")
    _body("10.3. Notwithstanding the aforesaid, the Applicant shall not be entitled to cancel the Order Form after Vendor has dispatched the RTS System to the Applicant Site.")

    _clause("15. GOVERNING LAW AND DISPUTE RESOLUTION:")
    _body("15.1. The interpretation and enforcement of this Agreement shall be governed by the laws of India.")
    _body("15.2. In the event of any dispute, controversy or difference between the Parties arising out of, or relating to this Agreement, both Parties shall make an effort to resolve the Dispute in good faith.")
    _body("15.3. The arbitration proceeding shall be governed by the provisions of the Arbitration and Conciliation Act, 1996 and shall be settled by a sole arbitrator mutually appointed by the Parties.")

    _docx_add_spacer(doc, 12)
    sig_tbl = doc.add_table(rows=2, cols=2)
    sig_tbl.style = 'Table Grid'
    _remove_tbl_borders(sig_tbl)
    sig_data = [
        (f"\n\n___________________________\n(Applicant)\n{client_name}",
         f"\n\n___________________________\n(Vendor)\n{company_name}"),
        ("", "\n\nOfficial Stamp / Seal"),
    ]
    for r_idx, (lt, rt) in enumerate(sig_data):
        row = sig_tbl.rows[r_idx]
        row.cells[0].width = _emu(9.0 * 360000)
        row.cells[1].width = _emu(9.0 * 360000)
        lp = row.cells[0].paragraphs[0]
        lp.add_run(lt).font.size = Pt(8.5)
        rp = row.cells[1].paragraphs[0]
        rp.add_run(rt).font.size = Pt(8.5)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_meter_testing_docx(client: dict, company: dict) -> bytes:
    """Meter Testing Request Word document mirroring the PDF layout."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = _build_docx_document(left_cm=1.2, right_cm=1.2, top_cm=1.0, bottom_cm=1.8)
    company_name = (company.get('company_name') or company.get('name') or '').strip()
    stages_dict = dict(client.get("stages") or {})
    ob_dict = dict(stages_dict.get("onboarding_data") or {})
    client_name  = (client.get('full_name') or client.get('name') or ob_dict.get('full_name') or '').strip()
    consumer_num = str(client.get('consumer_number') or ob_dict.get('consumer_number') or '').strip()
    client_addr  = (client.get('address') or ob_dict.get('address') or '').strip()
    city         = (client.get('city') or ob_dict.get('city') or '').strip()
    pincode      = str(client.get('pincode') or ob_dict.get('pincode') or '').strip()
    location_parts = []
    if client_addr:
        location_parts.append(client_addr)
    if city and pincode:
        location_parts.append(f"{city} - {pincode}")
    elif city:
        location_parts.append(city)
    elif pincode:
        location_parts.append(pincode)
    location_str = ", ".join(location_parts)

    def _clean(val):
        return "" if str(val).upper() in ("NA", "N/A", "0", "DEFAULT", "GROWATT", "NULL", "NONE") else val

    gen_make   = _clean((client.get('gen_meter_make') or ob_dict.get('gen_meter_make') or '').strip())
    gen_serial = _clean((client.get('gen_meter_serial') or ob_dict.get('gen_meter_serial') or '').strip())
    net_make   = _clean((client.get('net_meter_make') or ob_dict.get('net_meter_make') or '').strip())
    net_serial = _clean((client.get('net_meter_serial') or ob_dict.get('net_meter_serial') or '').strip())

    _docx_header_block(doc, company)
    _docx_add_spacer(doc, 8)

    def _body_bold(text, size=10):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    def _body(text, size=10, justify=True):
        p = doc.add_paragraph()
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    discom_code = client.get("discom_code") or client.get("discom") or "MSEDCL"
    licensee_title = client.get("distribution_licensee") or "Additional Executive Engineer"
    city_line = f"{discom_code} Meter Lab {city}." if city else f"{discom_code} Meter Lab."
    _body_bold(f"To,\n{licensee_title}\n{city_line}" + (f"\n{pincode}" if pincode else ""))
    _docx_add_spacer(doc, 6)
    _body_bold("Sub: Request for Gen-meter Letter.")
    _docx_add_spacer(doc, 4)
    _body("Dear Sir,")
    _docx_add_spacer(doc, 2)
    _body("I hope this letter finds you well. I am writing to request meter testing services for my solar photovoltaic (PV) system. "
          "As a responsible solar PV system owner, I understand the importance of accurate and reliable meter readings to ensure the "
          "system's optimal performance and compliance with regulatory standards.")
    _body(f"Customer Name: {client_name}  C NO. {consumer_num}  I currently have a solar PV system installed at the following "
          f"location: {location_str}")
    _body("To ensure the system's efficiency and adherence to industry standards, I am seeking a comprehensive meter testing service "
          "for the following meters within the system:")

    meter_tbl = doc.add_table(rows=2, cols=2)
    meter_tbl.style = 'Table Grid'
    _remove_tbl_borders(meter_tbl)
    meter_data = [
        (f"Generation Meter - Make- {gen_make}", f"SERIAL NO- {gen_serial}"),
        (f"NET METER - MAKE - {net_make}", f"SERIAL NO - {net_serial}"),
    ]
    for r_idx, (lt, rt) in enumerate(meter_data):
        row = meter_tbl.rows[r_idx]
        row.cells[0].width = _emu(9.3 * 360000)
        row.cells[1].width = _emu(9.3 * 360000)
        lp = row.cells[0].paragraphs[0]
        lp.add_run(lt).font.size = Pt(10)
        rp = row.cells[1].paragraphs[0]
        rp.add_run(rt).font.size = Pt(10)

    _docx_add_spacer(doc, 6)
    _body("I kindly request that the meter testing service be conducted by a certified and accredited organization, ensuring accurate "
          "and unbiased results. The testing should include a thorough assessment of the meters' functionality, calibration, and accuracy, "
          "as well as verification of their compliance with relevant industry standards and regulations.")
    _body("Thank you for your attention to this matter. I look forward to receiving your response and arranging the necessary meter "
          "testing for my solar PV system.")
    _docx_add_spacer(doc, 8)
    _body("Thanks & Regards,")
    _docx_add_spacer(doc, 12)
    _body_bold(company_name.upper())
    owner_name = (company.get('owner_name') or company.get('proprietor_name') or company.get('authorized_signatory') or company.get('manager_name') or '').strip()
    if owner_name:
        _body_bold(owner_name)
    _docx_add_spacer(doc, 16)
    encl_p = doc.add_paragraph()
    encl_run = encl_p.add_run("Encl:\n1. Gen-meter\n2. Test report of meter\n3. Electricity Bill\n4. Solar PV System Approval Latter Copy.")
    encl_run.font.size = Pt(10)
    encl_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_docx(doc_type: str, client: dict, company: dict) -> bytes:
    """
    Generate a Word (.docx) document for the given doc_type.
    Each document type has its own faithful DOCX generator that mirrors the PDF layout.
    Does NOT modify any existing PDF generator.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc_type_clean = (doc_type or "").lower().strip()

    # Annexure: use DOCX template flow (unchanged)
    if doc_type_clean == "annexure":
        try:
            import annexure_generator
            template_bytes = annexure_generator._load_template_bytes()
            replacements = annexure_generator.resolve_annexure_values(client, company)
            filled_docx = annexure_generator._fill_template(template_bytes, replacements)
            return filled_docx
        except Exception as _e:
            import logging as _logging
            _logging.getLogger(__name__).warning(f"Annexure DOCX template path failed: {_e}.")

    # Per-document faithful DOCX generators
    if doc_type_clean == "wcr":
        return _generate_wcr_docx(client, company)
    if doc_type_clean == "sldr":
        return _generate_sldr_docx(client, company)
    if doc_type_clean == "net_meter_agreement":
        return _generate_net_meter_docx(client, company)
    if doc_type_clean in ("vendor_agreement", "vendor"):
        return _generate_vendor_docx(client, company)
    if doc_type_clean in ("meter_testing_request", "meter_testing"):
        return _generate_meter_testing_docx(client, company)
    if doc_type_clean in ("tax_invoice", "proforma", "payment_receipt", "credit_note", "debit_note", "invoice"):
        return generate_invoice_docx(client, company)
    if doc_type_clean in ("purchase_order", "po"):
        return generate_po_docx(client, company)

    # Generic fallback for unknown types
    doc = _build_docx_document()
    _docx_header_block(doc, company)
    title_map = {
        "wcr": "WORK COMPLETION REPORT (WCR)",
        "sldr": "SINGLE LINE DIAGRAM REPORT (SLDR)",
        "net_meter_agreement": "NET METER AGREEMENT",
        "vendor_agreement": "VENDOR AGREEMENT",
        "meter_testing_request": "METER TESTING REQUEST",
        "meter_testing": "METER TESTING REQUEST",
        "annexure": "ANNEXURE - Material & Site Details",
    }
    title = title_map.get(doc_type_clean, doc_type_clean.upper())
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = datetime.now(timezone.utc).strftime("%d %b %Y")
    date_para = doc.add_paragraph(f"Date: {date_str}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    client_name = client.get("full_name") or client.get("name") or ""
    address_parts = [client.get("address") or "", client.get("city") or "",
                     client.get("state") or "", client.get("pincode") or ""]
    full_address = ", ".join(p for p in address_parts if p)
    pan_num = str(client.get("pan_number") or client.get("pan_card_number") or "").strip()
    aadhaar_num = str(client.get("aadhaar") or client.get("aadhaar_number") or "").strip()
    id_label = "PAN Card" if pan_num else "Aadhaar (last 4)"
    id_val = pan_num if pan_num else (aadhaar_num[-4:] if aadhaar_num else "-")
    _docx_kv_table(doc, [
        ["Client Name", client_name],
        ["Mobile", client.get("mobile") or "-"],
        ["Consumer Number", client.get("consumer_number") or "-"],
        ["Address", full_address or "-"],
        [id_label, id_val],
    ])
    doc.add_paragraph()
    doc.add_paragraph()
    sig_tbl = doc.add_table(rows=2, cols=2)
    sig_tbl.rows[0].cells[0].text = "Consumer Signature"
    sig_tbl.rows[0].cells[1].text = f"Vendor Signature ({company.get('company_name') or ''})"
    sig_tbl.rows[1].cells[0].text = "\n\n_________________________"
    sig_tbl.rows[1].cells[1].text = "\n\n_________________________"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_po_docx(data: dict, company: dict) -> bytes:
    """Generate faithful Word (.docx) document for Purchase Orders matching PDF layout."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    norm = _normalize_po_document_data(data, company)
    comp = norm["company"]
    po = norm["poDetails"]
    vendor = norm["vendor"]
    ship_to = norm["shipTo"]
    shipping = norm["shipping"]
    items = norm["lineItems"]
    fin = norm["financials"]
    notes = norm["notes"]

    doc = _build_docx_document(left_cm=1.2, right_cm=1.2, top_cm=1.0, bottom_cm=1.5)

    # 1. HEADER (Logo/Company Left, Title/Meta Right)
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.style = 'Table Grid'
    _remove_tbl_borders(hdr_table)

    cell_l = hdr_table.rows[0].cells[0]
    cell_r = hdr_table.rows[0].cells[1]

    logo_bytes = comp.get("logo_bytes")
    if logo_bytes:
        try:
            p_logo = cell_l.paragraphs[0]
            p_logo.add_run().add_picture(BytesIO(logo_bytes), width=Inches(1.8))
        except Exception:
            pass

    p_comp = cell_l.add_paragraph() if logo_bytes else cell_l.paragraphs[0]
    r_comp = p_comp.add_run(comp["name"])
    r_comp.bold = True
    r_comp.font.size = Pt(13)
    r_comp.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)

    if comp["tagline"]:
        p_tag = cell_l.add_paragraph()
        r_tag = p_tag.add_run(comp["tagline"])
        r_tag.font.size = Pt(8.5)
        r_tag.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    if comp["gstin"]:
        p_gst = cell_l.add_paragraph()
        r_gst = p_gst.add_run(f"GST: {comp['gstin']}")
        r_gst.bold = True
        r_gst.font.size = Pt(9)

    p_title = cell_r.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_title = p_title.add_run("PURCHASE ORDER\n")
    r_title.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)

    r_meta = p_title.add_run(f"DATE: {po['po_date']}\nP.O. NUMBER: {po['po_number']}\nVENDOR ID: {po['vendor_id'] or '—'}")
    r_meta.font.size = Pt(9)
    r_meta.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph()

    # 2. TWO-COLUMN PARTY SECTION (VENDOR | SHIP TO)
    party_tbl = doc.add_table(rows=2, cols=2)
    party_tbl.style = 'Table Grid'

    c_vh = party_tbl.rows[0].cells[0]
    c_sh = party_tbl.rows[0].cells[1]
    c_vh.text = "VENDOR"
    c_sh.text = "SHIP TO"
    _docx_set_cell_bg(c_vh, "1e3a8a")
    _docx_set_cell_bg(c_sh, "1e3a8a")
    for c in (c_vh, c_sh):
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    c_vb = party_tbl.rows[1].cells[0]
    c_sb = party_tbl.rows[1].cells[1]

    p_v = c_vb.paragraphs[0]
    rv_name = p_v.add_run(f"{vendor['name']}\n")
    rv_name.bold = True
    rv_name.font.size = Pt(9.5)
    v_details = []
    if vendor['address']: v_details.append(vendor['address'])
    if vendor['phone']: v_details.append(f"Phone: {vendor['phone']}")
    if vendor['email']: v_details.append(f"Email: {vendor['email']}")
    if vendor['gstin']: v_details.append(f"GSTIN: {vendor['gstin']}")
    rv_body = p_v.add_run("\n".join(v_details))
    rv_body.font.size = Pt(8.5)

    p_s = c_sb.paragraphs[0]
    rs_name = p_s.add_run(f"{ship_to['name']}\n")
    rs_name.bold = True
    rs_name.font.size = Pt(9.5)
    s_details = []
    if ship_to['address']: s_details.append(ship_to['address'])
    if ship_to['phone']: s_details.append(f"Phone: {ship_to['phone']}")
    if ship_to['email']: s_details.append(f"Email: {ship_to['email']}")
    if ship_to['gstin']: s_details.append(f"GSTIN: {ship_to['gstin']}")
    rs_body = p_s.add_run("\n".join(s_details))
    rs_body.font.size = Pt(8.5)

    doc.add_paragraph()

    # 3. SHIPPING BAR TABLE (4 COLUMNS)
    ship_tbl = doc.add_table(rows=2, cols=4)
    ship_tbl.style = 'Table Grid'
    s_headers = ["SHIP VIA", "SHIPPING METHOD", "SHIPPING TERM", "DELIVERY DATE"]
    s_vals = [shipping["ship_via"] or "FOR", shipping["shipping_method"] or "PAID", shipping["shipping_term"] or "DOOR DELIVERY", shipping["delivery_date"] or "—"]

    for i, h in enumerate(s_headers):
        cell = ship_tbl.rows[0].cells[i]
        cell.text = h
        _docx_set_cell_bg(cell, "1e3a8a")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    for i, v in enumerate(s_vals):
        cell = ship_tbl.rows[1].cells[i]
        cell.text = v
        _docx_set_cell_bg(cell, "f8fafc")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)

    doc.add_paragraph()

    # 4. LINE ITEMS TABLE
    items_tbl = doc.add_table(rows=len(items) + 1, cols=6)
    items_tbl.style = 'Table Grid'

    i_headers = ["CODE", "PRODUCT NAME / DESCRIPTION", "QTY", "UNIT", "UNIT PRICE (Rs.)", "TOTAL (Rs.)"]
    for i, h in enumerate(i_headers):
        cell = items_tbl.rows[0].cells[i]
        cell.text = h
        _docx_set_cell_bg(cell, "1e3a8a")
        for p in cell.paragraphs:
            if i in (0, 2, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i in (4, 5):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    for idx, item in enumerate(items, 1):
        row_cells = items_tbl.rows[idx].cells
        p_name = item["product_name"]
        if item["size"] and item["size"] not in p_name:
            p_name += f" ({item['size']})"
        qty = item["quantity"]
        unit = item["unit"]
        price = item["unit_price"]
        amount = item["amount"]

        row_cells[0].text = item["code"]
        row_cells[1].text = p_name
        row_cells[2].text = f"{qty:g}"
        row_cells[3].text = unit
        row_cells[4].text = f"{price:,.2f}"
        row_cells[5].text = f"{amount:,.2f}"

        for i, c in enumerate(row_cells):
            for p in c.paragraphs:
                if i in (0, 2, 3):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif i in (4, 5):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.size = Pt(8.5)

    doc.add_paragraph()

    # 5. BOTTOM SECTION (NOTES & TOTALS)
    bot_tbl = doc.add_table(rows=1, cols=2)
    bot_tbl.style = 'Table Grid'
    c_n = bot_tbl.rows[0].cells[0]
    c_t = bot_tbl.rows[0].cells[1]

    p_n = c_n.paragraphs[0]
    rn_hdr = p_n.add_run("NOTES AND INSTRUCTION\n\n")
    rn_hdr.bold = True
    rn_hdr.font.size = Pt(9)
    rn_bdy = p_n.add_run(notes)
    rn_bdy.font.size = Pt(8)

    totals_list = [
        ("SUBTOTAL", f"Rs. {fin['subtotal']:,.2f}")
    ]
    if fin["discount"] > 0:
        totals_list.append(("DISCOUNT", f"Rs. {fin['discount']:,.2f}"))
    if fin["cgst_amount"] > 0 or fin["cgst_rate"] > 0:
        totals_list.append((f"CGST ({fin['cgst_rate']}%)", f"Rs. {fin['cgst_amount']:,.2f}"))
    if fin["sgst_amount"] > 0 or fin["sgst_rate"] > 0:
        totals_list.append((f"SGST ({fin['sgst_rate']}%)", f"Rs. {fin['sgst_amount']:,.2f}"))
    if fin["igst_amount"] > 0 or fin["igst_rate"] > 0:
        totals_list.append((f"IGST ({fin['igst_rate']}%)", f"Rs. {fin['igst_amount']:,.2f}"))
    if fin["freight"] > 0:
        totals_list.append(("S&H / FREIGHT", f"Rs. {fin['freight']:,.2f}"))
    totals_list.append(("GRAND TOTAL", f"Rs. {fin['grand_total']:,.2f}"))

    for lbl, val in totals_list:
        p_row = c_t.add_paragraph()
        r_lbl = p_row.add_run(f"{lbl}: ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(8.5)
        r_val = p_row.add_run(val)
        r_val.font.size = Pt(8.5)
        if lbl == "GRAND TOTAL":
            r_lbl.font.size = Pt(9.5)
            r_lbl.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
            r_val.bold = True
            r_val.font.size = Pt(10)
            r_val.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)

    doc.add_paragraph()

    # Amount in Words
    p_w_hdr = doc.add_paragraph()
    r_w_hdr = p_w_hdr.add_run("Amount in Words:")
    r_w_hdr.bold = True
    r_w_hdr.font.size = Pt(9)
    r_w_hdr.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    p_w_body = doc.add_paragraph()
    r_w_body = p_w_body.add_run(_amount_to_words(fin.get("grand_total") or 0))
    r_w_body.font.size = Pt(8.5)

    doc.add_paragraph()

    # Authorized Signature Table
    sig_tbl = doc.add_table(rows=2, cols=2)
    sig_tbl.style = 'Table Grid'
    _remove_tbl_borders(sig_tbl)
    sig_tbl.rows[0].cells[0].text = "Vendor Acceptance Signature"
    sig_tbl.rows[0].cells[1].text = f"For {comp['name']}"
    sig_tbl.rows[1].cells[0].text = "\n\n_________________________"
    sig_tbl.rows[1].cells[1].text = "\n\n_________________________\nAuthorized Signatory"
    for r in sig_tbl.rows:
        r.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # 6. FOOTER & NOT TAX INVOICE
    p_ftr = doc.add_paragraph()
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ftr = p_ftr.add_run(f"Enquiries: {comp['email']} | Contact: {comp['phone']} | Office: {comp['address']}")
    r_ftr.font.size = Pt(8)
    r_ftr.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    p_not_tax = doc.add_paragraph()
    p_not_tax.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_not_tax = p_not_tax.add_run("THIS IS NOT TAX INVOICE !")
    r_not_tax.bold = True
    r_not_tax.font.size = Pt(10)
    r_not_tax.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

