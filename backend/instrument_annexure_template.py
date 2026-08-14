"""
Instrument the Annexure-I DOCX template by replacing live values with {{ placeholder }} tags.

Run once:
    python instrument_annexure_template.py

Produces: annexure_template.docx  (committed alongside server code)
"""
import copy
import re
from docx import Document
from docx.oxml.ns import qn
import io

INPUT_DOCX = "../02_Annexure_I-1.docx"
OUTPUT_DOCX = "annexure_template.docx"

# Map: (table_index, row_index, col_index_of_value_cell) → placeholder text
# Col 2 (index 2) is the "As Commissioned" column for table 0
TABLE_REPLACEMENTS = {
    # (table, row, col): placeholder
    (0, 1, 2): "{{ client_full_name }}",
    (0, 2, 2): "{{ consumer_no }}",
    (0, 3, 2): "{{ mobile_no }}",
    (0, 4, 2): "{{ email_id }}",
    (0, 5, 2): "{{ address }}",
    (0, 6, 2): "{{ net_metering_arrangement }}",
    (0, 7, 2): "{{ re_source }}",
    # Row 8 is blank (number 8, particulars blank) — leave as is
    (0, 9, 2): "{{ capacity_type }}",
    (0, 10, 2): "{{ project_model }}",
    (0, 11, 2): "{{ panel_in_kw_kw }}",
    (0, 12, 2): "NA",
    (0, 13, 2): "NA",
    (0, 14, 2): "{{ installation_date }}",
    (0, 15, 2): "{{ solar_pv_details }}",
    (0, 16, 2): "{{ inverter_in_kw_kw }}",
    (0, 17, 2): "{{ inverter_brand }}",
    (0, 18, 2): "{{ solar_panels_in_nos_nos }}",
    (0, 19, 2): "{{ panel_in_wp_wp }}",
}

# Paragraph-level replacements (replace matching substrings)
# The last paragraph: "Date: 29/07/2026 DESAI STOCK SYNDICATE    GVP SOLAR ENERGY,..."
PARAGRAPH_INDEX_17_REPLACEMENT = "Date: {{ installation_date }}     {{ client_full_name }}                    {{ company_details }}"


def _set_cell_text_preserve_format(cell, new_text: str):
    """Clear all paragraphs in a cell and write new_text into first paragraph, preserving run formatting."""
    # Get font info from the first run of the first paragraph (if any)
    first_para = cell.paragraphs[0] if cell.paragraphs else None
    ref_run = None
    if first_para and first_para.runs:
        ref_run = first_para.runs[0]

    # Clear all paragraphs content
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""

    # Write to first paragraph's first run (or add a run)
    if first_para:
        if first_para.runs:
            first_para.runs[0].text = new_text
        else:
            run = first_para.add_run(new_text)
            if ref_run:
                run.bold = ref_run.bold
                run.italic = ref_run.italic
                run.font.size = ref_run.font.size
    else:
        # Fallback: clear via XML
        for p in cell._tc.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    t.text = new_text


def main():
    doc = Document(INPUT_DOCX)

    # 1. Table replacements
    for (ti, ri, ci), placeholder in TABLE_REPLACEMENTS.items():
        cell = doc.tables[ti].rows[ri].cells[ci]
        _set_cell_text_preserve_format(cell, placeholder)

    # 2. Last paragraph (Date + signatures)
    # Find paragraph 17
    para = doc.paragraphs[17]
    # Replace all runs with combined text first
    full = "".join(r.text for r in para.runs)
    if para.runs:
        para.runs[0].text = PARAGRAPH_INDEX_17_REPLACEMENT
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(PARAGRAPH_INDEX_17_REPLACEMENT)

    doc.save(OUTPUT_DOCX)
    print(f"✓ Wrote {OUTPUT_DOCX}")

    # Verify: read back and show all placeholders
    doc2 = Document(OUTPUT_DOCX)
    import re
    found = set()
    for p in doc2.paragraphs:
        for m in re.finditer(r"\{\{([^}]+)\}\}", p.text):
            found.add(m.group(0).strip())
    for tbl in doc2.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for m in re.finditer(r"\{\{([^}]+)\}\}", p.text):
                        found.add(m.group(0).strip())
    print(f"Placeholders found in output: {sorted(found)}")


if __name__ == "__main__":
    main()
