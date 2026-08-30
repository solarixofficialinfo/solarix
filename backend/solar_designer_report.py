from __future__ import annotations
import os
import io
import re
import base64
import logging
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, Image as RLImage, HRFlowable
)
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)

# Register PDF-safe Unicode font supporting the Indian Rupee symbol
FONTS_DIR = os.path.join(os.path.dirname(__file__), "fonts")
REGULAR_FONT_PATH = os.path.join(FONTS_DIR, "NotoSans-Regular.ttf")
BOLD_FONT_PATH = os.path.join(FONTS_DIR, "NotoSans-Bold.ttf")

PDF_FONT = "Helvetica"
PDF_FONT_BOLD = "Helvetica-Bold"

try:
    if os.path.exists(REGULAR_FONT_PATH) and os.path.exists(BOLD_FONT_PATH):
        pdfmetrics.registerFont(TTFont("NotoSans", REGULAR_FONT_PATH))
        pdfmetrics.registerFont(TTFont("NotoSans-Bold", BOLD_FONT_PATH))
        pdfmetrics.registerFontFamily("NotoSans", normal="NotoSans", bold="NotoSans-Bold")
        PDF_FONT = "NotoSans"
        PDF_FONT_BOLD = "NotoSans-Bold"
    elif os.path.exists("/Library/Fonts/Arial Unicode.ttf"):
        pdfmetrics.registerFont(TTFont("NotoSans", "/Library/Fonts/Arial Unicode.ttf"))
        pdfmetrics.registerFont(TTFont("NotoSans-Bold", "/Library/Fonts/Arial Unicode.ttf"))
        pdfmetrics.registerFontFamily("NotoSans", normal="NotoSans", bold="NotoSans-Bold")
        PDF_FONT = "NotoSans"
        PDF_FONT_BOLD = "NotoSans-Bold"
except Exception as e:
    logger.warning(f"Font registration warning: {e}")

styles = getSampleStyleSheet()

TITLE_STYLE = ParagraphStyle('DocTitle', parent=styles['Heading1'], fontSize=16, fontName=PDF_FONT_BOLD, textColor=colors.HexColor('#0f172a'), spaceAfter=4, leading=20)
SUBTITLE_STYLE = ParagraphStyle('DocSub', parent=styles['Normal'], fontSize=9, fontName=PDF_FONT, textColor=colors.HexColor('#64748b'), spaceAfter=8, leading=12)
SECTION_HEADING = ParagraphStyle('SecHead', parent=styles['Heading2'], fontSize=11, fontName=PDF_FONT_BOLD, textColor=colors.HexColor('#1e40af'), spaceBefore=8, spaceAfter=4, leading=14)
BODY_STYLE = ParagraphStyle('Body', parent=styles['BodyText'], fontSize=8.5, fontName=PDF_FONT, textColor=colors.HexColor('#334155'), leading=12)
BODY_BOLD = ParagraphStyle('BodyBold', parent=styles['BodyText'], fontSize=8.5, fontName=PDF_FONT_BOLD, textColor=colors.HexColor('#0f172a'), leading=12)
SMALL_MUTED = ParagraphStyle('SmallMuted', parent=styles['Normal'], fontSize=7.5, fontName=PDF_FONT, textColor=colors.HexColor('#64748b'), leading=10)
DISCLAIMER_STYLE = ParagraphStyle('Disclaimer', parent=styles['Normal'], fontSize=7.5, fontName=PDF_FONT, textColor=colors.HexColor('#991b1b'), leading=10)
KPI_NUM_STYLE = ParagraphStyle('KPINum', parent=styles['Normal'], fontSize=13, fontName=PDF_FONT_BOLD, textColor=colors.HexColor('#1e40af'), alignment=1, leading=16)
KPI_LBL_STYLE = ParagraphStyle('KPILbl', parent=styles['Normal'], fontSize=7.5, fontName=PDF_FONT, textColor=colors.HexColor('#475569'), alignment=1, leading=10)

class SolarReportCanvas(canvas.Canvas):
    """Canvas with header, footer, page count, and brand lines."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.pages = []

    def showPage(self):
        self.pages.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self.pages)
        for page in self.pages:
            self.__dict__.update(page)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count: int):
        self.saveState()
        page_w, page_h = A4
        
        # Header Accent Line
        self.setStrokeColor(colors.HexColor('#2563eb'))
        self.setLineWidth(2)
        self.line(1.5 * cm, page_h - 1.2 * cm, page_w - 1.5 * cm, page_h - 1.2 * cm)
        
        # Footer
        self.setStrokeColor(colors.HexColor('#e2e8f0'))
        self.setLineWidth(0.75)
        self.line(1.5 * cm, 1.4 * cm, page_w - 1.5 * cm, 1.4 * cm)
        
        self.setFont(PDF_FONT, 7.5)
        self.setFillColor(colors.HexColor('#64748b'))
        self.drawString(1.5 * cm, 1.0 * cm, "SOLARIX 3D Solar Designer — Preliminary Technical Report")
        
        pg_str = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(page_w - 1.5 * cm, 1.0 * cm, pg_str)
        self.restoreState()


def _decode_base64_image(data_uri_or_b64: str) -> Optional[io.BytesIO]:
    """Safely decode data:image/*;base64,... into BytesIO."""
    if not data_uri_or_b64 or not isinstance(data_uri_or_b64, str):
        return None
    try:
        if "base64," in data_uri_or_b64:
            b64_str = data_uri_or_b64.split("base64,")[1]
        else:
            b64_str = data_uri_or_b64
        img_bytes = base64.b64decode(b64_str.strip())
        return io.BytesIO(img_bytes)
    except Exception as e:
        logger.warning(f"Error decoding base64 image: {e}")
        return None


def generate_solar_design_pdf(design: Dict[str, Any], company: Dict[str, Any]) -> bytes:
    """Generate complete multi-page PDF technical solar rooftop report."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
    )

    story = []
    page_w, _ = A4
    content_w = page_w - 3.0 * cm

    # 1. Company & Header Section
    company_name = company.get('company_name') or company.get('name') or 'Solar Energy Solutions'
    owner_name = company.get('owner_name') or company.get('proprietor_name') or ''
    mobile = company.get('mobile') or company.get('phone') or ''
    email = company.get('email') or ''
    gst = company.get('gst_number') or company.get('gstin') or ''

    header_left = [
        Paragraph(f"<b>{company_name.upper()}</b>", BODY_BOLD),
        Paragraph(f"{owner_name} {('· ' + mobile) if mobile else ''}", BODY_STYLE),
        Paragraph(f"{email} {('· GSTIN: ' + gst) if gst else ''}", SMALL_MUTED),
    ]

    design_title = design.get('site_name') or design.get('name') or "Solar Rooftop Layout Design"
    created_date = design.get('created_at', datetime.now(timezone.utc).isoformat())[:10]
    version_tag = f"v{design.get('version', 1)}"

    header_right = [
        Paragraph("<b>SOLAR ROOFTOP DESIGN REPORT</b>", ParagraphStyle('HR1', parent=TITLE_STYLE, fontSize=12, alignment=2, textColor=colors.HexColor('#1e40af'))),
        Paragraph(f"Design Ref: <b>{design.get('design_number') or str(design.get('id', 'SD-001'))[:10]}</b> ({version_tag})", ParagraphStyle('HR2', parent=BODY_STYLE, alignment=2)),
        Paragraph(f"Date: <b>{created_date}</b>", ParagraphStyle('HR3', parent=BODY_STYLE, alignment=2)),
    ]

    header_table = Table(
        [[header_left, header_right]],
        colWidths=[content_w * 0.55, content_w * 0.45]
    )
    header_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('PADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(header_table)
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cbd5e1'), spaceBefore=2, spaceAfter=8))

    # 2. Site & Client Location Banner
    client_name = design.get('client_name') or "Direct Client"
    formatted_address = design.get('formatted_address') or design.get('address') or "Site address not specified"
    lat = design.get('latitude')
    lng = design.get('longitude')
    coords_str = f"{lat:.6f}, {lng:.6f}" if (lat is not None and lng is not None) else "Coordinates not captured"

    site_info_data = [
        [
            Paragraph("<b>Client / Project:</b>", BODY_STYLE),
            Paragraph(f"<b>{client_name}</b>", BODY_BOLD),
            Paragraph("<b>Site Name:</b>", BODY_STYLE),
            Paragraph(f"{design_title}", BODY_BOLD)
        ],
        [
            Paragraph("<b>Site Address:</b>", BODY_STYLE),
            Paragraph(f"{formatted_address}", BODY_STYLE),
            Paragraph("<b>Coordinates:</b>", BODY_STYLE),
            Paragraph(f"{coords_str}", BODY_STYLE)
        ]
    ]
    site_table = Table(site_info_data, colWidths=[2.5 * cm, 6.0 * cm, 2.5 * cm, content_w - (11.0 * cm)])
    site_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
        ('BOX', (0,0), (-1,-1), 0.75, colors.HexColor('#e2e8f0')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#f1f5f9')),
        ('PADDING', (0,0), (-1,-1), 4),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(site_table)
    story.append(Spacer(1, 8))

    # 3. High-Level KPI Summary Cards
    panel_count = int(design.get('panel_count') or len(design.get('panels', [])) or 0)
    panel_wattage = float(design.get('panel_wattage') or 550)
    system_kw = float(design.get('system_kw') or (panel_count * panel_wattage / 1000.0))
    roof_area = float(design.get('roof_area_sqm') or design.get('roof_area') or 0)
    usable_area = float(design.get('usable_area_sqm') or design.get('usable_area') or roof_area * 0.85)
    coverage_pct = float(design.get('coverage_pct') or ((panel_count * 2.3 / max(usable_area, 1.0)) * 100.0 if usable_area > 0 else 0))
    coverage_pct = min(100.0, max(0.0, coverage_pct))

    kpi_cards = [
        [
            [Paragraph(f"{system_kw:.2f} kWp", KPI_NUM_STYLE), Paragraph("DC SYSTEM CAPACITY", KPI_LBL_STYLE)],
            [Paragraph(f"{panel_count} Nos", KPI_NUM_STYLE), Paragraph(f"PANELS ({int(panel_wattage)}W)", KPI_LBL_STYLE)],
            [Paragraph(f"{roof_area:.1f} m²", KPI_NUM_STYLE), Paragraph("ROOF BOUNDARY", KPI_LBL_STYLE)],
            [Paragraph(f"{usable_area:.1f} m²", KPI_NUM_STYLE), Paragraph("USABLE ROOFTOP", KPI_LBL_STYLE)],
            [Paragraph(f"{coverage_pct:.1f}%", KPI_NUM_STYLE), Paragraph("ROOF COVERAGE", KPI_LBL_STYLE)],
        ]
    ]
    kpi_table = Table(kpi_cards, colWidths=[content_w / 5.0] * 5)
    kpi_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#eff6ff')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#bfdbfe')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#dbeafe')),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(kpi_table)
    story.append(Spacer(1, 10))

    # 4. Technical System Specifications Table
    story.append(Paragraph("1. Technical & Engineering Specifications", SECTION_HEADING))
    panel_make = design.get('panel_make') or design.get('panel_model') or "High-Efficiency Mono PERC / TOPCon"
    orientation = str(design.get('orientation', 'Portrait')).title()
    tilt_deg = design.get('tilt_angle') or design.get('tilt', 15)
    azimuth_deg = design.get('azimuth_angle') or design.get('azimuth', 180)
    structure_type = str(design.get('structure_type', 'Elevated Mounting')).title()
    mounting_height = design.get('mounting_height_m') or design.get('mounting_height', 1.8)
    setback_m = design.get('setback_m', 0.5)
    walkway_m = design.get('walkway_m', 0.6)

    spec_rows = [
        [
            Paragraph("<b>Solar Module Make / Model</b>", BODY_STYLE),
            Paragraph(f"{panel_make}", BODY_BOLD),
            Paragraph("<b>Module Rating & Tech</b>", BODY_STYLE),
            Paragraph(f"{int(panel_wattage)} Wp · Tier-1 Mono / Bi-facial", BODY_BOLD)
        ],
        [
            Paragraph("<b>Panel Layout Orientation</b>", BODY_STYLE),
            Paragraph(f"{orientation}", BODY_STYLE),
            Paragraph("<b>Installation Tilt Angle</b>", BODY_STYLE),
            Paragraph(f"{tilt_deg}° (Fixed South Optimized)", BODY_STYLE)
        ],
        [
            Paragraph("<b>Azimuth / Orientation</b>", BODY_STYLE),
            Paragraph(f"{azimuth_deg}° (True South)", BODY_STYLE),
            Paragraph("<b>Mounting Structure Type</b>", BODY_STYLE),
            Paragraph(f"{structure_type} ({mounting_height}m clearance)", BODY_STYLE)
        ],
        [
            Paragraph("<b>Perimeter Setback Clearance</b>", BODY_STYLE),
            Paragraph(f"{setback_m} meters", BODY_STYLE),
            Paragraph("<b>Maintenance Walkway Width</b>", BODY_STYLE),
            Paragraph(f"{walkway_m} meters", BODY_STYLE)
        ],
        [
            Paragraph("<b>Total Area / Usable Area</b>", BODY_STYLE),
            Paragraph(f"{roof_area:.1f} m² / {usable_area:.1f} m²", BODY_STYLE),
            Paragraph("<b>Estimated Annual Generation</b>", BODY_STYLE),
            Paragraph(f"~{(system_kw * 1450):,.0f} kWh / Year", BODY_BOLD)
        ]
    ]

    specs_table = Table(spec_rows, colWidths=[4.2 * cm, content_w * 0.5 - (4.2 * cm), 4.2 * cm, content_w * 0.5 - (4.2 * cm)])
    specs_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#ffffff')),
        ('BOX', (0,0), (-1,-1), 0.75, colors.HexColor('#cbd5e1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
        ('PADDING', (0,0), (-1,-1), 4),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(specs_table)
    story.append(Spacer(1, 10))

    # 5. Visual Layout Snapshots (2D Plan & 3D Rooftop View)
    snap_2d_b64 = design.get('layout_snapshot_2d') or design.get('snapshot_2d')
    snap_3d_b64 = design.get('layout_snapshot_3d') or design.get('snapshot_3d')
    sat_b64 = design.get('satellite_snapshot')

    img_2d_stream = _decode_base64_image(snap_2d_b64)
    img_3d_stream = _decode_base64_image(snap_3d_b64)
    sat_stream = _decode_base64_image(sat_b64)

    if img_2d_stream or img_3d_stream or sat_stream:
        story.append(Paragraph("2. Rooftop Layout & 3D Visualization", SECTION_HEADING))
        
        if img_2d_stream and img_3d_stream:
            rl_2d = RLImage(img_2d_stream, width=content_w * 0.48, height=6.0 * cm)
            rl_3d = RLImage(img_3d_stream, width=content_w * 0.48, height=6.0 * cm)
            img_cells = [[
                [Paragraph("<b>2D Roof Plan & Panel Grid</b>", SMALL_MUTED), Spacer(1, 2), rl_2d],
                [Paragraph("<b>3D Rooftop Simulation View</b>", SMALL_MUTED), Spacer(1, 2), rl_3d]
            ]]
            img_table = Table(img_cells, colWidths=[content_w * 0.5, content_w * 0.5])
            img_table.setStyle(TableStyle([
                ('PADDING', (0,0), (-1,-1), 2),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
            ]))
            story.append(img_table)
        elif img_3d_stream:
            rl_3d = RLImage(img_3d_stream, width=content_w, height=7.5 * cm)
            story.append(rl_3d)
        elif img_2d_stream:
            rl_2d = RLImage(img_2d_stream, width=content_w, height=7.5 * cm)
            story.append(rl_2d)
        story.append(Spacer(1, 10))

    # 6. Preliminary Bill of Materials (BOM) & Structural Estimate
    story.append(Paragraph("3. Preliminary Material & Component Estimate", SECTION_HEADING))

    # Calculate preliminary engineering quantities based on panel count & structure
    rails_len_m = round(panel_count * 2.3 * 1.05, 1)
    mid_clamps = max(0, (panel_count - 2) * 2) if panel_count > 2 else panel_count * 2
    end_clamps = 4 * max(1, (panel_count // 10) + 1)
    fasteners = (mid_clamps + end_clamps) * 2
    structure_sets = max(1, int(round(panel_count / 4.0)))
    walkway_panels = max(1, int(round(roof_area * 0.08)))
    dc_cable_m = int(round(panel_count * 4.5 + 20))

    bom_rows = [
        [
            Paragraph("<b>#</b>", BODY_BOLD),
            Paragraph("<b>Component / Material Description</b>", BODY_BOLD),
            Paragraph("<b>Specification / Standard</b>", BODY_BOLD),
            Paragraph("<b>Estimated Qty</b>", BODY_BOLD),
            Paragraph("<b>Unit</b>", BODY_BOLD)
        ],
        [
            Paragraph("1", BODY_STYLE),
            Paragraph(f"Solar PV Modules ({panel_make})", BODY_BOLD),
            Paragraph(f"{int(panel_wattage)}W Mono PERC / Bi-facial", BODY_STYLE),
            Paragraph(f"<b>{panel_count}</b>", BODY_BOLD),
            Paragraph("Nos", BODY_STYLE)
        ],
        [
            Paragraph("2", BODY_STYLE),
            Paragraph("Aluminium Mounting Rails / Purlins", BODY_STYLE),
            Paragraph("Anodized Al 6063-T6 / HDGI Channel", BODY_STYLE),
            Paragraph(f"{rails_len_m}", BODY_STYLE),
            Paragraph("Mtrs", BODY_STYLE)
        ],
        [
            Paragraph("3", BODY_STYLE),
            Paragraph(f"Mounting Structure Framework ({structure_type})", BODY_STYLE),
            Paragraph(f"Elevated Legs / Columns ({mounting_height}m clearance)", BODY_STYLE),
            Paragraph(f"{structure_sets}", BODY_STYLE),
            Paragraph("Sets", BODY_STYLE)
        ],
        [
            Paragraph("4", BODY_STYLE),
            Paragraph("Module Mid Clamps & End Clamps", BODY_STYLE),
            Paragraph("Aluminium Anodized with SS304 Allen Bolts", BODY_STYLE),
            Paragraph(f"{mid_clamps + end_clamps} ({mid_clamps}M / {end_clamps}E)", BODY_STYLE),
            Paragraph("Nos", BODY_STYLE)
        ],
        [
            Paragraph("5", BODY_STYLE),
            Paragraph("Anchor Fasteners / Chemical Anchors", BODY_STYLE),
            Paragraph("M10/M12 SS304 Anchor Fasteners", BODY_STYLE),
            Paragraph(f"{fasteners}", BODY_STYLE),
            Paragraph("Nos", BODY_STYLE)
        ],
        [
            Paragraph("6", BODY_STYLE),
            Paragraph("Solar DC Cable (Red & Black)", BODY_STYLE),
            Paragraph("4 / 6 sq.mm Tinned Copper UV Protected", BODY_STYLE),
            Paragraph(f"{dc_cable_m}", BODY_STYLE),
            Paragraph("Mtrs", BODY_STYLE)
        ],
        [
            Paragraph("7", BODY_STYLE),
            Paragraph("Maintenance Walkways & Safety Grating", BODY_STYLE),
            Paragraph("FRP / Galvanized Anti-Slip Grating (0.6m width)", BODY_STYLE),
            Paragraph(f"{walkway_panels}", BODY_STYLE),
            Paragraph("Mtrs", BODY_STYLE)
        ]
    ]

    bom_table = Table(bom_rows, colWidths=[0.8 * cm, 6.0 * cm, 5.8 * cm, 2.8 * cm, content_w - (15.4 * cm)])
    bom_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f1f5f9')),
        ('BOX', (0,0), (-1,-1), 0.75, colors.HexColor('#cbd5e1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
        ('PADDING', (0,0), (-1,-1), 3.5),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f8fafc')])
    ]))
    story.append(bom_table)
    story.append(Spacer(1, 10))

    # 7. Obstacles & Exclusions Summary (if present)
    obstacles = design.get('obstacles', [])
    if obstacles:
        story.append(Paragraph("4. Identified Obstacles & Excluded Zones", SECTION_HEADING))
        obs_rows = [
            [
                Paragraph("<b>Obstacle Name</b>", BODY_BOLD),
                Paragraph("<b>Type / Category</b>", BODY_BOLD),
                Paragraph("<b>Dimensions (L × W × H)</b>", BODY_BOLD),
                Paragraph("<b>Excluded Area</b>", BODY_BOLD)
            ]
        ]
        for obs in obstacles[:8]:
            o_name = obs.get('name') or "Structure"
            o_type = obs.get('type') or "Obstruction"
            l = float(obs.get('length') or obs.get('width_m') or 1.5)
            w = float(obs.get('width') or obs.get('height_m') or 1.5)
            h = float(obs.get('height') or 1.0)
            area = l * w
            obs_rows.append([
                Paragraph(f"{o_name}", BODY_STYLE),
                Paragraph(f"{o_type}", BODY_STYLE),
                Paragraph(f"{l:.1f}m × {w:.1f}m × {h:.1f}m", BODY_STYLE),
                Paragraph(f"{area:.1f} m²", BODY_STYLE),
            ])
        obs_table = Table(obs_rows, colWidths=[content_w * 0.3, content_w * 0.25, content_w * 0.25, content_w * 0.2])
        obs_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#fef2f2')),
            ('BOX', (0,0), (-1,-1), 0.75, colors.HexColor('#fecaca')),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#fee2e2')),
            ('PADDING', (0,0), (-1,-1), 3),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(obs_table)
        story.append(Spacer(1, 10))

    # 8. Prominent Engineering Disclaimer & Signature Block
    story.append(KeepTogether([
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0'), spaceBefore=4, spaceAfter=6),
        Table([
            [
                Paragraph("<b>IMPORTANT ENGINEERING & ACCURACY NOTICE</b>", ParagraphStyle('WN', parent=DISCLAIMER_STYLE, fontName=PDF_FONT_BOLD, fontSize=8)),
            ],
            [
                Paragraph(
                    "This document is a preliminary solar layout and planning estimate prepared via geospatial analysis and computer-assisted modeling. "
                    "Dimensions, roof slopes, and usable areas derived from satellite imagery are approximate unless verified by on-site physical survey. "
                    "This report does NOT constitute a structural load certificate, sanctioned electrical SLD, or construction clearance. "
                    "Physical roof load testing and electrical safety audit must be completed by licensed structural and electrical engineers before installation.",
                    DISCLAIMER_STYLE
                )
            ]
        ], colWidths=[content_w], style=[
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#fff1f2')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#fda4af')),
            ('PADDING', (0,0), (-1,-1), 6),
        ]),
        Spacer(1, 12),
        Table([
            [
                Paragraph(f"<b>Prepared By:</b> {company.get('owner_name') or 'Solar EPC Engineer'}<br/><font color='#64748b'>Authorized Technical Signatory</font>", BODY_STYLE),
                Paragraph("<b>Customer Acknowledgement:</b><br/><font color='#64748b'>Signature / Acceptance Seal</font>", BODY_STYLE),
            ]
        ], colWidths=[content_w * 0.5, content_w * 0.5], style=[
            ('PADDING', (0,0), (-1,-1), 4),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ])
    ]))

    doc.build(story, canvasmaker=SolarReportCanvas)
    return buf.getvalue()


def generate_solar_design_docx(design: Dict[str, Any], company: Dict[str, Any]) -> bytes:
    """Generate Word (.docx) technical report for solar rooftop design."""
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    company_name = company.get('company_name') or company.get('name') or 'Solar Energy Solutions'
    design_title = design.get('site_name') or design.get('name') or "Solar Rooftop Layout Design"
    client_name = design.get('client_name') or "Client"
    formatted_address = design.get('formatted_address') or "Site Address"
    panel_count = int(design.get('panel_count') or len(design.get('panels', [])) or 0)
    panel_wattage = float(design.get('panel_wattage') or 550)
    system_kw = float(design.get('system_kw') or (panel_count * panel_wattage / 1000.0))
    roof_area = float(design.get('roof_area_sqm') or 0)
    usable_area = float(design.get('usable_area_sqm') or roof_area * 0.85)

    p_title = doc.add_paragraph()
    r_company = p_title.add_run(f"{company_name.upper()}\n")
    r_company.bold = True
    r_company.font.size = Pt(14)
    r_company.font.color.rgb = RGBColor(30, 64, 175)

    r_title = p_title.add_run("SOLAR ROOFTOP DESIGN REPORT\n")
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph(f"Site: {design_title} | Client: {client_name} | Ref: {str(design.get('id', 'SD-001'))[:10]}")
    p_sub.runs[0].font.size = Pt(9)
    p_sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    doc.add_heading("1. Site & Key Parameters", level=2)
    t_site = doc.add_table(rows=5, cols=2)
    t_site.style = 'Table Grid'
    site_data = [
        ("Client Name", client_name),
        ("Site Address", formatted_address),
        ("Total DC System Size", f"{system_kw:.2f} kWp ({panel_count} × {int(panel_wattage)}W)"),
        ("Roof Area / Usable Area", f"{roof_area:.1f} m² / {usable_area:.1f} m²"),
        ("Mounting Structure", f"{design.get('structure_type', 'Elevated')} (Tilt: {design.get('tilt_angle', 15)}°)"),
    ]
    for idx, (label, val) in enumerate(site_data):
        row = t_site.rows[idx]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = str(val)

    doc.add_paragraph()
    doc.add_heading("2. Bill of Materials (BOM) Estimate", level=2)
    t_bom = doc.add_table(rows=1, cols=4)
    t_bom.style = 'Table Grid'
    hdr = t_bom.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Description"
    hdr[2].text = "Qty"
    hdr[3].text = "Unit"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True

    bom_items = [
        ("Solar PV Modules", f"{design.get('panel_make', 'Tier-1')} {int(panel_wattage)}W", str(panel_count), "Nos"),
        ("Mounting Structure Rails", "Aluminium 6063-T6 / HDGI Channel", f"{round(panel_count * 2.4, 1)}", "Mtrs"),
        ("Structure Sets / Columns", f"{design.get('structure_type', 'Elevated')}", f"{max(1, int(round(panel_count/4.0)))}", "Sets"),
        ("Module Clamps (Mid/End)", "Aluminium with SS304 Fasteners", f"{panel_count * 2}", "Nos"),
        ("Solar DC Cable", "4/6 sq.mm Tinned Copper UV Cable", f"{int(round(panel_count * 4.5 + 20))}", "Mtrs"),
    ]
    for item in bom_items:
        row = t_bom.add_row().cells
        for col_idx, text in enumerate(item):
            row[col_idx].text = text

    doc.add_paragraph()
    p_disc = doc.add_paragraph()
    r_disc_hdr = p_disc.add_run("PRELIMINARY ENGINEERING DISCLAIMER:\n")
    r_disc_hdr.bold = True
    r_disc_hdr.font.color.rgb = RGBColor(185, 28, 28)
    r_disc_body = p_disc.add_run(
        "This report provides preliminary technical estimation based on geospatial imagery and user inputs. "
        "Physical structural and electrical audits must be conducted prior to installation. "
        "Solarix and the author do not certify civil roof load capacity."
    )
    r_disc_body.font.size = Pt(8.5)
    r_disc_body.font.color.rgb = RGBColor(100, 116, 139)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
